from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import math2docx
import math
from docx.enum.section import WD_SECTION
from copy import deepcopy
import os
from django.conf import settings
from types import SimpleNamespace


def set_heading_style(style, font_name="Times New Roman", size=14, bold=True, align=None,
                      f_l_indent=0.0, space_after=6, space_before=12, color=(0, 0, 0)):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor(*color)

    rFonts = OxmlElement('w:rFonts')
    for tag in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{tag}'), font_name)
    style.element.rPr.append(rFonts)

    fmt = style.paragraph_format
    if align:
        fmt.alignment = align
    fmt.first_line_indent = Cm(f_l_indent)
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)


def add_paragraph_with_indent(doc, text="", indent=3.0, italic=None, sub_text="", sup_text="", value=None, desc=None):
    left = indent
    first_line = -indent
    tab = indent

    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(left)
    fmt.first_line_indent = Cm(first_line)
    fmt.tab_stops.add_tab_stop(Cm(tab))
    if text:
        run = p.add_run(text)
        run.italic = italic
        sub = p.add_run(sub_text)
        if sub_text is not None or sub_text != '':
            sub.font.subscript = True
        sup = p.add_run(sup_text)
        if sup_text is not None or sup_text != '':
            sup.font.superscript = True
        if value is not None:
            p.add_run(f" = {value} – ")
        else:
            p.add_run(" – ")
        if desc:
            p.add_run("\t")
            p.add_run(desc)
    return p


def add_equation(doc, latex: str, number=None, f_l_indent=0.0, font_size_pt=14, min_line_pts=30):
    section = doc.sections[0]
    content_width_emu = section.page_width - section.left_margin - section.right_margin
    right_tab_pos_pt = content_width_emu / 12700.0  # EMU -> pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(f_l_indent)  # без красной строки
    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    fmt.line_spacing = Pt(min_line_pts)  # ключ к отсутствию обрезания
    fmt.space_after = fmt.space_before = Pt(2)

    math2docx.add_math(p, latex)

    if number is not None:
        fmt.tab_stops.add_tab_stop(Pt(right_tab_pos_pt), alignment=WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        p.add_run(f"({number})").font.size = Pt(font_size_pt)

    return p


def add_frame(main_doc, template_name, section_index=0):
    template_path = os.path.join(settings.BASE_DIR, 'templates', 'frames', template_name)
    template_doc = Document(template_path)
    section = main_doc.sections[section_index]
    section.header.is_linked_to_previous = False

    for p in section.header.paragraphs:
        p.clear()

    template_header = template_doc.sections[0].header
    for element in template_header._element:
        section.header._element.append(deepcopy(element))


def generate_report(result):
    d = SimpleNamespace(**result)
    print('d')
    for k in result.keys():
        print(k, "=", getattr(d, k), type(getattr(d, k)))

    f_counter = 0
    doc = Document()
    section = doc.sections[0]

    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin, section.right_margin = Cm(3), Cm(1.25)
    section.top_margin, section.bottom_margin = Cm(1.75), Cm(2.5)
    section.header_distance = section.footer_distance = Cm(0)

    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Times New Roman', Pt(14)

    fmt = style.paragraph_format
    fmt.line_spacing = Pt(21)
    fmt.first_line_indent = Cm(1.25)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_after = fmt.space_before = Pt(0)

    set_heading_style(doc.styles['Heading 1'], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_heading_style(doc.styles['Heading 2'], align=WD_ALIGN_PARAGRAPH.LEFT, f_l_indent=1.25)

    for _ in range(4):
        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn('w:val'), 'none')
        tblBorders.append(border_el)

    tblPr = tbl.xpath("./w:tblPr")[0]
    tblPr.append(tblBorders)
    table.columns[0].width = Cm(9.75)
    table.columns[1].width = Cm(7.25)
    cell = table.cell(0, 1)
    paragraphs = [
        ("УТВЕРЖДАЮ", True),
        ("Генеральный директор", False),
        ("ООО «ИНОТЭК»", False),
        ("____________ Р.Т. Фахретдинов", False),
        ("« ___ » ________________2025 г.", False),
    ]
    for text, bold in paragraphs:
        p = cell.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        run = p.runs[0]
        run.bold = bold
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph("РАСЧЕТ ТОЛЩИН СТЕНОК ПРИВАРНОГО РАЗРЕЗНОГО ТРОЙНИКА", style='Heading 1')
    t = title.add_run()
    t.add_break(WD_BREAK.LINE)
    t.add_text(f"DN{int(d.D_N_pipe)}/{int(d.D_N_b)}-PN{int(d.pressure * 10)}")
    for _ in range(14):
        doc.add_paragraph()
    doc.add_paragraph('Уфа 2025 г.')

    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("СОДЕРЖАНИЕ", style='Heading 1')
    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("1. Исходные данные", style='Heading 2')

    p1 = doc.add_paragraph(
        f"Проводится расчёт тройника Ду {int(d.D_N_pipe)} из стали класса прочности {d.steel_class}, "
        f"устанавливаемого на трубопровод наружным диаметром {d.D_n} мм из стали класса прочности "
        f"{d.steel_class}, ответвление тройника Ду {int(d.D_N_b)} с внутренним диаметром {d.D_vn} мм,"
        f" давление среды в трубопроводе {d.pressure} МПа. Минимальное значение временного сопротивления "
        f"для стали класса прочности {d.steel_class} – σ")
    p1.add_run("В").font.subscript = True
    p1.add_run(f" = {int(d.sigma_b)} МПа, а минимальное значение предела текучести σ")
    p1.add_run("Т").font.subscript = True
    p1.add_run(f" = {int(d.sigma_y)} МПа.")

    doc.add_paragraph("Расчёт проводится согласно методике, приведенной в "
                      "СТО Газпром 2-2.3-116-2016 «Правила производства работ на "
                      "газопроводах врезкой под давлением» [1].")

    doc.add_paragraph("2. Определение толщины стенки трубопровода", style='Heading 2')
    doc.add_paragraph("Толщина стенки трубопровода определяется по СП36.13330.2012 «Магистральные трубопроводы» [2] "
                      "формула (10):")
    f_counter += 1
    add_equation(doc, r"\delta = \frac{n \cdot p \cdot D_н}{2 \cdot (R_1 + n \cdot p)}",
                 f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, "n", italic=True, value=d.n, indent=2.0,
                              desc="коэффициент надежности по нагрузке - внутреннему рабочему давлению в "
                                   "трубопроводе ([2] табл. 14);")
    add_paragraph_with_indent(doc, "p", italic=True, value=d.pressure, indent=2.5,
                              desc="рабочее (нормативное) давление в трубопроводе, МПа;")
    add_paragraph_with_indent(doc, "D", italic=True, sub_text="н", value=d.D_n,
                              desc="наружный диаметр трубопровода, мм;")
    add_paragraph_with_indent(doc, "R", italic=True, sub_text="1", indent=1.25,
                              desc="расчетное сопротивление растяжению (сжатию) ([2] формула (2)):")

    f_counter += 1
    add_equation(doc, r"R_1 = \frac{R_1^н \cdot m}{k_1 \cdot k_н}", f_counter, f_l_indent=7.25)
    doc.add_paragraph("где:", style="Normal")
    add_paragraph_with_indent(doc, 'R', italic=True, sub_text='1', sup_text='н', value=d.sigma_b, indent=3.0,
                              desc="нормативное сопротивление растяжению (сжатию) металла трубопровода, принимается "
                                   "равным минимальному значению временного сопротивления металла, МПа ")
    add_paragraph_with_indent(doc, "m", italic=True, value=d.m,
                              desc="коэффициент условий работы трубопровода, принимаемый по исходным данным "
                                   "опросного листа Заказчика ([2] табл. 1);")
    add_paragraph_with_indent(doc, "k", italic=True, sub_text="1", value=d.k_1,
                              desc="коэффициент надежности по материалу ([2] табл. 10);")
    add_paragraph_with_indent(doc, "k", italic=True, sub_text="н", value=d.k_n,
                              desc="коэффициент надежности по ответственности трубопровода ([2] табл. 12), "
                                   "принимается по условному диаметру и давлению в трубопроводе.")
    R_1_calc = round(d.sigma_b * d.m / d.k_1 / d.k_n, 1)
    add_equation(doc, rf"R_1 = \frac{{{d.sigma_b} \cdot {d.m}}}{{{d.k_1} \cdot {d.k_n}}} = {R_1_calc} МПа")

    delta_calc = round((d.n * d.pressure * d.D_n / (2 * (R_1_calc + d.n * d.pressure))), 2)
    add_equation(doc, rf"\delta = \frac{{{d.n} \cdot {d.pressure} \cdot {d.D_n}}}"
                      rf"{{2 \cdot ({R_1_calc} + {d.n} \cdot {d.pressure})}} = {delta_calc} мм")

    doc.add_paragraph("3. Определение толщин стенок тройника", style='Heading 2')

    doc.add_paragraph("Расчетная толщина стенки магистральной части разрезного тройника определяется по формуле:")
    f_counter += 1
    f_c_mem_3 = f_counter
    add_equation(doc, r"T_h = \eta \cdot t_h", f_counter, f_l_indent=7.0)

    doc.add_paragraph("Расчетная толщина стенки ответвления разрезного тройника определяется по формуле:")
    f_counter += 1
    f_c_mem_4 = f_counter
    add_equation(doc, r"T_b = \xi \cdot T_h", f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")

    add_paragraph_with_indent(doc, "t", italic=True, sub_text="h", indent=1.25,
                              desc="расчетная толщина стерки условной трубы, имеющей диаметр магистрали и "
                                   "материал разрезного тройника ([2] формула (10)):")
    add_equation(doc, r"t_h = \frac{n \cdot p \cdot D_h}{2 \cdot (R_1 + n \cdot p)}", f_l_indent=7.0)
    add_paragraph_with_indent(doc, "t", italic=True, sub_text="b", indent=1.25,
                              desc="расчетная толщина стерки условной трубы, имеющей диаметр ответвления и "
                                   "материал разрезного тройника ([2] формула (10)):")
    add_equation(doc, r"t_b = \frac{n \cdot p \cdot D_b}{2 \cdot (R_1 + n \cdot p)}", f_l_indent=7.0)

    add_paragraph_with_indent(doc, "η", italic=True, indent=1.25, desc="коэффициент несущей способности тройника:")
    f_counter += 1
    f_c_mem_1 = f_counter
    add_equation(doc, r"\eta = \frac{1}{2 \cdot a} \cdot \left(-b + \sqrt{b^2 - 4 \cdot a \cdot c}\right)",
                 number=f_counter, f_l_indent=5.0)
    add_paragraph_with_indent(doc, "ξ", italic=True, indent=1.25, desc="коэффициент для сварного разрезного тройника:")
    f_counter += 1
    f_c_mem_2 = f_counter
    add_equation(doc, r"\xi = 0.45 + 0.55 \cdot \frac{D_b}{D_h}", number=f_counter, f_l_indent=6.0)

    p3 = doc.add_paragraph("Предварительно принимается наружный диаметр магистральной "
                           "части тройника и ответвления исходя из расчетной толщины стенки трубопровода ")
    p3.add_run("δ").italic = True
    p3.add_run(" в п.2.")

    if delta_calc <= 6.0:
        delta_calc = 6.0
        doc.add_paragraph("Минимальная толщина стенки из условия прочности тройника в соответствии "
                          "с принципом замещения площадей принимается равной 6 мм.")
    else:
        delta_calc = delta_calc

    doc.add_paragraph("Наружный диаметр магистральной части тройника:")
    f_counter += 1
    add_equation(doc, r"D_h = d_h + 2 \cdot \delta", number=f_counter, f_l_indent=7.0)

    doc.add_paragraph("Внутренний диаметр магистральной части:")
    if d.D_N_pipe < 1000:
        add = 4.0
    else:
        add = 6.0
    f_counter += 1
    add_equation(doc, rf"d_h = D_н + {add}", number=f_counter, f_l_indent=7.0)
    d_h_calc = d.D_n + add
    add_equation(doc, rf"d_h = {d.D_n} + {add} = {d_h_calc} мм")

    doc.add_paragraph("Наружный диаметр ответвления тройника:")
    f_counter += 1
    add_equation(doc, r"D_b = d_b + 2 \cdot \delta", number=f_counter, f_l_indent=7.0)
    D_b_calc = round(d.D_vn + 2 * delta_calc, 1)
    add_equation(doc, rf"D_b = {d.D_vn} + 2 \cdot {delta_calc} = {D_b_calc} мм")
    D_h_calc = round(d_h_calc + 2 * delta_calc, 1)
    add_equation(doc, rf"D_h = {d_h_calc} + 2 \cdot {delta_calc} = {D_h_calc} мм")

    doc.add_paragraph(f"Уточняющий расчет толщин стенок тройника (формулы ({f_c_mem_1}), ({f_c_mem_2}):")
    xi_calc = 0.45 + 0.55 * D_b_calc / D_h_calc
    add_equation(doc, rf"\xi = 0.45 + 0.55 \cdot \frac{{{D_b_calc}}}{{{D_h_calc}}} = {xi_calc:.3f}")

    p4 = doc.add_paragraph("Коэффициенты ")
    p4.add_run("a").italic = True
    p4.add_run(", ")
    p4.add_run("b").italic = True
    p4.add_run(", ")
    p4.add_run("c").italic = True
    p4.add_run(f" из выражения ({f_c_mem_1}) определяются по формулам:")
    add_equation(doc, r"a \cdot \eta^2 + b \cdot \eta + c = 0", f_l_indent=4.25)
    f_counter += 1
    add_equation(doc, r"a = (2 + 5 \cdot m_b - 4 \cdot \psi) \cdot \xi \cdot t_h",
                 number=f_counter, f_l_indent=4.75)
    f_counter += 1
    add_equation(doc, r"b = (2 \cdot \psi - 1) \cdot D_b + 4 \cdot \psi \cdot t_h - 5 \cdot m_b \cdot t_b",
                 number=f_counter, f_l_indent=3.5)
    f_counter += 1
    add_equation(doc, r"c = -2 \cdot \psi \cdot D_b",
                 number=f_counter, f_l_indent=6.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, "m", italic=True, sub_text="b", value=d.m_b, indent=2.25,
                              desc=f"коэффициент несущей способности тройника. Принимается равным 1, "
                                   f"если материал магистральной части и ответвления одинаковый – "
                                   f"сталь класса прочности {d.steel_class};")
    f_counter += 1
    add_equation(doc, r"\psi = \frac{L}{d_b}", number=f_counter, f_l_indent=7.0)
    add_paragraph_with_indent(doc, "L", italic=True, indent=1.25, desc="полудлина магистрали разрезного тройника;")
    psi_calc = d.L / d.D_vn
    add_equation(doc, rf"\psi = \frac{{{d.L}}}{{{d.D_vn}}} = {psi_calc:.3f}")

    p4 = doc.add_paragraph("Расчет коэффициентов ")
    p4.add_run("a").italic = True
    p4.add_run(", ")
    p4.add_run("b").italic = True
    p4.add_run(", ")
    p4.add_run("c").italic = True
    p4.add_run(" для определения несущей способности тройника:")

    t_h_calc = round((d.n * d.pressure * D_h_calc / (2 * (R_1_calc + d.n * d.pressure))), 2)
    t_b_calc = round((d.n * d.pressure * D_b_calc / (2 * (R_1_calc + d.n * d.pressure))), 2)
    add_equation(doc, rf"t_h = \frac{{{d.n} \cdot {d.pressure} \cdot {D_h_calc}}}"
                      rf"{{2 \cdot ({R_1_calc} + {d.n} \cdot {d.pressure})}} = {t_h_calc} мм")
    add_equation(doc, rf"t_b = \frac{{{d.n} \cdot {d.pressure} \cdot D_b}}"
                      rf"{{2 \cdot ({R_1_calc} + {d.n} \cdot {d.pressure})}} = {t_b_calc} мм")
    a_calc = (2 + 5 * d.m_b - 4 * psi_calc) * xi_calc * t_h_calc
    add_equation(doc, rf"a = (2 + 5 \cdot {d.m_b} - 4 \cdot {psi_calc:.3f}) \cdot {xi_calc:.3f} \cdot {t_h_calc}"
                      rf" = {a_calc:.3f}")
    b_calc = (2 * psi_calc - 1) * D_b_calc + 4 * psi_calc * t_h_calc - 5 * d.m_b * t_b_calc
    add_equation(doc, rf"b = (2 \cdot {psi_calc:.3f} - 1) \cdot {D_b_calc} + "
                      rf"4 \cdot {psi_calc:.3f} \cdot {t_h_calc} - 5 \cdot {d.m_b} \cdot {t_b_calc} = {b_calc:.3f}")
    c_calc = -2 * psi_calc * D_b_calc
    add_equation(doc, rf"c = -2 \cdot {psi_calc:.3f} \cdot {D_b_calc} = {c_calc:.3f}")

    doc.add_paragraph("Решение уравнения для коэффициента несущей способности тройника:")

    eta_calc = (1 / (2 * a_calc)) * (-b_calc + (b_calc ** 2 - 4 * a_calc * c_calc) ** 0.5)

    add_equation(doc, rf"\eta = 1/(2 \cdot {a_calc:.3f}) \cdot (-{b_calc:.3f} + "
                      rf"\sqrt{{{b_calc:.3f}^2 - 4 \cdot {a_calc:.3f} \cdot ({c_calc:.3f})}}) = {eta_calc:.3f}")

    doc.add_paragraph(f"Тогда, по формулам ({f_c_mem_3}), ({f_c_mem_4}):")

    T_h_calc_r = round(eta_calc * t_h_calc, 1)
    T_b_calc_r = round(xi_calc * T_h_calc_r, 1)
    add_equation(doc, rf"T_h = {eta_calc:.3f} \cdot {t_h_calc} = {T_h_calc_r} мм")
    add_equation(doc, rf"T_b = {xi_calc:.3f} \cdot {T_h_calc_r} = {T_b_calc_r} мм")

    if T_h_calc_r <= 6.0 or T_b_calc_r <= 6.0:
        T_h_calc = 6.0
        T_b_calc = xi_calc * T_h_calc
        if T_b_calc <= 6.0:
            T_b_calc = 6.0
            doc.add_paragraph("Минимальная толщина стенки из условия прочности тройника в соответствии "
                              "с принципом замещения площадей принимается равной 6 мм.")
            add_equation(doc, rf"T_h = {T_h_calc} мм")
            add_equation(doc, rf"T_b = {T_b_calc} мм")
        else:
            T_b_calc = T_b_calc
            add_equation(doc, rf"T_h = {T_h_calc} мм")
            add_equation(doc, rf"T_b = {T_b_calc} мм")
    else:
        T_h_calc = T_h_calc_r
        T_b_calc = T_b_calc_r
        add_equation(doc, rf"T_h = {T_h_calc} мм")
        add_equation(doc, rf"T_b = {T_b_calc} мм")

    doc.add_paragraph("После итерационного расчета принимаем толщины стенок:")
    add_equation(doc, rf"T_h = {d.T_h} мм")
    add_equation(doc, rf"T_b = {d.T_b} мм")
    T_h_calc = d.T_h
    T_b_calc = d.T_b

    doc.add_paragraph("Уточненный наружный диаметр магистральной части тройника и "
                      "ответвления по принятым толщинам стенок:")
    D_h_calc = d_h_calc + 2 * T_h_calc
    add_equation(doc, rf'D_h = {d_h_calc} + 2 \cdot {T_h_calc} = {D_h_calc} мм')
    D_b_calc = d.D_vn + 2 * T_b_calc
    add_equation(doc, rf'D_b = {d.D_vn} + 2 \cdot {T_b_calc} = {D_b_calc} мм')
    t_h_calc = round((d.n * d.pressure * D_h_calc / (2 * (R_1_calc + d.n * d.pressure))), 2)
    t_b_calc = round((d.n * d.pressure * D_b_calc / (2 * (R_1_calc + d.n * d.pressure))), 2)

    doc.add_paragraph("Полудлина тройника L должна удовлетворять неравенству:")

    f_counter += 1
    add_equation(doc, r"L >= 0.5 \cdot D_b + 2 \cdot T_h", number=f_counter, f_l_indent=6.0)
    add_equation(doc, rf"{d.L} >= 0.5 \cdot {D_b_calc} + 2 \cdot {T_h_calc} = {0.5 * D_b_calc + 2 * T_h_calc}")
    if d.L >= 0.5 * D_b_calc + 2 * T_h_calc:
        doc.add_paragraph("Условие полудлины тройника выполняется.")
    else:
        doc.add_paragraph("_____Условие полудлины тройника НЕ выполняется!!!_____")

    doc.add_paragraph("4. Условие прочности тройника в соответствии с принципом замещения площадей",
                      style='Heading 2')
    f_counter += 1
    add_equation(doc, r"A_1 + m_b \cdot A_2 >= A", number=f_counter, f_l_indent=6.0)
    doc.add_paragraph("где: ", style="Normal")
    A_calc = d.D_vn * t_h_calc
    add_equation(doc,
                 rf"A = d_b \cdot t_h = {d.D_vn} \cdot {t_h_calc} = {A_calc:.3f} мм^2")
    A1_calc = (2 * d.L - d.D_vn) * (T_h_calc - t_h_calc)
    add_equation(doc, rf"A_1 = (2 \cdot L - d_b) \cdot (T_h - t_h) = "
                      rf"(2 \cdot {d.L} - {d.D_vn}) \cdot ({T_h_calc} - {t_h_calc}) = {A1_calc:.3f} мм^2")
    H_1_calc = 2.5 * T_h_calc
    add_equation(doc, rf"H_1 = 2.5 \cdot T_h = 2.5 \cdot {T_h_calc} = {H_1_calc:.1f} мм")
    A2_calc = 2 * H_1_calc * (T_b_calc - t_b_calc)
    add_equation(doc,
                 rf"A_2 = 2 \cdot H_1 \cdot (T_b - t_b) = 2 \cdot {H_1_calc:.1f} \cdot ({T_b_calc} - {t_b_calc})"
                 rf" = {A2_calc:.3f} мм^2")
    doc.add_paragraph("Подставляем полученные значения в условие прочности:")
    if_A = A1_calc + d.m_b * A2_calc
    add_equation(doc, rf"{A1_calc:.3f} + {d.m_b} \cdot {A2_calc:.3f} = {if_A:.3f} >= {A_calc:.3f}")
    if if_A >= A_calc:
        doc.add_paragraph("Условие прочности тройника выполняется.")
    else:
        doc.add_paragraph("_____Условие прочности тройника НЕ выполняется!!!_____")

    doc.add_paragraph("5. Заключение", style='Heading 2')
    doc.add_paragraph("Значение расчетных и принятых толщин стенок тройника:")

    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[1].merge(hdr_cells[2])
    hdr_cells[3].merge(hdr_cells[4])

    hdr_cells[0].text = "Категория трубопровода"
    hdr_cells[1].text = "Th, мм"
    hdr_cells[3].text = "Tb, мм"

    row2 = table.rows[1].cells
    row2[1].text = "Расчетная"
    row2[2].text = "Принятая \n(с учетом технического допуска)"
    row2[3].text = "Расчетная"
    row2[4].text = "Принятая \n(с учетом технического допуска)"

    table.cell(0, 0).merge(table.cell(1, 0))

    row3 = table.rows[2].cells
    row3[0].text = "I"
    row3[1].text = f"{T_h_calc_r}"
    row3[2].text = f"{d.T_h}"
    row3[3].text = f"{T_b_calc_r}"
    row3[4].text = f"{d.T_b}"

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                run = p.runs[0]
                run.font.size = Pt(14)
                run.font.name = "Times New Roman"
                rFonts = run._element.rPr.rFonts
                rFonts.set(qn("w:eastAsia"), "Times New Roman")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing = Pt(16)

    p5 = doc.add_paragraph()
    run5 = p5.add_run()
    run5.add_break(WD_BREAK.PAGE)

    doc.add_paragraph("Перечень документов", style='Heading 1')
    doc.add_paragraph("1.	СТО Газпром 2-2.3-116-2016 – Правила производства работ на газопроводах врезкой под "
                      "давлением;")
    doc.add_paragraph(
        "2.	СП36.13330.2012 – Магистральные трубопроводы (актуализированная редакция СНиП 2.05.06-85*);")
    doc.add_paragraph("3.	ГОСТ 31447-2012 – Трубы стальные сварные для магистральных газопроводов, нефтепроводов и "
                      "нефтепродуктопроводов;")

    # Сохраняем файл
    filename = f"Расчет_толщин_стенок_тройника_{int(d.D_N_pipe)}_{int(d.D_N_b)}_{int(d.pressure * 10)}"
    doc.save(f'{filename}.docx')

    # main_doc = Document(f'{filename}.docx')
    # add_frame(main_doc, 'template_1.docx', section_index=0)
    # add_frame(main_doc, 'template_2.docx', section_index=1)
    # add_frame(main_doc, 'template_3.docx', section_index=2)
    #
    # main_doc.save(f'{filename}_w_f.docx')
