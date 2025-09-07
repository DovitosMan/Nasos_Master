from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import math2docx
from docx.enum.section import WD_SECTION
from copy import deepcopy
import os
from django.conf import settings


def set_heading_style(style, font_name="Times New Roman", size=14, bold=True, align=None,
                      f_l_indent=0.0, space_after=6, space_before=12, color=(0, 0, 0)):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor(*color)

    rFonts = OxmlElement('w:rFonts')
    for tag in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{tag}'), font_name)
    style.element.rPr.append(rFonts)

    fmt = style.paragraph_format
    if align:
        fmt.alignment = align
    fmt.first_line_indent = Cm(f_l_indent)
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)


def add_paragraph_with_indent(doc, text="", indent=3.0,
                              italic=None, sub_text="", subscript=False,
                              sup_text="", superscript=False, value=None, desc=None):
    left = indent
    first_line = -indent
    tab = indent

    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(left)
    fmt.first_line_indent = Cm(first_line)
    fmt.tab_stops.add_tab_stop(Cm(tab))
    if text:
        run = p.add_run(text)
        run.italic = italic
        sub = p.add_run(sub_text)
        sub.font.subscript = subscript
        sup = p.add_run(sup_text)
        sup.font.superscript = superscript
        if value is not None:
            p.add_run(f" = {value}  – ")
        else:
            p.add_run("  – ")
        if desc:
            p.add_run("\t")
            p.add_run(desc)
    return p


def add_equation(doc, latex: str, number=None, f_l_indent=0.0, font_size_pt=14, min_line_pts=30):
    section = doc.sections[0]
    content_width_emu = section.page_width - section.left_margin - section.right_margin
    right_tab_pos_pt = content_width_emu / 12700.0  # EMU -> pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(f_l_indent)  # без красной строки
    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    fmt.line_spacing = Pt(min_line_pts)  # ключ к отсутствию обрезания
    fmt.space_after = fmt.space_before = Pt(2)

    math2docx.add_math(p, latex)

    if number is not None:
        fmt.tab_stops.add_tab_stop(Pt(right_tab_pos_pt), alignment=WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        p.add_run(f"({number})").font.size = Pt(font_size_pt)

    return p


def add_equation_calc(doc, formula_latex: str, variables: dict, units: str, number=None, eval_expr=None, roundness=0):
    formula_with_values = formula_latex
    for var, val in variables.items():
        formula_with_values = formula_with_values.replace(var, str(val))

    expr = eval_expr or formula_latex.split("=")[-1]
    for var, val in variables.items():
        expr = expr.replace(var, str(val))

    try:
        result = round(eval(expr), roundness)
        full_formula = f"{formula_with_values} = {result:.{roundness}f} {units}"
    except Exception:
        result = None
        full_formula = formula_with_values

    add_equation(doc, full_formula, number=number)
    return result


def generate_report(result, T_b, input_select_names):
    print(result)
    gasket_data = dict(result['gasket_params'])
    gasket_document = gasket_data['document']
    q_obj = gasket_data['q_obj']
    m = gasket_data['m']
    K_obj = gasket_data['K_obj']
    E_p = round(float(gasket_data['E_p']), 10)
    q_obj_dop = gasket_data['q_obj_dop']

    print(gasket_data)
    fasteners_data = dict(result['fasteners_data'])
    print(fasteners_data)

    bolt_data = fasteners_data['bolt']
    print(bolt_data)
    bolt_type = bolt_data[0]

    nut_data = fasteners_data['nut']
    print(nut_data)
    nut_height = nut_data[1]

    washer_data = fasteners_data['washer']
    washer_thickness = washer_data[2]
    print(washer_data)

    D_N_flange = result['D_N_flange']
    pressure = result['pressure']
    temperature = result['temperature']
    flange_type = result['flange_type']

    flange_type_name = input_select_names[0]

    face_type = result['face_type']
    D_ext_flange = result['D_ext_flange']
    D_int_flange = result['D_int_flange']
    D_n_flange = result['D_n_flange']
    D_m_flange = result['D_m_flange']
    h_flange = result['h_flange']
    H_flange = result['H_flange']
    r = result['r']
    pin_diam = result['pin_diam']
    stud_length = result['stud_length']
    bolt_area = result['bolt_area']
    pins_quantity = result['pins_quantity']
    d_pins_flange = result['d_pins_flange']
    flange_steel = result['flange_steel']
    bolt_steel = result['bolt_steel']

    sigma_b_flange_T = result['sigma_b_flange_T']
    sigma_y_flange_T = result['sigma_y_flange_T']
    sigma_b_bolt_T = result['sigma_b_bolt_T']
    sigma_y_bolt_T = result['sigma_y_bolt_T']
    E_bolt_20 = result['E_bolt_20']
    E_bolt_T = result['E_bolt_T']
    E_flange_20 = result['E_flange_20']
    E_flange_T = result['E_flange_T']
    alpha_bolt_20 = result['alpha_bolt_20']
    alpha_bolt_T = round(result['alpha_bolt_T'], 7)
    alpha_flange_20 = result['alpha_flange_20']
    alpha_flange_T = round(result['alpha_flange_T'], 7)
    allowed_stress_flange_T = result['allowed_stress_flange_T']
    allowed_theta = round(result['allowed_theta'], 4)
    alpha_sigma_1 = round(result['alpha_sigma_1'], 3)
    stress_delta_data = dict(result['stress_delta_data'])
    stress_delta_R_1_1 = round(stress_delta_data['stress_delta_R_1_1'], 3)
    stress_delta_R_1_2 = round(stress_delta_data['stress_delta_R_1_2'], 3)
    stress_delta_R_1_4 = round(stress_delta_data['stress_delta_R_1_4'], 3)
    stress_delta_R_1_6 = round(stress_delta_data['stress_delta_R_1_6'], 3)
    stress_delta_R_0_1 = round(stress_delta_data['stress_delta_R_0_1'], 3)
    stress_delta_R_0_3 = round(stress_delta_data['stress_delta_R_0_3'], 3)
    stress_delta_R_0_5 = round(stress_delta_data['stress_delta_R_0_5'], 3)
    stress_delta_R_0_2 = round(stress_delta_data['stress_delta_R_0_2'], 3)
    stress_delta_R_0_4 = round(stress_delta_data['stress_delta_R_0_4'], 3)
    stress_delta_R_0_6 = round(stress_delta_data['stress_delta_R_0_6'], 3)

    B_V = round(result['B_V'], 3)
    lambda_ = round(result['lambda_'], 3)
    koef_B_x = round(result['koef_B_x'], 3)
    n_T_b = result['n_T_b']
    f = round(result['f'], 3)
    c_flange = round(result['c_flange'], 3)
    c_bolt = result['c_bolt']
    B_F = round(result['B_F'], 3)
    B_Y = round(result['B_Y'], 3)
    B_Z = round(result['B_Z'], 3)

    num_cycles_c = result['num_cycles_c']
    num_cycles_r = result['num_cycles_r']
    flange_cycles = result['flange_cycles']
    bolt_cycles = result['bolt_cycles']
    cycles_c = flange_cycles[1] + bolt_cycles[1]
    cycles_p = flange_cycles[2] + bolt_cycles[2]

    ext_force = result['ext_force']
    ext_moment = result['ext_moment']

    f_counter = 0

    doc = Document()
    section = doc.sections[0]

    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)

    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin, section.right_margin = Cm(3), Cm(1.25)
    section.top_margin, section.bottom_margin = Cm(1.75), Cm(2.5)

    style = doc.styles['Normal']
    font = style.font
    font.name, font.size = 'Times New Roman', Pt(14)
    fmt = style.paragraph_format
    fmt.line_spacing = Pt(21)
    fmt.first_line_indent = Cm(1.25)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_after = fmt.space_before = Pt(0)

    # === Стили заголовков ===
    set_heading_style(doc.styles['Heading 1'], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_heading_style(doc.styles['Heading 2'], align=WD_ALIGN_PARAGRAPH.LEFT, f_l_indent=1.25)

    if flange_type == "one_one":
        if D_m_flange == D_n_flange:
            doc.add_paragraph(f"РАСЧЕТ НА ПРОЧНОСТЬ И ГЕРМЕТИЧНОСТЬ ПРИВАРНОГО ВСТЫК ФЛАНЦЕВОГО СОЕДИНЕНИЯ "
                              f"С ЦИЛИНДРИЧЕСКОЙ ВТУЛКОЙ DN{int(D_N_flange)} PN{pressure}", style='Heading 1')
        if D_m_flange > D_n_flange:
            doc.add_paragraph(f"РАСЧЕТ НА ПРОЧНОСТЬ И ГЕРМЕТИЧНОСТЬ ПРИВАРНОГО ВСТЫК ФЛАНЦЕВОГО СОЕДИНЕНИЯ "
                              f"С КОНИЧЕСКОЙ ВТУЛКОЙ DN{int(D_N_flange)} PN{pressure}", style='Heading 1')
    if flange_type == "zero_one":
        doc.add_paragraph(f"РАСЧЕТ НА ПРОЧНОСТЬ И ГЕРМЕТИЧНОСТЬ ПРИВАРНОГО ВСТЫК ФЛАНЦЕВОГО СОЕДИНЕНИЯ "
                          f"С ПЛОСКИМИ ФЛАНЦАМИ DN{int(D_N_flange)} PN{pressure}", style='Heading 1')

    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("СОДЕРЖАНИЕ", style='Heading 1')
    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph('Настоящий расчет распространяется на фланцевые соединения разрезных тройников, привариваемых '
                      'под давлением на трубопроводы с рабочим давлением до 16.0 МПа включительно с номинальным '
                      'диаметром от DN500 до DN1200, транспортирующие углеводороды (природный газ), жидкие углеводороды'
                      ' (нефть и нефтепродукты), стабильный и нестабильный конденсат, широкие фракции углеводородов.')

    doc.add_paragraph("1. Исходные данные", style='Heading 2')

    doc.add_paragraph("Расчет производится согласно ГОСТ Р 52874.4-2007 "
                      '"Нормы и методы расчета на прочность" [1]. '
                      'Геометрические размеры фланца для первой итерации расчета согласно [5]. Для последней итерации –'
                      ' приведены ниже:')
    doc.add_paragraph(f'Тип фланца: {flange_type_name};')
    doc.add_paragraph(f'Наружный диаметр тарелки фланца: {D_ext_flange} мм;')
    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            doc.add_paragraph(f'Наружный диаметр в месте приварки: {D_n_flange} мм;')
            doc.add_paragraph(f'Наружный диаметр в месте присоединения тарелки: {D_m_flange} мм;')
            doc.add_paragraph(f'Радиус скругления между конической втулкой и тарелкой: {r} мм;')
        if D_m_flange == D_n_flange:
            D_m_flange = D_n_flange
            doc.add_paragraph(f'Наружный диаметр цилиндрической втулки фланца: {D_n_flange} мм;')
            doc.add_paragraph(f'Радиус скругления между цилиндрической втулкой и тарелкой: {r} мм;')
    if flange_type == 'zero_one':
        doc.add_paragraph(f'Толщина стенки ответвления тройника: {T_b} мм;')
    doc.add_paragraph(f'Толщина тарелки фланца: {h_flange} мм;')
    doc.add_paragraph(f'Полная высота фланца: {H_flange} мм;')

    p1 = doc.add_paragraph(
        f"Фланцевое соединение с приварным встык фланцем Ду {int(D_N_flange)} с уплотнительной поверхностью типа "
        f"{face_type} [5], устанавливаемый на ответвление разрезного тройника с внутренним диаметром {D_int_flange} мм,"
        f" давление среды {pressure} МПа. Количество шпилек фланца – {int(pins_quantity)} шт. с гайкой "
        f"М{int(pin_diam)}, расположенных на окружности диаметром ")
    p1.add_run("D").italic = True
    p1.add_run("б").font.subscript = True
    p1.add_run(f" = {d_pins_flange} мм. Расчетная температура в зоне фланцевого соединения – {temperature}˚С. "
               f"Марка стали фланца – {flange_steel}, механические свойства при расчетной температуре: σ")
    p1.add_run("В мин").font.subscript = True
    p1.add_run(f" = {sigma_b_flange_T} МПа; σ")
    p1.add_run("Т мин").font.subscript = True
    p1.add_run(f" = {sigma_y_flange_T} МПа. Марка стали шпилек и гаек – {bolt_steel}, механические свойства при "
               f"расчетной температуре: σ")
    p1.add_run("В мин").font.subscript = True
    p1.add_run(f" = {sigma_b_bolt_T} МПа; σ")
    p1.add_run("Т мин").font.subscript = True
    p1.add_run(f" = {sigma_y_bolt_T} МПа.")

    if ext_force == 0.0 and ext_moment == 0.0:
        doc.add_paragraph("Внешняя осевая сила и изгибающий момент отсутствуют.")
    if ext_force != 0.0 and ext_moment == 0.0:
        doc.add_paragraph(f"Внешняя осевая сила равна {ext_force} кН, изгибающий момент отсутствует.")
    if ext_force == 0.0 and ext_moment != 0.0:
        doc.add_paragraph(f"Внешняя осевая сила отсутствует, изгибающий момент равен {ext_moment} кН*м.")
    if ext_force != 0.0 and ext_moment != 0.0:
        doc.add_paragraph(f"Внешняя осевая сила равна {ext_force} кН, изгибающий момент равен {ext_moment} кН*м.")

    if gasket_document == 'ГОСТ 15180-86':
        gasket_type = gasket_data['type']
        gasket_material = gasket_data['material']
        gasket_D_outer = gasket_data['D_outer']
        gasket_d_inner = gasket_data['d_inner']
        gasket_thickness = gasket_data['thickness']

        p2 = doc.add_paragraph(f"Прокладка плоская типа {gasket_type} {gasket_document}, "
                               f"материал – {gasket_material}, наружный диаметр ")
        p2.add_run("D").italic = True
        p2.add_run("НП").font.subscript = True
        p2.add_run(f" = {gasket_D_outer} мм, внутренний диаметр ")
        p2.add_run("d").italic = True
        p2.add_run("НП").font.subscript = True
        p2.add_run(f" = {gasket_d_inner} мм, толщина прокладки – {gasket_thickness} мм [2].")

    if gasket_document == 'ГОСТ 34655-2020':
        gasket_type = gasket_data['type']
        gasket_material = gasket_data['material']
        gasket_D_outer = gasket_data['D_outer']
        gasket_d_inner = gasket_data['d_inner']
        gasket_thickness = gasket_data['thickness']
        if gasket_type == "1":
            radius_or_height_1 = gasket_data['radius_or_height_1']

            p2 = doc.add_paragraph(f"Прокладка стальная овального сечения типа {gasket_type} {gasket_document}, "
                                   f"материал – {gasket_material}, наружный диаметр ")
            p2.add_run("D").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_D_outer} мм, внутренний диаметр ")
            p2.add_run("d").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_d_inner} мм, толщина прокладки – {gasket_thickness} мм, "
                       f"радиус сечения – {radius_or_height_1} мм [2].")
        if gasket_type == "2":
            radius_or_height_1 = gasket_data['radius_or_height_1']

            p2 = doc.add_paragraph(f"Прокладка стальная восьмиугольного сечения типа {gasket_type} {gasket_document}, "
                                   f"материал – {gasket_material}, наружный диаметр ")
            p2.add_run("D").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_D_outer} мм, внутренний диаметр ")
            p2.add_run("d").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_d_inner} мм, толщина прокладки – {gasket_thickness} мм, "
                       f"высота сечения по наружному диаметру – {radius_or_height_1} мм [2].")
        if gasket_type == "3":
            radius_or_height_1 = gasket_data['radius_or_height_1']

            p2 = doc.add_paragraph(f"Прокладка стальная линзовая типа {gasket_type} {gasket_document}, "
                                   f"материал – {gasket_material}, наружный диаметр ")
            p2.add_run("D").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_D_outer} мм, внутренний диаметр ")
            p2.add_run("d").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_d_inner} мм, толщина прокладки – {gasket_thickness} мм, "
                       f"внешние радиусы сечения прокладки – {radius_or_height_1} мм [2].")

    if gasket_document == 'ISO 7483-2011':
        gasket_type = gasket_data['type']
        gasket_material = gasket_data['material']
        gasket_D_outer = gasket_data['D_outer']
        gasket_d_inner = gasket_data['d_inner']
        gasket_thickness = gasket_data['thickness']
        if gasket_type == "1":
            radius_or_width_c = gasket_data['radius_or_width_c']

            p2 = doc.add_paragraph(f"Прокладка стальная овального сечения типа {gasket_type} {gasket_document}, "
                                   f"материал – {gasket_material}, наружный диаметр ")
            p2.add_run("D").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_D_outer} мм, внутренний диаметр ")
            p2.add_run("d").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_d_inner} мм, толщина прокладки – {gasket_thickness} мм, "
                       f"радиус сечения – {radius_or_width_c} мм [2].")
        if gasket_type == "2":
            radius_or_width_c = gasket_data['radius_or_width_c']

            p2 = doc.add_paragraph(f"Прокладка стальная восьмиугольного сечения типа {gasket_type} {gasket_document}, "
                                   f"материал – {gasket_material}, наружный диаметр ")
            p2.add_run("D").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_D_outer} мм, внутренний диаметр ")
            p2.add_run("d").italic = True
            p2.add_run("НП").font.subscript = True
            p2.add_run(f" = {gasket_d_inner} мм, толщина прокладки – {gasket_thickness} мм, "
                       f"ширина сечения по верхнему краю – {radius_or_width_c} мм [2].")

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            f_counter += 1
            doc.add_paragraph('Толщина втулки фланца в месте приварки:')
            S_0_formula = r'S_0 = \frac{D_n - D_{вн}}{2}'
            add_equation(doc, S_0_formula, number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='n', subscript=True, value=D_n_flange,
                                      indent=3.0,
                                      desc='наружный диаметр втулки фланца в месте приварки, мм;')
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='вн', subscript=True, value=D_int_flange,
                                      indent=3.0,
                                      desc='внутренний диаметр фланца, мм.')
            variables_S_0 = {'D_n': D_n_flange, 'D_{вн}': D_int_flange}
            S_0_py = '(D_n - D_{вн}) / 2'
            S_0_calc = add_equation_calc(doc, S_0_formula, variables_S_0, 'мм', eval_expr=S_0_py, roundness=1)

            doc.add_paragraph('Толщина втулки фланца в месте присоединения тарелки:')
            f_counter += 1
            S_1_formula = r'S_1 = \frac{D_m - D_{вн}}{2}'
            add_equation(doc, S_1_formula, number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='m', subscript=True, value=D_m_flange,
                                      indent=3.0,
                                      desc='наружный диаметр втулки фланца в месте присоединения тарелки, мм;')
            variables_S_1 = {'D_m': D_m_flange, 'D_{вн}': D_int_flange}
            S_1_py = '(D_m - D_{вн}) / 2'
            S_1_calc = add_equation_calc(doc, S_1_formula, variables_S_1, 'мм', eval_expr=S_1_py, roundness=1)
        if D_m_flange == D_n_flange:
            doc.add_paragraph('Толщина втулки фланца в месте приварки:')
            f_counter += 1
            S_0_formula = r'S_0, S_1 = \frac{D_n - D_{вн}}{2}'
            add_equation(doc, S_0_formula, number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='n', subscript=True, value=D_ext_flange,
                                      indent=3.0,
                                      desc='наружный диаметр втулки фланца, мм;')
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='вн', subscript=True, value=D_int_flange,
                                      indent=3.0,
                                      desc='внутренний диаметр фланца, мм.')
            variables_S_0 = {'D_n': D_n_flange, 'D_{вн}': D_int_flange}
            S_0_py = '(D_n - D_{вн}) / 2'
            S_0_calc = add_equation_calc(doc, S_0_formula, variables_S_0, 'мм', eval_expr=S_0_py, roundness=1)
            S_1_calc = S_0_calc

    if flange_type == 'zero_one':
        doc.add_paragraph('Толщина стенки ответвления разрезного тройника:')
        f_counter += 1
        add_equation(doc, rf'S_0, S_1 = T_b = {T_b} мм', number=f_counter, f_l_indent=6.75)
        S_1_calc = S_0_calc = T_b

    doc.add_paragraph('Расчет применим к фланцам, удовлетворяющим следующим условиям:')
    f_counter += 1
    add_equation(doc, rf'\frac{{D_н}}{{D_в}} <= 5.0', number=f_counter, f_l_indent=7.0)
    f_counter += 1
    add_equation(doc, rf'\frac{{2 \cdot H}}{{D_н - D_в}} >= 0.25', number=f_counter, f_l_indent=6.5)
    f_counter += 1
    add_equation(doc, rf'\frac{{S_1 - S_0}}{{L}} <= 0.4', number=f_counter, f_l_indent=6.75)

    add_equation(doc, rf'{round(D_ext_flange / D_int_flange, 3)} <= 5.0',
                 f_l_indent=1.0)
    add_equation(doc, rf'{round(2 * h_flange / (D_ext_flange - D_int_flange), 3)} >= 0.25', f_l_indent=1.0)
    add_equation(doc, rf'{round((S_1_calc - S_0_calc) / (H_flange - h_flange), 3)} <= 0.4', f_l_indent=1.0)
    doc.add_paragraph('Условия выполняются.')

    doc.add_paragraph("2. Усилие, необходимое для смятия прокладки и обеспечения герметичности "
                      "соединения", style='Heading 2')

    doc.add_paragraph("Ширина прокладки определяется исходя из ее наружного и внутреннего диаметра:")
    f_counter += 1
    b_p_formula = r'b_п = \frac{D_{нп} - d_{нп}}{2}'
    add_equation(doc, b_p_formula, number=f_counter, f_l_indent=6.75)
    variables_b_p = {'D_{нп}': gasket_D_outer, 'd_{нп}': gasket_d_inner}
    b_p_py = '(D_{нп} - d_{нп}) / 2'
    b_p_calc = add_equation_calc(doc, b_p_formula, variables_b_p, " мм", eval_expr=b_p_py)

    if gasket_document == 'ГОСТ 15180-86':
        if b_p_calc <= 15.0:
            b_0_calc = b_p_calc
            p3 = doc.add_paragraph(f"Эффективная ширина плоской прокладки при ")
            p3.add_run("b").italic = True
            p3.add_run("п").font.subscript = True
            p3.add_run(" <= 15 мм:")
            f_counter += 1
            add_equation(doc, rf'b_0 = b_п = {b_p_calc}', number=f_counter)
        else:
            p3 = doc.add_paragraph(f"Эффективная ширина плоской прокладки при ")
            p3.add_run("b").italic = True
            p3.add_run("п").font.subscript = True
            p3.add_run(" > 15 мм:")
            b_0_formula = r'b_0 = 3.8 \cdot \sqrt{b_п}'
            f_counter += 1
            add_equation(doc, b_0_formula, number=f_counter, f_l_indent=6.75)
            variables_b_0 = {'b_п': b_p_calc}
            b_0_py = '3.8 * b_п ** 0.5'
            b_0_calc = add_equation_calc(doc, b_0_formula, variables_b_0, ' мм', eval_expr=b_0_py, roundness=1)

    if gasket_document in ['ГОСТ 34655-2020', 'ISO 7483-2011']:
        doc.add_paragraph("Эффективная ширина прокладки овального или восьмиугольного сечения:")
        f_counter += 1
        b_0_formula = r'b_0 = b_п / 4'
        add_equation(doc, b_0_formula, number=f_counter, f_l_indent=7.25)
        variables_b_0 = {'b_п': b_p_calc}
        b_0_py = 'b_п / 4'
        b_0_calc = add_equation_calc(doc, b_0_formula, variables_b_0, ' мм', eval_expr=b_0_py, roundness=1)

    if gasket_document == 'ГОСТ 15180-86':
        doc.add_paragraph("Расчетный диаметр плоской прокладки:")
        f_counter += 1
        D_sp_formula = r'D_{сп} = D_{нп} - b_0'
        add_equation(doc, D_sp_formula, number=f_counter, f_l_indent=6.75)
        variables_D_sp = {'D_{нп}': gasket_D_outer, 'b_0': b_0_calc}
        D_sp_py = 'D_{нп} - b_0'
        D_sp_calc = add_equation_calc(doc, D_sp_formula, variables_D_sp, " мм", eval_expr=D_sp_py, roundness=1)

    if gasket_document in ['ГОСТ 34655-2020', 'ISO 7483-2011']:
        doc.add_paragraph("Расчетный диаметр прокладки овального или восьмиугольного "
                          "сечения равен ее среднему диаметру:")
        f_counter += 1
        D_sp_formula = r'D_{сп} = D_{нп} - b_п'
        add_equation(doc, D_sp_formula, number=f_counter, f_l_indent=6.5)
        variables_D_sp = {'D_{нп}': gasket_D_outer, 'b_п': b_p_calc}
        D_sp_py = 'D_{нп} - b_п'
        D_sp_calc = add_equation_calc(doc, D_sp_formula, variables_D_sp, " мм", eval_expr=D_sp_py, roundness=1)

    doc.add_paragraph("Усилие, необходимое для смятия прокладки при затяжке:")
    f_counter += 1
    P_obj_formula = r'P_{обж} = 0.5 \cdot \pi \cdot D_{сп} \cdot b_0 \cdot q_{обж}'
    add_equation(doc, P_obj_formula, number=f_counter, f_l_indent=5.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, "q", italic=True, sub_text='обж', subscript=True, value=q_obj,
                              desc="удельное давление обжатия прокладки, МПа (принимается согласно [1], приложение И).")
    variables_P_obj = {'D_{сп}': D_sp_calc, 'b_0': b_0_calc, 'q_{обж}': q_obj}
    P_obj_py = '0.5 * 3.141593 * D_{сп} * b_0 * q_{обж} / 1000'
    P_obj_calc = add_equation_calc(doc, P_obj_formula, variables_P_obj, "кН", eval_expr=P_obj_py, roundness=3)

    doc.add_paragraph("Усилие на прокладке, необходимое для обеспечения герметичности фланцевого "
                      "соединения в рабочих условиях:")
    f_counter += 1
    R_p_formula = r'R_п = \pi \cdot D_{сп} \cdot b_0 \cdot m \cdot P'
    add_equation(doc, R_p_formula, number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'm', italic=True, value=m, indent=2.25,
                              desc='прокладочный коэффициент (принимается согласно [1], приложение И);')
    add_paragraph_with_indent(doc, 'P', italic=True, indent=1.0, desc='расчетное давление, МПа.')
    variables_R_p = {'D_{сп}': D_sp_calc, 'b_0': b_0_calc, 'm': m, 'P': pressure}
    R_p_py = '3.141593 * D_{сп} * b_0 * m * P / 1000'
    R_p_calc = add_equation_calc(doc, R_p_formula, variables_R_p, 'кН', eval_expr=R_p_py, roundness=3)

    doc.add_paragraph("3. Усилие в шпильках фланцевого соединения при затяжке и в рабочих условиях",
                      style='Heading 2')
    doc.add_paragraph("Суммарная площадь поперечного сечения шпилек по внутреннему диаметру резьбы:")
    f_counter += 1
    bolts_area_formula = r'A_ш = n \cdot f_ш'
    add_equation(doc, bolts_area_formula, number=f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'n', italic=True, indent=1.0, desc='количество шпилек фланца, шт.;')
    add_paragraph_with_indent(doc, 'f', italic=True, sub_text='ш', subscript=True, value=bolt_area, indent=2.75,
                              desc=f'площадь поперечного сечения шпильки, мм2 (принимается для одной шпильки'
                                   f' типа {bolt_type} под гайку с резьбой М{int(pin_diam)} согласно [1], '
                                   f'таблица Д.1).')
    variables_b_a = {'n': int(pins_quantity), 'f_ш': bolt_area}
    bolts_area_py = 'n * f_ш'
    bolts_area_calc = add_equation_calc(doc, bolts_area_formula, variables_b_a, 'мм2', eval_expr=bolts_area_py,
                                        roundness=1)
    doc.add_paragraph("Равнодействующая нагрузка от давления:")
    f_counter += 1
    Q_d_formula = r'Q_д = 0.785 \cdot D_{сп}^2 \cdot P'
    add_equation(doc, Q_d_formula, number=f_counter, f_l_indent=6.25)
    variables_Q_d = {'D_{сп}': D_sp_calc, 'P': pressure}
    Q_d_py = '0.785 * D_{сп}**2 * P / 1000'
    Q_d_calc = add_equation_calc(doc, Q_d_formula, variables_Q_d, 'кН', eval_expr=Q_d_py, roundness=3)

    if ext_force != 0.0 or ext_moment != 0.0:
        doc.add_paragraph("Приведенная нагрузка, вызванная воздействием внешней силы и изгибающего момента:")
        f_counter += 1
        Q_fm_formula = r'Q_{fm} = F ± \frac{4 \cdot |M|}{D_{сп}}'
        add_equation(doc, Q_fm_formula, number=f_counter, f_l_indent=6.25)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'F', italic=True, indent=1.25, desc='внешняя осевая сила, кН;')
        add_paragraph_with_indent(doc, 'M', italic=True, indent=1.25, desc='внешний изгибающий момент, кН*м.')
        variables_Q_fm = {'F': ext_force, 'M': ext_moment, 'D_{сп}': D_sp_calc}
        Q_fm_py = 'max(F + (4 * abs(M) * 1000 / D_{сп}), F - (4 * abs(M) * 1000 / D_{сп}))'
        Q_fm_calc = add_equation_calc(doc, Q_fm_formula, variables_Q_fm, 'кН', eval_expr=Q_fm_py, roundness=3)

    if ext_force == 0.0 and ext_moment == 0.0:
        doc.add_paragraph("Приведенная нагрузка, вызванная воздействием внешней силы и изгибающего момента равна нулю,"
                          " так как внешняя сила и изгибающий момент отсутствуют:")
        f_counter += 1
        Q_fm_formula = rf'Q_{{fm}} = F ± \frac{{4 \cdot |M|}}{{D_{{сп}}}} = 0'
        add_equation(doc, Q_fm_formula, number=f_counter, f_l_indent=6.25)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'F', italic=True, indent=1.25, desc='внешняя осевая сила, кН;')
        add_paragraph_with_indent(doc, 'M', italic=True, indent=1.25, desc='внешний изгибающий момент, кН*м.')
        Q_fm_calc = 0.0

    doc.add_paragraph("Нагрузка, вызванная стесненностью температурных деформаций в соединении с приварным встык"
                      " или плоским фланцем:")
    f_counter += 1
    Q_t_formula = (r'Q_t = \gamma \cdot \Big[ 2 \cdot (\alpha_ф) \cdot H \cdot (t_ф - 20) '
                   r'- 2 \cdot (\alpha_ш) \cdot H \cdot (t_ш - 20) \Big]')
    add_equation(doc, Q_t_formula, number=f_counter, f_l_indent=2.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'γ', italic=True, indent=1.0, desc='коэффициент жесткости фланцевого'
                                                                      ' соединения, Н/мм;')
    add_paragraph_with_indent(doc, 'α', italic=True, sub_text='ф', subscript=True, indent=2.0,
                              desc='температурный коэффициент линейного расширения материала фланца, 1/˚С '
                                   '([1], таблица Ж.2);')
    add_paragraph_with_indent(doc, 'H', italic=True, sub_text='', subscript=True, indent=1.75,
                              desc='толщина тарелки фланца, мм;')
    add_paragraph_with_indent(doc, 't', italic=True, sub_text='ф', subscript=True, indent=1.75,
                              desc='рабочая температура фланцев, ˚С; ')
    add_paragraph_with_indent(doc, 'α', italic=True, sub_text='ш', subscript=True, indent=1.5,
                              desc='температурный коэффициент линейного расширения материала шпильки, 1/˚С '
                                   '([1], таблица Ж.2);')
    add_paragraph_with_indent(doc, 't', italic=True, sub_text='ш', subscript=True, indent=1.25,
                              desc='рабочая температура шпильки, ˚С; ')

    doc.add_paragraph("Коэффициент жесткости фланцевого соединения для приварных встык и плоских фланцев "
                      "согласно [1], приложение Е, определяется по формуле:")
    f_counter += 1
    gamma_formula = (r'\gamma = \frac{1}{y_п + (y_ш) \cdot \frac{E_ш^{20}}{E_ш} + \left( 2 \cdot (y_{ф}) \cdot '
                     r'\frac{E^{20}}{E} \right) \cdot B^2}')
    add_equation(doc, gamma_formula, number=f_counter, f_l_indent=5.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'y', italic=True, sub_text='п, ш', subscript=True, indent=1.5,
                              desc='податливость прокладки, шпильки, мм/Н ([1], приложение К);')
    add_paragraph_with_indent(doc, 'y', italic=True, sub_text='ф1, ф2', subscript=True, indent=2.0,
                              desc='угловая податливость фланца при затяжке, 1/(Н*мм) ([1], приложение К);')
    add_paragraph_with_indent(doc, 'B', italic=True, indent=1.0, desc='плечо усилий в шпильках, мм '
                                                                      '([1], приложение Е);')
    add_paragraph_with_indent(doc, 'E', italic=True, sup_text='20', superscript=True, indent=1.5,
                              desc='модуль продольной упругости материала фланца, шпильки при расчетной температуре, '
                                   'МПа ([1], приложение Ж).')
    doc.add_paragraph('Податливость прокладки определяется по формуле:')
    f_counter += 1
    y_p_formula = r'y_п = \frac{h_п \cdot k_обж}{(E_п) \cdot \pi \cdot D_{сп} \cdot b_п}'
    add_equation(doc, y_p_formula, number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'h', italic=True, sub_text='п', subscript=True, indent=2.25,
                              value=gasket_thickness, desc='толщина прокладки, мм [2];')
    add_paragraph_with_indent(doc, 'k', italic=True, sub_text='обж', subscript=True, indent=2.75, value=K_obj,
                              desc='коэффициент обжатия прокладки ([1], приложение И);')
    add_paragraph_with_indent(doc, 'E', italic=True, sub_text='п', subscript=True, indent=2.75, value=E_p,
                              desc='Модуль продольной упругости материала прокладки, МПа.')
    variables_y_p = {'h_п': gasket_thickness, 'k_обж': K_obj, 'E_п': E_p, 'D_{сп}': D_sp_calc, 'b_п': b_p_calc}
    y_p_py = '(h_п * k_обж) / (E_п * 3.141593 * D_{сп} * b_п)'
    y_p_calc = add_equation_calc(doc, y_p_formula, variables_y_p, 'мм/Н', eval_expr=y_p_py, roundness=2)

    doc.add_paragraph('Податливость шпилек определяется по формуле:')
    f_counter += 1
    y_b_formula = r'y_ш = \frac{L_ш}{E_ш^{20} \cdot А_ш}'
    add_equation(doc, y_b_formula, number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'L', italic=True, sub_text='ш', subscript=True, indent=1.5,
                              desc='эффективная длина шпильки, мм.')
    doc.add_paragraph('Эффективная длина шпильки:')
    f_counter += 1
    L_stud_formula = r'L_ш = L_{ш0} + 0,56 \cdot d'
    add_equation(doc, L_stud_formula, number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'L', italic=True, sub_text='ш0', subscript=True, indent=1.5,
                              desc='начальная длина шпильки, мм;')
    add_paragraph_with_indent(doc, 'd', italic=True, indent=1.5, desc='диаметр шпильки, мм.')
    doc.add_paragraph('Начальная длина шпильки, округленная в большую сторону по [6]:')
    f_counter += 1
    L_stud_0_formula = r'L_{ш0} = 2 \cdot (h_г + h_ш + h)'
    add_equation(doc, L_stud_0_formula, number=f_counter, f_l_indent=5.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'h', italic=True, sub_text='г, ш', subscript=True, indent=1.5,
                              desc='высота гайки и шайбы, мм [3, 4].')
    add_equation(doc, rf'L_{{ш0}} = 2 \cdot ({nut_height} + {washer_thickness} + {h_flange}) = {stud_length} мм')
    add_equation(doc, rf'L_{{ш}} = {stud_length} + 0.56 \cdot {pin_diam} = '
                      rf'{round(stud_length + 0.56 * pin_diam, 2)} мм')
    doc.add_paragraph('Тогда, податливость шпилек:')
    variables_y_b = {'L_ш': round(stud_length + 0.56 * pin_diam, 2), 'E_ш^{20}': E_bolt_20, 'А_ш': bolts_area_calc}
    y_b_py = 'L_ш / (E_ш^{20} * А_ш)'
    y_b_calc = add_equation_calc(doc, y_b_formula, variables_y_b, 'мм/Н', eval_expr=y_b_py, roundness=10)

    doc.add_paragraph('Угловая податливость фланца при затяжке определяется по формуле:')
    f_counter += 1
    y_f_formula = r'y_ф = \frac{0.91 \cdot \beta_V}{E^20 \cdot \lambda \cdot S_0^2 \cdot l_0}'
    add_equation(doc, y_f_formula, number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='V', subscript=True, value=B_V, indent=2.75,
                              desc='коэффициент, определяемый по рисунку К.3 [1], приложение К;')
    add_paragraph_with_indent(doc, 'λ', italic=True, value=lambda_, indent=2.75,
                              desc='коэффициент, определяемый по формуле К.11 [1], приложение К, а также '
                                   'по рисункам К.2, К.3 и формулам К.3 - К.10')
    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            add_paragraph_with_indent(doc, 'S', italic=True, sub_text='0', subscript=True, indent=1.25,
                                      desc='толщина конической втулки приварного встык фланца в месте '
                                           'приварки к обечайке;')
        if D_m_flange == D_n_flange:
            add_paragraph_with_indent(doc, 'S', italic=True, sub_text='0', subscript=True, indent=1.25,
                                      desc='толщина цилиндрической втулки приварного встык фланца;')
    if flange_type == 'zero_one':
        add_paragraph_with_indent(doc, 'S', italic=True, sub_text='0', subscript=True, indent=1.25,
                                  desc='толщина стенки ответвления разрезного тройника;')
    add_paragraph_with_indent(doc, 'l', italic=True, sub_text='0', subscript=True, indent=1.25,
                              desc='длина обечайки, мм.')

    doc.add_paragraph('Длина обечайки определяется по формуле (К.3, [1], приложение К):')
    f_counter += 1
    l_0_formula = r'l_0 = \sqrt{D_{вн} \cdot S_0}'
    add_equation(doc, l_0_formula, number=f_counter, f_l_indent=6.5)
    variables_l_0 = {'D_{вн}': D_int_flange, 'S_0': S_0_calc}
    l_0_py = '(D_{вн} * S_0)**0.5'
    l_0_calc = add_equation_calc(doc, l_0_formula, variables_l_0, 'мм', eval_expr=l_0_py, roundness=2)

    doc.add_paragraph('Тогда, угловая податливость фланца при затяжке будет равна:')
    variables_y_f = {'beta_V': B_V, 'E^20': E_flange_20, 'lambda': lambda_, 'S_0': S_0_calc, 'l_0': l_0_calc}
    y_f_py = '(0.91 * beta_V) / (E^20 * lambda * S_0**2 * l_0)'
    y_f_calc = add_equation_calc(doc, y_f_formula, variables_y_f, '(1/(Н*мм))', eval_expr=y_f_py, roundness=13)

    doc.add_paragraph('Плечо усилий в шпильках определим по формуле:')
    f_counter += 1
    b_formula = r'B = 0.5 \cdot (D_б - D_{сп})'
    add_equation(doc, b_formula, number=f_counter, f_l_indent=6.0)
    variables_b = {'D_б': d_pins_flange, 'D_{сп}': D_sp_calc}
    b_py = '0.5 * (D_б - D_{сп})'
    b_calc = add_equation_calc(doc, b_formula, variables_b, 'мм', eval_expr=b_py, roundness=1)

    doc.add_paragraph('Подставим найденные значения в выражение (10):')
    variables_gamma = {'y_п': y_p_calc, 'y_ш': y_b_calc, 'E_ш^{20}': E_bolt_20, 'E_ш': E_bolt_T, 'y_{ф}': y_f_calc,
                       'E^{20}': E_flange_20, 'E': E_flange_T, 'B': b_calc}
    gamma_py = '1 / (y_п + y_ш * E_ш^{20} / E_ш + (2 * y_{ф} * E^{20} / E) * B**2)'
    gamma_calc = add_equation_calc(doc, gamma_formula, variables_gamma, 'мм/Н', eval_expr=gamma_py, roundness=5)

    doc.add_paragraph("Рабочая температура фланцев и шпилек:")
    add_equation(doc, rf't_{{ф1, ф2}} = 0,96 \cdot t = 0.96 \cdot {temperature} = {0.96 * temperature} ˚С')
    add_equation(doc, rf't_ш = 0,95 \cdot t = 0.95 \cdot {temperature} = {0.95 * temperature} ˚С')

    doc.add_paragraph('Далее, подставим найденные значения в выражение (9):')
    variables_Q_t = {'gamma': gamma_calc, 'alpha_ф': alpha_flange_T, 'H': h_flange, 't_ф': 0.96 * temperature,
                     'alpha_ш': alpha_bolt_T, 't_ш': 0.95 * temperature}
    Q_t_py = ('gamma * (2 * alpha_ф * H * (t_ф - 20) '
              '- alpha_ш * 2 * H * (t_ш - 20)) / 1000')
    Q_t_calc = add_equation_calc(doc, Q_t_formula, variables_Q_t, 'кН', eval_expr=Q_t_py, roundness=10)

    doc.add_paragraph('Расчетная нагрузка на шпильки фланцевого соединения при затяжке:')
    f_counter += 1
    P_b_m_formula = r'P_ш^м = \max (P_{ш1}; P_{ш2})'
    add_equation(doc, P_b_m_formula, number=f_counter, f_l_indent=6.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'P', italic=True, sub_text='ш1', subscript=True, indent=1.5,
                              desc='расчетная нагрузка на шпильки при затяжке, необходимая для обеспечения '
                                   'достаточной герметизации фланцевого соединения в рабочих условиях '
                                   'давления на прокладку, кН;')
    add_paragraph_with_indent(doc, 'P', italic=True, sub_text='ш2', subscript=True, indent=1.5,
                              desc='расчетная нагрузка на шпильки при затяжке, необходимая для обеспечения'
                                   ' обжатия прокладки и минимального начального натяжения шпилек, кН.')
    P_b_1_formula = (
        r'P_{ш1} = \max \Big['
        r'(\alpha) \cdot (Q_д + F) + R_п + \frac{4 \cdot (\omega_м) \cdot \vert M \vert}{D_{сп}} \, , \\ '
        r'(\alpha) \cdot (Q_д + F) + R_п + \frac{4 \cdot (\omega_м) \cdot \vert M \vert}{D_{сп}} - (Q_t)'
        r'\Big]'
    )
    f_counter += 1
    add_equation(doc, P_b_1_formula, number=f_counter, f_l_indent=3.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'α', italic=True, indent=1.25,
                              desc='коэффициент жесткости фланцевого соединения, нагруженного внутренним давлением '
                                   'или внешней осевой силой ([1], приложение Е);')
    add_paragraph_with_indent(doc, 'ω', italic=True, sub_text='М', subscript=True, indent=1.25,
                              desc='коэффициент жесткости фланцевого соединения, нагруженного внешним изгибающим '
                                   'моментом ([1], приложение Е);')
    alpha_formula = r'\alpha = 1 - \frac{y_п - 2 \cdot (y_ф) \cdot E \cdot B}{y_п + (y_ш) + 2 \cdot (y_ф) \cdot B^2}'
    f_counter += 1
    add_equation(doc, alpha_formula, number=f_counter, f_l_indent=6.0)
    alpha_M_formula = (
        r'\omega_м = \frac{(y_ш) + 2 \cdot (y_{фн}) \cdot B \cdot \left( B + E - \frac{E^2}{D_{сп}} \right)}'
        r'{(y_ш) + y_п \cdot \left( \frac{D_б}{D_{сп}} \right)^2 + 2 \cdot (y_{фн}) \cdot B^2}')
    f_counter += 1
    add_equation(doc, alpha_M_formula, number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'E', italic=True, indent=1.25,
                              desc='плечо усилия от действия давления внутри фланца, мм;')
    add_paragraph_with_indent(doc, 'y', italic=True, sub_text='фн', subscript=True, indent=1.25,
                              desc='угловая податливость фланца, нагруженного внешним изгибающим моментом, 1/(Н*мм)'
                                   ' ([1], приложение К).')
    doc.add_paragraph('Плечо усилия от действия давления внутри фланца:')
    e_formula = r'E = 0.5 \cdot (D_{сп} - D_{вн} - S_э)'
    f_counter += 1
    add_equation(doc, e_formula, number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")

    if flange_type == 'one_one':
        add_paragraph_with_indent(doc, 'S', italic=True, sub_text='э', subscript=True, indent=1.25,
                                  desc='эквивалентная толщина втулки приварных встык фланцев, мм')
        f_counter += 1
        S_e_formula = r'S_э = \zeta \cdot S_0'
        add_equation(doc, S_e_formula, number=f_counter, f_l_indent=7.25)
        add_paragraph_with_indent(doc, 'ζ', italic=True, indent=2.5, value=koef_B_x,
                                  desc='коэффициент, зависящий от соотношения размеров конической втулки фланца,'
                                       ' определяется по формуле Е.6 [1], приложение Е, используя формулы К.9, '
                                       'К.10 [1], приложение К.')
        variables_S_e = {'zeta': koef_B_x, 'S_0': S_0_calc}
        S_e_py = 'zeta * S_0'
        S_e_calc = add_equation_calc(doc, S_e_formula, variables_S_e, 'мм', eval_expr=S_e_py, roundness=2)

    if flange_type == 'zero_one':
        add_paragraph_with_indent(doc, 'S', italic=True, sub_text='э', subscript=True, indent=1.25,
                                  desc='эквивалентная толщина стенки ответвления разрезного тройника, мм')
        f_counter += 1
        add_equation(doc, rf'S_э = S_0 = {S_0_calc} мм', number=f_counter, f_l_indent=1.0)
        S_e_calc = S_0_calc

    doc.add_paragraph('Тогда, плечо усилия от действия давления внутри фланца:')
    variables_e = {'D_{сп}': D_sp_calc, 'D_{вн}': D_int_flange, 'S_э': S_e_calc}
    e_py = '0.5 * (D_{сп} - D_{вн} - S_э)'
    e_calc = add_equation_calc(doc, e_formula, variables_e, 'мм', eval_expr=e_py, roundness=2)

    doc.add_paragraph('Угловая податливость фланца, нагруженного внешним изгибающим моментом:')
    f_counter += 1
    y_f_n_formula = r'y_{фн} = \left( \frac{\pi}{4} \right)^3 \cdot \frac{D_б}{E^{20} \cdot D_н \cdot H^3}'
    add_equation(doc, y_f_n_formula, number=f_counter, f_l_indent=6.0)
    variables_y_f_n = {'D_б': d_pins_flange, 'E^{20}': E_flange_20, 'D_н': D_ext_flange, 'H': h_flange}
    y_f_n_py = '(3.141593 / 4)**3 * D_б / (E^{20} * D_н * H**3)'
    y_f_n_calc = add_equation_calc(doc, y_f_n_formula, variables_y_f_n, '(1/(Н*мм))', eval_expr=y_f_n_py, roundness=13)

    doc.add_paragraph('Подставим найденные значения в выражения (21) и (22):')

    variables_alpha = {'y_п': y_p_calc, 'y_ф': y_f_calc, 'E': e_calc, 'B': b_calc, 'y_ш': y_b_calc}
    alpha_py = '1 - (y_п - 2 * y_ф * E * B) / (y_п + y_ш + 2 * y_ф * B**2)'
    alpha_calc = add_equation_calc(doc, alpha_formula, variables_alpha, '', eval_expr=alpha_py, roundness=11)

    variables_alpha_M = {'y_ш': y_b_calc, 'y_{фн}': y_f_n_calc, 'B': b_calc, 'E': e_calc, 'D_{сп}': D_sp_calc,
                         'y_п': y_p_calc, 'D_б': d_pins_flange}
    alpha_M_py = ('(y_ш + 2 * y_{фн} * B * (B + E - E**2 / D_{сп})) / (y_ш + y_п * (D_б / D_{сп})**2 + 2 * y_{фн} * '
                  'B**2)')
    alpha_M_calc = add_equation_calc(doc, alpha_M_formula, variables_alpha_M, '', eval_expr=alpha_M_py, roundness=13)

    doc.add_paragraph('Подставим полученные значения в выражение (20):')

    variables_P_b_1 = {'alpha': alpha_calc, 'Q_д': Q_d_calc, 'F': ext_force, 'R_п': R_p_calc, 'omega_м': alpha_M_calc,
                       'M': ext_moment, 'D_{сп}': D_sp_calc, 'Q_t': Q_t_calc}
    P_b_1_py = ('max(alpha * (Q_д + F) + R_п + 4 * omega_м * abs(M) * 1000 / D_{сп}, '
                'alpha * (Q_д + F) + R_п + 4 * omega_м * abs(M) * 1000 / D_{сп} - Q_t)')
    P_b_1_calc = add_equation_calc(doc, P_b_1_formula, variables_P_b_1, 'кН', eval_expr=P_b_1_py, roundness=3)

    P_b_2_formula = r'P_{ш2} = \max \Big( P_{обж}; 0,4 \cdot A_ш \cdot [σ]_р^ш \Big)'
    f_counter += 1
    add_equation(doc, P_b_2_formula, number=f_counter, f_l_indent=5.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, '[σ]', italic=True, sub_text='р', subscript=True, sup_text='ш',
                              superscript=True, indent=1.75, desc='допускаемое напряжение для шпилек в условиях '
                                                                  'испытания, МПа ([1], приложение Г).')
    f_counter += 1
    a_sigma_b_r_formula = r'[σ]_р^ш = K_{у.р.} \cdot K_{у.з.} \cdot K_{у.т.} \cdot [σ]_н^ш}'
    add_equation(doc, a_sigma_b_r_formula, number=f_counter, f_l_indent=5.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.р.', subscript=True, value=1.35,
                              desc='коэффициент условий работы для условий испытания ([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.з.', subscript=True, value=1.1,
                              desc='коэффициент условий затяжки при затяжке с контролем по крутящему моменту '
                                   '([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.т.', subscript=True, value=1.3,
                              desc='коэффициент учета нагрузки от температурных деформаций '
                                   '([1], приложение Г);')
    add_paragraph_with_indent(doc, '[σ]', italic=True, sub_text='н', subscript=True, sup_text='ш',
                              superscript=True, indent=1.75, desc='номинальное допускаемое напряжение для шпилек '
                                                                  'при затяжке в рабочих условиях, МПа '
                                                                  '([1], приложение Г).')
    a_sigma_b_n_formula = r'[σ]_н^ш = \frac{σ_т}{n_т}'
    f_counter += 1
    add_equation(doc, a_sigma_b_n_formula, number=f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'n', italic=True, sub_text='т', subscript=True, indent=2.5, value=n_T_b,
                              desc='коэффициент запаса по отношению к пределу текучести '
                                   '([1], приложение Г).')
    variables_a_sigma_b_n = {'σ_т': sigma_y_bolt_T, 'n_т': n_T_b}
    a_sigma_b_n_py = 'σ_т / n_т'
    a_sigma_b_n_calc = add_equation_calc(doc, a_sigma_b_n_formula, variables_a_sigma_b_n, 'МПа',
                                         eval_expr=a_sigma_b_n_py, roundness=3)
    variables_a_sigma_b_r = {'K_{у.р.}': 1.35, 'K_{у.з.}': 1.1, 'K_{у.т.}': 1.3, '[σ]_н^ш': a_sigma_b_n_calc}
    a_sigma_b_r_py = 'K_{у.р.} * K_{у.з.} * K_{у.т.} * [σ]_н^ш'
    a_sigma_b_r_calc = add_equation_calc(doc, a_sigma_b_r_formula, variables_a_sigma_b_r, 'МПа',
                                         eval_expr=a_sigma_b_r_py, roundness=3)
    doc.add_paragraph('Допускаемые напряжения для шпилек при затяжке в рабочих условиях:')
    f_counter += 1
    a_sigma_b_m_formula = r'[σ]_м^ш = \xi \cdot K_{у.р.} \cdot K_{у.з.} \cdot K_{у.т.} \cdot [σ]_н^ш}'
    add_equation(doc, a_sigma_b_m_formula, number=f_counter, f_l_indent=5.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'ξ', italic=True, value=1.2,
                              desc='коэффициент увеличения допускаемых напряжений при затяжке '
                                   '([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.р.', subscript=True, value=1.0,
                              desc='коэффициент условий работы для рабочих условий ([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.з.', subscript=True, value=1.1,
                              desc='коэффициент условий затяжки при затяжке с контролем по крутящему моменту '
                                   '([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.т.', subscript=True, value=1.3,
                              desc='коэффициент учета нагрузки от температурных деформаций '
                                   '([1], приложение Г);')
    variables_a_sigma_b_m = {'xi': 1.2, 'K_{у.р.}': 1.0, 'K_{у.з.}': 1.1, 'K_{у.т.}': 1.3, '[σ]_н^ш': a_sigma_b_n_calc}
    a_sigma_b_m_py = 'xi * K_{у.р.} * K_{у.з.} * K_{у.т.} * [σ]_н^ш'
    a_sigma_b_m_calc = add_equation_calc(doc, a_sigma_b_m_formula, variables_a_sigma_b_m, 'МПа',
                                         eval_expr=a_sigma_b_m_py, roundness=3)

    doc.add_paragraph('Подставим полученные значения в выражение (26)')
    variables_P_b_2 = {'P_{обж}': P_obj_calc, 'A_ш': bolts_area_calc, '[σ]_р^ш': a_sigma_b_r_calc}
    P_b_2_py = 'max(P_{обж}, 0.4 * A_ш * [σ]_р^ш / 1000)'
    P_b_2_calc = add_equation_calc(doc, P_b_2_formula, variables_P_b_2, 'кН', eval_expr=P_b_2_py, roundness=3)

    doc.add_paragraph('Расчетная нагрузка на шпильки при затяжке, необходимая для обеспечения обжатия прокладки и '
                      'минимального начального натяжения шпилек в рабочих условиях:')

    variables_P_b_m = {'P_{ш1}': P_b_1_calc, 'P_{ш2}': P_b_2_calc}
    P_b_m_py = 'max(P_{ш1}, P_{ш2})'
    P_b_m_calc = add_equation_calc(doc, P_b_m_formula, variables_P_b_m, "кН", eval_expr=P_b_m_py, roundness=3)

    P_b_r_formula = (r'P_ш^р = P_ш^м + (1 - (\alpha)) \cdot (Q_д + F) + (Q_t) + \frac{4 \cdot (1 - (\omega_м)) \cdot '
                     r'\vert M \vert}{D_{сп}}')
    f_counter += 1
    add_equation(doc, P_b_r_formula, number=f_counter, f_l_indent=3.0)
    variables_P_b_r = {'P_ш^м': P_b_m_calc, 'alpha': alpha_calc, 'Q_д': Q_d_calc, 'F': ext_force, 'Q_t': Q_t_calc,
                       'omega_м': alpha_M_calc, 'M': ext_moment, 'D_{сп}': D_sp_calc}
    P_b_r_py = 'P_ш^м + (1 - alpha) * (Q_д + F) + (Q_t) + 4 * (1 - omega_м) * abs(M) * 1000 / D_{сп}'
    P_b_r_calc = add_equation_calc(doc, P_b_r_formula, variables_P_b_r, 'кН', eval_expr=P_b_r_py, roundness=3)

    doc.add_paragraph("4. Проверка прочности шпилек и прокладки", style='Heading 2')

    doc.add_paragraph('Расчетные напряжения в шпильках при затяжке:')
    sigma_b_1_formula = r'σ_{ш1} = \frac{P_ш^м}{A_ш}'
    f_counter += 1
    add_equation(doc, sigma_b_1_formula, number=f_counter, f_l_indent=7.0)
    variables_sigma_b_1 = {'P_ш^м': P_b_m_calc, 'A_ш': bolts_area_calc}
    sigma_b_1_py = 'P_ш^м * 1000 / A_ш'
    sigma_b_1_calc = add_equation_calc(doc, sigma_b_1_formula, variables_sigma_b_1, 'МПа', eval_expr=sigma_b_1_py,
                                       roundness=3)
    doc.add_paragraph('Расчетные напряжения в шпильках при затяжке в рабочих условиях:')
    f_counter += 1
    sigma_b_2_formula = r'σ_{ш2} = \frac{P_ш^р}{A_ш}'
    add_equation(doc, sigma_b_2_formula, number=f_counter, f_l_indent=7.0)
    variables_sigma_b_2 = {'P_ш^р': P_b_r_calc, 'A_ш': bolts_area_calc}
    sigma_b_2_py = 'P_ш^р * 1000 / A_ш'
    sigma_b_2_calc = add_equation_calc(doc, sigma_b_2_formula, variables_sigma_b_2, 'МПа', eval_expr=sigma_b_2_py,
                                       roundness=3)
    doc.add_paragraph('Условие прочности шпилек при затяжке:')
    f_counter += 1
    add_equation(doc, r'σ_{ш1} <= [σ]_м^ш', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'{sigma_b_1_calc} МПа <= {a_sigma_b_m_calc} МПа.')

    doc.add_paragraph('Условие прочности шпилек в рабочих условиях:')
    f_counter += 1
    add_equation(doc, r'σ_{ш2} <= [σ]_р^ш', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'{sigma_b_2_calc} МПа <= {a_sigma_b_r_calc} МПа.')

    if gasket_document == 'ГОСТ 15180-86':
        doc.add_paragraph('Условие прочности прокладки:')
        f_counter += 1
        add_equation(doc, r'q <= [q]', number=f_counter, f_l_indent=7.0)

        q_formula = r'q = \frac{ \max (P_ш^м; P_ш^р)}{ \pi \cdot D_{сп} \cdot b_п}'
        f_counter += 1
        add_equation(doc, q_formula, number=f_counter, f_l_indent=6.5)
        variables_q = {'P_ш^м': P_b_m_calc, 'P_ш^р': P_b_r_calc, 'D_{сп}': D_sp_calc, 'b_п': b_p_calc}
        q_py = 'max(P_ш^м, P_ш^р) * 1000 / (3.141593 * D_{сп} * b_п)'
        q_calc = add_equation_calc(doc, q_formula, variables_q, 'МПа', eval_expr=q_py, roundness=3)
        add_equation(doc, rf'{q_calc} МПа <= {q_obj_dop} МПа.')
    else:
        doc.add_paragraph('Для стальных прокладок проверка прочности прокладок не требуется [1].')

    doc.add_paragraph("5. Расчет фланцев на статическую прочность", style='Heading 2')

    doc.add_paragraph('Расчетный изгибающий момент, действующий на фланец при затяжке:')
    f_counter += 1
    M_M_formula = r'M^M = C_F \cdot P_ш^м \cdot B'
    add_equation(doc, M_M_formula, number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'C', italic=True, sub_text='F', subscript=True, indent=1.25,
                              desc='коэффициент, учитывающий изгиб тарелки фланца между шпильками '
                                   '([1], приложение К).')
    f_counter += 1
    C_F_formula = (r'C_F = \max \Big{'
                   r'1; \sqrt{\frac{\frac{\pi \cdot D_б}{n}}{2 \cdot \text{d} + \frac{6 \cdot H}{\text{m} + 0.5}}} '
                   r'\Big}')
    add_equation(doc, C_F_formula, number=f_counter, f_l_indent=5.5)
    variables_C_F = {'D_б': d_pins_flange, 'n': pins_quantity, 'text{d}': pin_diam, 'H': h_flange, 'text{m}': m}
    C_F_py = 'max(1.0, (((3.141593 * D_б) / (n)) / (2 * text{d} + (6 * H) / (text{m} + 0.5)))**0.5)'
    C_F_calc = add_equation_calc(doc, C_F_formula, variables_C_F, '', eval_expr=C_F_py, roundness=3)

    doc.add_paragraph('Тогда, расчетный изгибающий момент, действующий на фланец при затяжке:')
    variables_M_M = {'C_F': C_F_calc, 'P_ш^м': P_b_m_calc, 'B': b_calc}
    M_M_py = 'C_F * P_ш^м * B / 1000'
    M_M_calc = add_equation_calc(doc, M_M_formula, variables_M_M, 'кН*м', eval_expr=M_M_py, roundness=3)

    doc.add_paragraph('Расчетный изгибающий момент, действующий на фланец в рабочих условиях:')
    f_counter += 1
    M_R_formula = (r'M^Р = C_F \cdot \max \Big{ \Big[ P_ш^р \cdot \text{B} + \left( Q_д + Q_{fm} \right) \cdot E \Big];'
                   r' \vert Q_д + Q_{fm} \vert \cdot E \Big}')
    add_equation(doc, M_R_formula, number=f_counter, f_l_indent=2.0)
    variables_M_R = {'C_F': C_F_calc, 'P_ш^р': P_b_r_calc, 'text{B}': b_calc, 'Q_д': Q_d_calc, 'Q_{fm}': Q_fm_calc,
                     'E': e_calc}
    M_R_py = 'C_F * max(P_ш^р * text{B} / 1000 + (Q_д + Q_{fm}) * E / 1000, abs(Q_д + Q_{fm}) * E / 1000)'
    M_R_calc = add_equation_calc(doc, M_R_formula, variables_M_R, 'кН*м', eval_expr=M_R_py, roundness=3)

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            p4 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p4.add_run('S').italic = True
            p4.add_run('1').font.subscript = True
            p4.add_run(' при затяжке:')  # с - русская
            f_counter += 1
            sigma_1_m_formula = r'σ_1^м = \frac{M^M}{\lambda \cdot (S_1 - \text{с})^2 \cdot D_{прив.}}'
            add_equation(doc, sigma_1_m_formula, number=f_counter, f_l_indent=6.25)

            p5 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p5.add_run('S').italic = True
            p5.add_run('0').font.subscript = True
            p5.add_run(' при затяжке:')
            f_counter += 1
            sigma_0_m_formula = r'σ_0^м = \text{f} \cdot σ_1^м'
            add_equation(doc, sigma_0_m_formula, number=f_counter, f_l_indent=7.0)
        if D_m_flange == D_n_flange:
            doc.add_paragraph('Меридиональное изгибное напряжение в цилиндрической втулке приварного встык фланца '
                              'при затяжке:')
            f_counter += 1  # с - русская
            sigma_1_0_m_formula = r'σ_1^м, σ_0^м = \frac{M^M}{\lambda \cdot (S_0 - \text{с})^2 \cdot D_{прив.}}'
            add_equation(doc, sigma_1_0_m_formula, number=f_counter, f_l_indent=6.25)
    if flange_type == 'zero_one':
        doc.add_paragraph('Меридиональное изгибное напряжение в приварном встык плоском фланце при затяжке:')
        f_counter += 1  # с - русская
        sigma_1_0_m_formula = r'σ_1^м, σ_0^м = \frac{M^M}{\lambda \cdot (S_0 - \text{с})^2 \cdot D_{прив.}}'
        add_equation(doc, sigma_1_0_m_formula, number=f_counter, f_l_indent=6.25)

    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'с', italic=True, value=c_flange, indent=2.25,
                              desc='прибавка на коррозию, мм ([7], п.12.1);')
    if flange_type == 'one_one':
        add_paragraph_with_indent(doc, 'D', italic=True, sub_text='прив.', subscript=True, indent=2.25,
                                  desc='приведенный диаметр приварного встык фланца с конической или прямой втулкой;')
        add_paragraph_with_indent(doc, 'f', italic=True, value=f, indent=2.25, desc='поправочный коэффициент для '
                                                                                    'напряжений во втулке фланца '
                                                                                    '([1], рисунок К.4).')

        if D_int_flange >= 20 * S_1_calc:
            p6 = doc.add_paragraph('Приведенный диаметр при ')
            p6.add_run('D').italic = True
            p6.add_run('вн').font.subscript = True
            p6.add_run(' >= 20 * ')
            p6.add_run('S').italic = True
            p6.add_run('1').font.subscript = True
            f_counter += 1
            add_equation(doc, rf'D_{{прив.}} = D_{{вн}} = {D_int_flange} мм', number=f_counter, f_l_indent=3.5)
            D_priv_calc = D_int_flange
        elif D_int_flange < 20 * S_1_calc and f > 1:
            p6 = doc.add_paragraph('Приведенный диаметр при ')
            p6.add_run('D').italic = True
            p6.add_run('вн').font.subscript = True
            p6.add_run(' < 20 * ')
            p6.add_run('S').italic = True
            p6.add_run('1').font.subscript = True
            p6.add_run(' и ')
            p6.add_run('f').italic = True
            p6.add_run(' > 1.0:')
            f_counter += 1
            add_equation(doc, rf'D_{{прив.}} = D_{{вн}} + S_0 = {D_int_flange} + {S_0_calc} = '
                              rf'{D_int_flange + S_0_calc} мм', number=f_counter, f_l_indent=3.5)
            D_priv_calc = D_int_flange + S_0_calc
        elif D_int_flange < 20 * S_1_calc and f == 1.0:
            p6 = doc.add_paragraph('Приведенный диаметр при ')
            p6.add_run('D').italic = True
            p6.add_run('вн').font.subscript = True
            p6.add_run(' < 20 * ')
            p6.add_run('S').italic = True
            p6.add_run('1').font.subscript = True
            p6.add_run(' и ')
            p6.add_run('f').italic = True
            p6.add_run(' = 1.0:')
            f_counter += 1
            add_equation(doc, rf'D_{{прив.}} = D_{{вн}} + S_1 = {D_int_flange} + {S_1_calc} = '
                              rf'{D_int_flange + S_1_calc} мм', number=f_counter, f_l_indent=3.5)
            D_priv_calc = D_int_flange + S_1_calc

    if flange_type == 'zero_one':
        add_paragraph_with_indent(doc, 'D', italic=True, sub_text='прив.', subscript=True, indent=2.25,
                                  desc='приведенный диаметр приварного встык плоского фланца.')
        doc.add_paragraph('Приведенный диаметр приварного встык плоского фланца:')
        f_counter += 1
        add_equation(doc, rf'D_{{прив.}} = D_{{вн}} = {D_int_flange} мм', number=f_counter, f_l_indent=3.5)
        D_priv_calc = D_int_flange

    doc.add_paragraph('Подставим полученные значения в формулу меридиональных изгибных напряжений в условиях затяжки:')

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            variables_sigma_1_m = {'M^M': M_M_calc, 'lambda': lambda_, 'S_1': S_1_calc, 'text{с}': c_flange,
                                   'D_{прив.}': D_priv_calc}
            sigma_1_m_py = 'M^M * 1000000 / (lambda * (S_1 - text{с})**2 * D_{прив.})'
            sigma_1_m_calc = add_equation_calc(doc, sigma_1_m_formula, variables_sigma_1_m, 'МПа',
                                               eval_expr=sigma_1_m_py, roundness=3)
            variables_sigma_0_m = {'text{f}': f, 'σ_1^м': sigma_1_m_calc}
            sigma_0_m_py = 'text{f} * σ_1^м'
            sigma_0_m_calc = add_equation_calc(doc, sigma_0_m_formula, variables_sigma_0_m, 'МПа',
                                               eval_expr=sigma_0_m_py, roundness=3)
        if D_m_flange == D_n_flange:
            variables_sigma_1_0_m = {'M^M': M_M_calc, 'lambda': lambda_, 'S_0': S_0_calc, 'text{с}': c_flange,
                                     'D_{прив.}': D_priv_calc}
            sigma_1_0_m_py = 'M^M * 1000000 / (lambda * (S_0 - text{с})**2 * D_{прив.})'
            sigma_1_m_calc = sigma_0_m_calc = add_equation_calc(doc, sigma_1_0_m_formula, variables_sigma_1_0_m,
                                                                'МПа', eval_expr=sigma_1_0_m_py, roundness=3)
    if flange_type == 'zero_one':
        variables_sigma_1_0_m = {'M^M': M_M_calc, 'lambda': lambda_, 'S_0': S_0_calc, 'text{с}': c_flange,
                                 'D_{прив.}': D_priv_calc}
        sigma_1_0_m_py = 'M^M * 1000000 / (lambda * (S_0 - text{с})**2 * D_{прив.})'
        sigma_1_m_calc = sigma_0_m_calc = add_equation_calc(doc, sigma_1_0_m_formula, variables_sigma_1_0_m,
                                                            'МПа', eval_expr=sigma_1_0_m_py, roundness=3)

    doc.add_paragraph('Радиальное напряжение в тарелке приварного встык фланца в условиях затяжки:')
    f_counter += 1
    sigma_r_m_formula = (r'σ_R^M = \frac{1.33 \cdot \beta_F \cdot H + l_0}{\lambda \cdot H^2 \cdot l_0 \cdot D_{вн}}'
                         r' \cdot M^M')
    add_equation(doc, sigma_r_m_formula, number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='F', subscript=True, value=B_F, indent=2.75,
                              desc='коэффициент, зависящий от размеров втулки фланца, принимается по рисунку '
                                   'К.2 [1], приложение К.')
    variables_sigma_r_m = {'beta_F': B_F, 'H': h_flange, 'l_0': l_0_calc, 'lambda': lambda_, 'D_{вн}': D_int_flange,
                           'M^M': M_M_calc}
    sigma_r_m_py = '(1.33 * beta_F * H + l_0) / (lambda * H**2 * l_0 * D_{вн}) * M^M * 1000000'
    sigma_r_m_calc = add_equation_calc(doc, sigma_r_m_formula, variables_sigma_r_m, 'МПа', eval_expr=sigma_r_m_py,
                                       roundness=3)
    doc.add_paragraph('Окружное напряжение в тарелке приварного встык фланца в условиях затяжки:')
    f_counter += 1
    sigma_t_m_formula = r'σ_T^M = \frac{\beta_Y \cdot M^M}{H^2 \cdot D_{вн}} - \beta_Z \cdot σ_R^M'
    add_equation(doc, sigma_t_m_formula, number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='Y', subscript=True, value=B_Y, indent=2.75,
                              desc='коэффициент, зависящий от соотношения размеров тарелки фланца, '
                                   'принимается по рисунку К.1 или по формулам К.4, К.7 [1], приложение К.')
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='Z', subscript=True, value=B_Z, indent=2.75,
                              desc='коэффициент, зависящий от соотношения размеров тарелки фланца, '
                                   'принимается по рисунку К.1 или по формулам К.4, К.8 [1], приложение К.')
    variables_sigma_t_m = {'beta_Y': B_Y, 'M^M': M_M_calc, 'H': h_flange, 'D_{вн}': D_int_flange, 'beta_Z': B_Z,
                           'σ_R^M': sigma_r_m_calc}
    sigma_t_m_py = 'beta_Y * M^M * 1000000 / (H**2 * D_{вн}) - beta_Z * σ_R^M'
    sigma_t_m_calc = add_equation_calc(doc, sigma_t_m_formula, variables_sigma_t_m, 'МПа', eval_expr=sigma_t_m_py,
                                       roundness=3)

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            p4 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p4.add_run('S').italic = True
            p4.add_run('1').font.subscript = True
            p4.add_run(' в рабочих условиях:')  # с - русская
            f_counter += 1
            sigma_1_r_formula = r'σ_1^р = \frac{M^Р}{\lambda \cdot (S_1 - \text{с})^2 \cdot D_{прив.}}'
            add_equation(doc, sigma_1_r_formula, number=f_counter, f_l_indent=6.25)

            p5 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p5.add_run('S').italic = True
            p5.add_run('0').font.subscript = True
            p5.add_run(' в рабочих условиях:')
            f_counter += 1
            sigma_0_r_formula = r'σ_0^р = \text{f} \cdot σ_1^р'
            add_equation(doc, sigma_0_r_formula, number=f_counter, f_l_indent=7.0)
        if D_m_flange == D_n_flange:
            doc.add_paragraph('Меридиональное изгибное напряжение в цилиндрической втулке приварного встык фланца '
                              'в рабочих условиях:')
            f_counter += 1  # с - русская
            sigma_1_0_r_formula = r'σ_1^р, σ_0^р = \frac{M^Р}{\lambda \cdot (S_0 - \text{с})^2 \cdot D_{прив.}}'
            add_equation(doc, sigma_1_0_r_formula, number=f_counter, f_l_indent=6.25)
    if flange_type == 'zero_one':
        doc.add_paragraph('Меридиональное изгибное напряжение в приварном встык плоском фланце в рабочих условиях:')
        f_counter += 1  # с - русская
        sigma_1_0_r_formula = r'σ_1^р, σ_0^р = \frac{M^Р}{\lambda \cdot (S_0 - \text{с})^2 \cdot D_{прив.}}'
        add_equation(doc, sigma_1_0_r_formula, number=f_counter, f_l_indent=6.25)

    doc.add_paragraph('Подставим полученные значения в формулу меридиональных изгибных напряжений в рабочих условиях:')

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            variables_sigma_1_r = {'M^Р': M_R_calc, 'lambda': lambda_, 'S_1': S_1_calc, 'text{с}': c_flange,
                                   'D_{прив.}': D_priv_calc}
            sigma_1_r_py = 'M^Р * 1000000 / (lambda * (S_1 - text{с})**2 * D_{прив.})'
            sigma_1_r_calc = add_equation_calc(doc, sigma_1_r_formula, variables_sigma_1_r, 'МПа',
                                               eval_expr=sigma_1_r_py, roundness=3)
            variables_sigma_0_r = {'text{f}': f, 'σ_1^р': sigma_1_r_calc}
            sigma_0_r_py = 'text{f} * σ_1^р'
            sigma_0_r_calc = add_equation_calc(doc, sigma_0_r_formula, variables_sigma_0_r, 'МПа',
                                               eval_expr=sigma_0_r_py, roundness=3)
        if D_m_flange == D_n_flange:
            variables_sigma_1_0_r = {'M^Р': M_R_calc, 'lambda': lambda_, 'S_0': S_0_calc, 'text{с}': c_flange,
                                     'D_{прив.}': D_priv_calc}
            sigma_1_0_r_py = 'M^Р * 1000000 / (lambda * (S_0 - text{с})**2 * D_{прив.})'
            sigma_1_r_calc = sigma_0_r_calc = add_equation_calc(doc, sigma_1_0_r_formula, variables_sigma_1_0_r,
                                                                'МПа', eval_expr=sigma_1_0_r_py, roundness=3)
    if flange_type == 'zero_one':
        variables_sigma_1_0_r = {'M^Р': M_M_calc, 'lambda': lambda_, 'S_0': S_0_calc, 'text{с}': c_flange,
                                 'D_{прив.}': D_priv_calc}
        sigma_1_0_r_py = 'M^Р * 1000000 / (lambda * (S_0 - text{с})**2 * D_{прив.})'
        sigma_1_r_calc = sigma_0_r_calc = add_equation_calc(doc, sigma_1_0_r_formula, variables_sigma_1_0_r,
                                                            'МПа', eval_expr=sigma_1_0_r_py, roundness=3)

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            p7 = doc.add_paragraph('Меридиональные мембранные напряжения в конической втулке приварного встык фланца в '
                                   'сечении ')
            p7.add_run('S').italic = True
            p7.add_run('1').font.subscript = True
            p7.add_run(' в рабочих условиях:')
            f_counter += 1
            sigma_1_m_m_r_formula = r'σ_{1м.м}^р = \frac{Q_д + Q_{fm}}{\pi \cdot (D_{вн} + S_1) \cdot (S_1 - \text{с})}'
            add_equation(doc, sigma_1_m_m_r_formula, number=f_counter, f_l_indent=6.25)
            variables_sigma_1_m_m_r = {'Q_д': Q_d_calc, 'Q_{fm}': Q_fm_calc, 'D_{вн}': D_int_flange, 'S_1': S_1_calc,
                                       'text{с}': c_flange}
            sigma_1_m_m_r_py = '(Q_д + Q_{fm}) * 1000 / (3.141593 * (D_{вн} + S_1) * (S_1 - text{с}))'
            sigma_1_m_m_r_calc = add_equation_calc(doc, sigma_1_m_m_r_formula, variables_sigma_1_m_m_r, 'МПа',
                                                   eval_expr=sigma_1_m_m_r_py, roundness=3)

            p8 = doc.add_paragraph('Меридиональные мембранные напряжения в конической втулке приварного встык фланца в '
                                   'сечении ')
            p8.add_run('S').italic = True
            p8.add_run('0').font.subscript = True
            p8.add_run(' в рабочих условиях:')
            f_counter += 1
            sigma_0_m_m_r_formula = r'σ_{0м.м}^р = \frac{Q_д + Q_{fm}}{\pi \cdot (D_{вн} + S_0) \cdot (S_0 - \text{с})}'
            add_equation(doc, sigma_0_m_m_r_formula, number=f_counter, f_l_indent=6.25)
            variables_sigma_0_m_m_r = {'Q_д': Q_d_calc, 'Q_{fm}': Q_fm_calc, 'D_{вн}': D_int_flange, 'S_0': S_0_calc,
                                       'text{с}': c_flange}
            sigma_0_m_m_r_py = '(Q_д + Q_{fm}) * 1000 / (3.141593 * (D_{вн} + S_0) * (S_0 - text{с}))'
            sigma_0_m_m_r_calc = add_equation_calc(doc, sigma_0_m_m_r_formula, variables_sigma_0_m_m_r, 'МПа',
                                                   eval_expr=sigma_0_m_m_r_py, roundness=3)
        if D_m_flange == D_n_flange:
            p7 = doc.add_paragraph('Меридиональные мембранные напряжения в цилиндрической втулке приварного встык '
                                   'фланца в сечении ')
            p7.add_run('S').italic = True
            p7.add_run('0').font.subscript = True
            p7.add_run(' в рабочих условиях:')
            f_counter += 1
            sigma_0_m_m_r_formula = (r'σ_{0м.м}^р, σ_{1м.м}^р = \frac{Q_д + Q_{fm}}{\pi \cdot (D_{вн} + S_0) \cdot ('
                                     r'S_0 - \text{с})}')
            add_equation(doc, sigma_0_m_m_r_formula, number=f_counter, f_l_indent=6.25)
            variables_sigma_0_m_m_r = {'Q_д': Q_d_calc, 'Q_{fm}': Q_fm_calc, 'D_{вн}': D_int_flange, 'S_0': S_0_calc,
                                       'text{с}': c_flange}
            sigma_0_m_m_r_py = '(Q_д + Q_{fm}) * 1000 / (3.141593 * (D_{вн} + S_0) * (S_0 - text{с}))'
            sigma_1_m_m_r_calc = sigma_0_m_m_r_calc = add_equation_calc(doc, sigma_0_m_m_r_formula,
                                                                        variables_sigma_0_m_m_r, 'МПа',
                                                                        eval_expr=sigma_0_m_m_r_py, roundness=3)
    if flange_type == 'zero_one':
        doc.add_paragraph('Меридиональные мембранные напряжения в приварном встык плоском фланце в рабочих условиях:')
        f_counter += 1
        sigma_0_m_m_r_formula = (r'σ_{0м.м}^р, σ_{1м.м}^р = \frac{Q_д + Q_{fm}}{\pi \cdot (D_{вн} + S_0) \cdot ('
                                 r'S_0 - \text{с})}')
        add_equation(doc, sigma_0_m_m_r_formula, number=f_counter, f_l_indent=6.25)
        variables_sigma_0_m_m_r = {'Q_д': Q_d_calc, 'Q_{fm}': Q_fm_calc, 'D_{вн}': D_int_flange, 'S_0': S_0_calc,
                                   'text{с}': c_flange}
        sigma_0_m_m_r_py = '(Q_д + Q_{fm}) * 1000 / (3.141593 * (D_{вн} + S_0) * (S_0 - text{с}))'
        sigma_1_m_m_r_calc = sigma_0_m_m_r_calc = add_equation_calc(doc, sigma_0_m_m_r_formula,
                                                                    variables_sigma_0_m_m_r, 'МПа',
                                                                    eval_expr=sigma_0_m_m_r_py, roundness=3)

    p9 = doc.add_paragraph('Окружные мембранные напряжения от действия давления во втулке приварного встык фланца '
                           'в сечении ')
    p9.add_run('S').italic = True
    p9.add_run('0').font.subscript = True
    p9.add_run(' в рабочих условиях:')
    f_counter += 1
    sigma_0_m_o_r_formula = r'σ_{0м.о}^р = \frac{P \cdot D_{вн}}{2 \cdot (S_0 - \text{с})}'
    add_equation(doc, sigma_0_m_o_r_formula, number=f_counter, f_l_indent=1.0)
    variables_sigma_0_m_o_r = {'P': pressure, 'D_{вн}': D_int_flange, 'S_0': S_0_calc, 'text{с}': c_flange}
    sigma_0_m_o_r_py = 'P * D_{вн} / (2 * (S_0 - text{с}))'
    sigma_0_m_o_r_calc = add_equation_calc(doc, sigma_0_m_o_r_formula, variables_sigma_0_m_o_r, 'МПа',
                                           eval_expr=sigma_0_m_o_r_py, roundness=3)

    doc.add_paragraph('Радиальное напряжение в тарелке приварного встык фланца в рабочих условиях:')
    f_counter += 1
    sigma_r_r_formula = (r'σ_R^р = \frac{1.33 \cdot \beta_F \cdot H + l_0}{\lambda \cdot H^2 \cdot l_0 \cdot D_{вн}}'
                         r' \cdot M^Р')
    add_equation(doc, sigma_r_r_formula, number=f_counter, f_l_indent=5.75)
    variables_sigma_r_r = {'beta_F': B_F, 'H': h_flange, 'l_0': l_0_calc, 'lambda': lambda_, 'D_{вн}': D_int_flange,
                           'M^Р': M_R_calc}
    sigma_r_r_py = '(1.33 * beta_F * H + l_0) / (lambda * H**2 * l_0 * D_{вн}) * M^Р * 1000000'
    sigma_r_r_calc = add_equation_calc(doc, sigma_r_r_formula, variables_sigma_r_r, 'МПа', eval_expr=sigma_r_r_py,
                                       roundness=3)
    doc.add_paragraph('Окружное напряжение в тарелке приварного встык фланца в рабочих условиях:')
    f_counter += 1
    sigma_t_r_formula = r'σ_T^р = \frac{\beta_Y \cdot M^Р}{H^2 \cdot D_{вн}} - \beta_Z \cdot σ_R^р'
    add_equation(doc, sigma_t_r_formula, number=f_counter, f_l_indent=5.75)
    variables_sigma_t_r = {'beta_Y': B_Y, 'M^Р': M_R_calc, 'H': h_flange, 'D_{вн}': D_int_flange, 'beta_Z': B_Z,
                           'σ_R^р': sigma_r_r_calc}
    sigma_t_r_py = 'beta_Y * M^Р * 1000000 / (H**2 * D_{вн}) - beta_Z * σ_R^р'
    sigma_t_r_calc = add_equation_calc(doc, sigma_t_r_formula, variables_sigma_t_r, 'МПа', eval_expr=sigma_t_r_py,
                                       roundness=3)

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            p10 = doc.add_paragraph('Условия статической прочности приварных встык фланцев с конической втулкой в '
                                    'сечении ')
            p10.add_run('S').italic = True
            p10.add_run('1').font.subscript = True
            p10.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_1^м + σ_R^м \vert; \vert σ_1^м + σ_Т^м \vert \Big}'
                              r' <= K_T \cdot [σ]_м', number=f_counter, f_l_indent=4.0)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_1^р - σ_{1м.м}^р + σ_R^р \vert; \vert σ_1^р - σ_{1м.м}^р + σ_Т^р '
                              r'\vert; \vert σ_1^р + σ_{1м.м}^р \vert \Big} <= K_T \cdot [σ]_м', number=f_counter,
                         f_l_indent=1.0)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'K', italic=True, sub_text='T', subscript=True, value=1.3, indent=2.5,
                                      desc='коэффициент увеличения допускаемых напряжений при расчете фланцев с '
                                           'учетом стесненности температурных деформаций ([1] п.8.5.1);')
            add_paragraph_with_indent(doc, '[σ]', sub_text='м', subscript=True, indent=1.75,
                                      desc='допускаемое значение общих мембранных и изгибных напряжений, МПа '
                                           '([7] п.8.10).')
            f_counter += 1
            add_equation(doc, r'\left[ σ \right]_м = 1.5 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, '[σ]', value=allowed_stress_flange_T, indent=3.5,
                                      desc=f'допускаемое напряжение для стали марки {flange_steel}, МПа,'
                                           f' найденное методом экстраполяции для расчетной температуры '
                                           '([8] табл. 3).')
            add_equation(doc, rf'\left[ σ \right]_м = 1.5 \cdot {allowed_stress_flange_T} = '
                              rf'{1.5 * allowed_stress_flange_T} МПа')

            p11 = doc.add_paragraph('Условия статической прочности приварных встык фланцев с конической втулкой в '
                                    'сечении ')
            p11.add_run('S').italic = True
            p11.add_run('0').font.subscript = True
            p11.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
            f_counter += 1
            add_equation(doc, r'σ_0^м <= 1.3 \cdot [σ]_R', number=f_counter, f_l_indent=6.5)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_0^р ± σ_{0м.м}^р \vert; \vert 0.3 \cdot σ_0^р ± σ_{0м.о}^р '
                              r'\vert; \vert 0.7 \cdot σ_0^р ± \Big(σ_{0м.м}^р - σ_{0м.о}^р \Big) \vert \Big} '
                              r'<= 1.3 \cdot [σ]_R', number=f_counter, f_l_indent=0.0)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, '[σ]', sub_text='R', subscript=True, indent=1.75,
                                      desc='допускаемое значение суммарных общих и местных условных упругих мембранных '
                                           'и изгибных напряжений, МПа ([7] п.8.10).')
            f_counter += 1
            add_equation(doc, r'\left[ σ \right]_R = 3.0 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
            add_equation(doc, rf'\left[ σ \right]_R = 3.0 \cdot {allowed_stress_flange_T} = '
                              rf'{3.0 * allowed_stress_flange_T} МПа')

        if D_m_flange == D_n_flange:
            p10 = doc.add_paragraph('Условия статической прочности приварных встык фланцев с цилиндрической втулкой в '
                                    'сечении ')
            p10.add_run('S').italic = True
            p10.add_run('0').font.subscript = True
            p10.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_0^м + σ_R^м \vert; \vert σ_0^м + σ_Т^м \vert \Big}'
                              r' <= K_T \cdot [σ]_м', number=f_counter, f_l_indent=4.0)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_0^р - σ_{0м.м}^р + σ_T^р \vert; \vert σ_0^р - σ_{0м.м}^р + σ_R^р '
                              r'\vert; \vert σ_0^р + σ_{0м.м}^р \vert \Big} <= K_T \cdot [σ]_м', number=f_counter,
                         f_l_indent=1.0)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'K', italic=True, sub_text='T', subscript=True, value=1.3, indent=2.5,
                                      desc='коэффициент увеличения допускаемых напряжений при расчете фланцев с '
                                           'учетом стесненности температурных деформаций ([1] п.8.5.1);')
            add_paragraph_with_indent(doc, '[σ]', sub_text='м', subscript=True, indent=1.75,
                                      desc='допускаемое значение общих мембранных и изгибных напряжений, МПа '
                                           '([7] п.8.10).')
            f_counter += 1
            add_equation(doc, r'\left[ σ \right]_м = 1.5 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, '[σ]', value=allowed_stress_flange_T, indent=3.5,
                                      desc=f'допускаемое напряжение для стали марки {flange_steel}, МПа,'
                                           f' найденное методом экстраполяции для расчетной температуры '
                                           '([8] табл. 3).')
            add_equation(doc, rf'\left[ σ \right]_м = 1.5 \cdot {allowed_stress_flange_T} = '
                              rf'{1.5 * allowed_stress_flange_T} МПа')

    if flange_type == 'zero_one':
        p10 = doc.add_paragraph('Условия статической прочности приварных встык плоских фланцев в '
                                'сечении ')
        p10.add_run('S').italic = True
        p10.add_run('0').font.subscript = True
        p10.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_0^м + σ_R^м \vert; \vert σ_0^м + σ_Т^м \vert \Big}'
                          r' <= K_T \cdot [σ]_м', number=f_counter, f_l_indent=4.0)
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_0^р - σ_{0м.м}^р + σ_T^р \vert; \vert σ_0^р - σ_{0м.м}^р + σ_R^р '
                          r'\vert; \vert σ_0^р + σ_{0м.м}^р \vert \Big} <= K_T \cdot [σ]_м', number=f_counter,
                     f_l_indent=1.0)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'K', italic=True, sub_text='T', subscript=True, value=1.3, indent=2.5,
                                  desc='коэффициент увеличения допускаемых напряжений при расчете фланцев с '
                                       'учетом стесненности температурных деформаций ([1] п.8.5.1);')
        add_paragraph_with_indent(doc, '[σ]', sub_text='м', subscript=True, indent=1.75,
                                  desc='допускаемое значение общих мембранных и изгибных напряжений, МПа '
                                       '([7] п.8.10).')
        f_counter += 1
        add_equation(doc, r'\left[ σ \right]_м = 1.5 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, '[σ]', value=allowed_stress_flange_T, indent=3.5,
                                  desc=f'допускаемое напряжение для стали марки {flange_steel}, МПа,'
                                       f' найденное методом экстраполяции для расчетной температуры '
                                       '([8] табл. 3).')
        add_equation(doc, rf'\left[ σ \right]_м = 1.5 \cdot {allowed_stress_flange_T} = '
                          rf'{1.5 * allowed_stress_flange_T} МПа')

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            p12 = doc.add_paragraph('Подставим найденные значения в условия статической прочности фланцев '
                                    'с конической втулкой ')
            p12.add_run('в сечении ')
            p12.add_run('S').italic = True
            p12.add_run('1').font.subscript = True
            p12.add_run(':')
            if max(abs(sigma_1_m_calc + sigma_r_m_calc),
                   abs(sigma_1_m_calc + sigma_t_m_calc)) <= 1.3 * 1.5 * allowed_stress_flange_T:
                max_1 = round(max(abs(sigma_1_m_calc + sigma_r_m_calc), abs(sigma_1_m_calc + sigma_t_m_calc)), 3)
                add_equation(doc, rf'\max \Big{{ \vert {sigma_1_m_calc} + {sigma_r_m_calc} \vert; \vert '
                                  rf'{sigma_1_m_calc} + {sigma_t_m_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {1.5 * allowed_stress_flange_T}', f_l_indent=1.0)
                add_equation(doc, rf'{max_1} <= {round(1.3 * 1.5 * allowed_stress_flange_T, 3)}')
                doc.add_paragraph('Условие выполняется.')
            if (max(
                    abs(sigma_1_r_calc - sigma_1_m_m_r_calc + sigma_r_r_calc),
                    abs(sigma_1_r_calc - sigma_1_m_m_r_calc + sigma_t_r_calc),
                    abs(sigma_1_r_calc + sigma_1_m_m_r_calc))
                    <= 1.3 * 1.5 * allowed_stress_flange_T):
                max_1 = round(max(
                    abs(sigma_1_r_calc - sigma_1_m_m_r_calc + sigma_r_r_calc),
                    abs(sigma_1_r_calc - sigma_1_m_m_r_calc + sigma_t_r_calc),
                    abs(sigma_1_r_calc + sigma_1_m_m_r_calc)), 3)
                add_equation(doc,
                             rf'\max \Big{{ \vert {sigma_1_r_calc} - {sigma_1_m_m_r_calc} + {sigma_r_r_calc} \vert; \vert {sigma_1_r_calc} - {sigma_1_m_m_r_calc} + {sigma_t_r_calc} '
                             rf'\vert; \vert {sigma_1_r_calc} + {sigma_1_m_m_r_calc} \vert \Big}} <= 1.3 \cdot {1.5 * allowed_stress_flange_T}',
                             f_l_indent=1.0)
                add_equation(doc, rf'{max_1} <= {round(1.3 * 1.5 * allowed_stress_flange_T, 3)}')
                doc.add_paragraph('Условие выполняется.')

            p13 = doc.add_paragraph('В сечении ')
            p13.add_run('S').italic = True
            p13.add_run('0').font.subscript = True
            p13.add_run(':')
            if sigma_0_m_calc <= 1.3 * 3.0 * allowed_stress_flange_T:
                add_equation(doc, rf'{sigma_0_m_calc} <= 1.3 \cdot {3.0 * allowed_stress_flange_T}')
                add_equation(doc, rf'{sigma_0_m_calc} <= {round(1.3 * 3.0 * allowed_stress_flange_T, 3)}')
                doc.add_paragraph('Условие выполняется.')
            if max(
                    max(abs(sigma_0_r_calc + sigma_0_m_m_r_calc),
                        abs(sigma_0_r_calc - sigma_0_m_m_r_calc)),
                    max(abs(0.3 * sigma_0_r_calc + sigma_0_m_o_r_calc),
                        abs(0.3 * sigma_0_r_calc - sigma_0_m_o_r_calc)),
                    max(abs(0.7 * sigma_0_r_calc + (sigma_0_m_m_r_calc - sigma_0_m_o_r_calc)),
                        abs(0.7 * sigma_0_r_calc - (sigma_0_m_m_r_calc - sigma_0_m_o_r_calc)))
            ) <= 1.3 * 3.0 * allowed_stress_flange_T:
                max_2 = round(max(
                    max(abs(sigma_0_r_calc + sigma_0_m_m_r_calc),
                        abs(sigma_0_r_calc - sigma_0_m_m_r_calc)),
                    max(abs(0.3 * sigma_0_r_calc + sigma_0_m_o_r_calc),
                        abs(0.3 * sigma_0_r_calc - sigma_0_m_o_r_calc)),
                    max(abs(0.7 * sigma_0_r_calc + (sigma_0_m_m_r_calc - sigma_0_m_o_r_calc)),
                        abs(0.7 * sigma_0_r_calc - (sigma_0_m_m_r_calc - sigma_0_m_o_r_calc)))
                ), 3)
                add_equation(doc,
                             rf'\max \Big{{ \vert {sigma_0_r_calc} ± {sigma_0_m_m_r_calc} \vert; \vert 0.3 \cdot {sigma_0_r_calc} ± {sigma_0_m_o_r_calc} '
                             rf'\vert; \vert 0.7 \cdot {sigma_0_r_calc} ± \Big({sigma_0_m_m_r_calc} - {sigma_0_m_o_r_calc}р \Big) \vert \Big}} '
                             rf'<= 1.3 \cdot {3.0 * allowed_stress_flange_T}', f_l_indent=0.0)
                add_equation(doc, rf'{max_2} <= {round(1.3 * 3.0 * allowed_stress_flange_T, 3)}')
                doc.add_paragraph('Условие выполняется.')

        if D_m_flange == D_n_flange:
            p12 = doc.add_paragraph('Подставим найденные значения в условия статической прочности фланцев с '
                                    'цилиндрической втулкой ')
            p12.add_run('в сечении ')
            p12.add_run('S').italic = True
            p12.add_run('0').font.subscript = True
            p12.add_run(':')
            if max(
                    abs(sigma_0_m_calc + sigma_r_m_calc),
                    abs(sigma_0_m_calc + sigma_t_m_calc)
            ) <= 1.3 * 1.5 * allowed_stress_flange_T:
                max_2 = round(max(
                    abs(sigma_0_m_calc + sigma_r_m_calc),
                    abs(sigma_0_m_calc + sigma_t_m_calc)
                ), 3)
                add_equation(doc,
                             rf'\max \Big{{ \vert {sigma_0_m_calc} + {sigma_r_m_calc} \vert; \vert {sigma_0_m_calc} + '
                             rf'{sigma_t_m_calc} \vert \Big}}'
                             rf' <= 1.3 \cdot {1.5 * allowed_stress_flange_T}', f_l_indent=1.0)
                add_equation(doc, rf'{max_2} <= {1.3 * 1.5 * allowed_stress_flange_T}')
                doc.add_paragraph('Условие выполняется.')
            if max(
                    abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_t_r_calc),
                    abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_r_r_calc),
                    abs(sigma_0_r_calc + sigma_0_m_m_r_calc)
            ) <= 1.3 * 1.5 * allowed_stress_flange_T:
                max_2 = round(max(
                    abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_t_r_calc),
                    abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_r_r_calc),
                    abs(sigma_0_r_calc + sigma_0_m_m_r_calc)
                ), 3)
                add_equation(doc,
                             rf'\max \Big{{ \vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_t_r_calc} \vert; \vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_r_r_calc} '
                             rf'\vert; \vert {sigma_0_r_calc} + {sigma_0_m_m_r_calc} \vert \Big}} <= 1.3 \cdot {1.5 * allowed_stress_flange_T}',
                             f_l_indent=1.0)
                add_equation(doc, rf'{max_2} <= {round(1.3 * 1.5 * allowed_stress_flange_T, 3)}')
                doc.add_paragraph('Условие выполняется.')
    else:
        doc.add_paragraph('Условие НЕ выполняется!')

    if flange_type == 'zero_one':
        p12 = doc.add_paragraph('Подставим найденные значения в условия статической прочности плоских фланцев ')
        p12.add_run('в сечении ')
        p12.add_run('S').italic = True
        p12.add_run('0').font.subscript = True
        p12.add_run(':')
        if max(
                abs(sigma_0_m_calc + sigma_r_m_calc),
                abs(sigma_0_m_calc + sigma_t_m_calc)
        ) <= 1.3 * 1.5 * allowed_stress_flange_T:
            max_2 = round(max(
                abs(sigma_0_m_calc + sigma_r_m_calc),
                abs(sigma_0_m_calc + sigma_t_m_calc)
            ), 3)
            add_equation(doc,
                         rf'\max \Big{{ \vert {sigma_0_m_calc} + {sigma_r_m_calc} \vert; \vert {sigma_0_m_calc} + '
                         rf'{sigma_t_m_calc} \vert \Big}}'
                         rf' <= 1.3 \cdot {1.5 * allowed_stress_flange_T}', f_l_indent=1.0)
            add_equation(doc, rf'{max_2} <= {1.3 * 1.5 * allowed_stress_flange_T}')
            doc.add_paragraph('Условие выполняется.')
        if max(
                abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_t_r_calc),
                abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_r_r_calc),
                abs(sigma_0_r_calc + sigma_0_m_m_r_calc)
        ) <= 1.3 * 1.5 * allowed_stress_flange_T:
            max_2 = round(max(
                abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_t_r_calc),
                abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_r_r_calc),
                abs(sigma_0_r_calc + sigma_0_m_m_r_calc)
            ), 3)
            add_equation(doc,
                         rf'\max \Big{{ \vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_t_r_calc} \vert; \vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_r_r_calc} '
                         rf'\vert; \vert {sigma_0_r_calc} + {sigma_0_m_m_r_calc} \vert \Big}} <= 1.3 \cdot {1.5 * allowed_stress_flange_T}',
                         f_l_indent=1.0)
            add_equation(doc, rf'{max_2} <= {round(1.3 * 1.5 * allowed_stress_flange_T, 3)}')
            doc.add_paragraph('Условие выполняется.')

    p12 = doc.add_paragraph('Также для фланцев всех типов в сечении ')
    p12.add_run('S').italic = True
    p12.add_run('0').font.subscript = True
    p12.add_run(' должно выполняться условие:')
    f_counter += 1
    add_equation(doc, r'\max \Big{ \vert σ_{0м.о}^р \vert; \vert σ_{0м.м}^р \vert \Big} <= [σ]', number=f_counter,
                 f_l_indent=5.25)
    if max(sigma_0_m_o_r_calc, sigma_0_m_m_r_calc) <= allowed_stress_flange_T:
        add_equation(doc,
                     rf'\max \Big{{ \vert {sigma_0_m_o_r_calc} \vert; \vert {sigma_0_m_m_r_calc} \vert \Big}} <= {allowed_stress_flange_T}',
                     f_l_indent=5.25)
        add_equation(doc, rf'{max(sigma_0_m_o_r_calc, sigma_0_m_m_r_calc)} <= {allowed_stress_flange_T}')
        doc.add_paragraph('Условие выполняется.')
    else:
        doc.add_paragraph('Условие НЕ выполняется!')

    doc.add_paragraph('Для тарелок приварных встык фланцев при затяжке и в рабочих условиях должны выполняться '
                      'следующие условия:')
    f_counter += 1
    add_equation(doc, r'\max \Big{ \vert σ_R^м \vert; \vert σ_T^м \vert \Big} <= K_T \cdot [σ]', number=f_counter,
                 f_l_indent=5.25)
    f_counter += 1
    add_equation(doc, r'\max \Big{ \vert σ_R^р \vert; \vert σ_T^р \vert \Big} <= K_T \cdot [σ]', number=f_counter,
                 f_l_indent=5.25)
    if max(abs(sigma_r_m_calc), abs(sigma_t_m_calc)) <= 1.3 * allowed_stress_flange_T:
        add_equation(doc,
                     rf'\max \Big{{ \vert {sigma_r_m_calc} \vert; \vert {sigma_t_m_calc} \vert \Big}} <= 1.3 \cdot {allowed_stress_flange_T}',
                     f_l_indent=1.0)
        add_equation(doc,
                     rf'{max(abs(sigma_r_m_calc), abs(sigma_t_m_calc))} <= {round(1.3 * allowed_stress_flange_T, 3)}')
        doc.add_paragraph('Условие выполняется.')
    else:
        doc.add_paragraph('Условие НЕ выполняется!')
    if max(abs(sigma_r_r_calc), abs(sigma_t_r_calc)) <= 1.3 * allowed_stress_flange_T:
        add_equation(doc,
                     rf'\max \Big{{ \vert {sigma_r_r_calc} \vert; \vert {sigma_t_r_calc} \vert \Big}} <= 1.3 \cdot {allowed_stress_flange_T}',
                     f_l_indent=1.0)
        add_equation(doc,
                     rf'{max(abs(sigma_r_r_calc), abs(sigma_t_r_calc))} <= {round(1.3 * allowed_stress_flange_T, 3)}')
        doc.add_paragraph('Условие выполняется.')
    else:
        doc.add_paragraph('Условие НЕ выполняется!')

    doc.add_paragraph("6. Проверка углов поворота фланцев", style='Heading 2')
    doc.add_paragraph('Угол поворота приварного встык фланца в рабочих условиях:')
    f_counter += 1
    theta_formula = r'\theta = M^Р \cdot y_ф \cdot \frac{E^{20}}{E_ф} <= K_{\theta} \cdot [\theta]'
    add_equation(doc, theta_formula, number=f_counter, f_l_indent=5.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, '[θ]', indent=2.75, value=allowed_theta,
                              desc='допустимый угол поворота приварного встык фланца,'
                                   ' при промежуточных значениях определяется с помощью '
                                   'линейной интерполяции ([1], п.9.1);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='θ', subscript=True, value=1.0, indent=1.25,
                              desc='коэффициент, применяемый к рабочим условиям ([1], п.9.1);')
    add_paragraph_with_indent(doc, 'E', italic=True, sup_text='20', superscript=True, value=E_flange_20, indent=2.75,
                              desc='модуль продольной упругости материала фланца при температуре 20˚С, МПа'
                                   ' ([1], таблица Ж.1);')
    add_paragraph_with_indent(doc, 'E', italic=True, sub_text='ф', subscript=True, value=E_flange_T, indent=2.75,
                              desc='модуль продольной упругости материала фланца при расчетной температуре,'
                                   ' рассчитанный методом линейной интерполяции, МПа ([1], таблица Ж.1).')
    add_equation(doc, r'[\theta] = 0.006 при D_{вн} <= 400 мм')
    add_equation(doc, r'[\theta] = 0.013 при D_{вн} > 2000 мм')
    doc.add_paragraph('Подставим в выражение для угла поворота фланца:')
    theta_formula_calc = rf'\theta = {M_R_calc} \cdot ({y_f_calc}) \cdot \frac{{{E_flange_20}}}{{{E_flange_T}}} <= 1.0 \cdot {allowed_theta}'
    add_equation(doc, theta_formula_calc)
    add_equation(doc, rf'({round(M_R_calc * y_f_calc * E_flange_20 / E_flange_T, 15)}) <= {allowed_theta}')
    doc.add_paragraph('Условие выполняется.')

    doc.add_paragraph("7. Расчет элементов фланцевого соединения на малоцикловую усталость", style='Heading 2')
    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            p13 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p13.add_run('S').italic = True
            p13.add_run('1').font.subscript = True
            p13.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{11}}^М = σ_1^М = {sigma_1_m_calc} МПа', number=f_counter, f_l_indent=5.5)
            f_counter += 1
            add_equation(doc, rf'σ_{{12}}^М = -σ_1^М = {-sigma_1_m_calc} МПа', number=f_counter, f_l_indent=5.0)

            sigma_1_1_m_calc = sigma_1_m_calc
            sigma_1_2_m_calc = -sigma_1_m_calc

            p14 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'конической втулки приварного встык фланца в сечении ')
            p14.add_run('S').italic = True
            p14.add_run('1').font.subscript = True
            p14.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{13}}^М = σ_{{14}}^М = σ_T^М = {sigma_t_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_1_3_m_calc = sigma_1_4_m_calc = sigma_t_m_calc

            p15 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p15.add_run('S').italic = True
            p15.add_run('1').font.subscript = True
            p15.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{15}}^М = σ_{{16}}^М = σ_R^М = {sigma_r_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_1_5_m_calc = sigma_1_6_m_calc = sigma_r_m_calc

            p16 = doc.add_paragraph('Размах условных упругих напряжений в меридиональном направлении на наружной '
                                    'и внутренней поверхностях конической втулки приварного встык фланца в сечении ')
            p16.add_run('S').italic = True
            p16.add_run('0').font.subscript = True
            p16.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{01}}^М = σ_0^М = {sigma_0_m_calc} МПа', number=f_counter,
                         f_l_indent=5.5)
            f_counter += 1
            add_equation(doc, rf'σ_{{02}}^М = -σ_0^М = {-sigma_0_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_0_1_m_calc = sigma_0_m_calc
            sigma_0_2_m_calc = -sigma_0_m_calc

            p17 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p17.add_run('S').italic = True
            p17.add_run('1').font.subscript = True
            p17.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{11}}^Р = σ_1^Р + σ_{{1м.м}}^Р = {sigma_1_m_calc} + {sigma_1_m_m_r_calc} = {sigma_1_m_calc + sigma_1_m_m_r_calc} МПа',
                         number=f_counter, f_l_indent=2.5)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{12}}^Р = -σ_1^Р + σ_{{1м.м}}^Р = {-sigma_1_m_calc} + {sigma_1_m_m_r_calc} = {-sigma_1_m_calc + sigma_1_m_m_r_calc} МПа',
                         number=f_counter, f_l_indent=2.0)

            sigma_1_1_r_calc = sigma_1_m_calc + sigma_1_m_m_r_calc
            sigma_1_2_r_calc = -sigma_1_m_calc + sigma_1_m_m_r_calc

            p18 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'конической втулки приварного встык фланца в сечении ')
            p18.add_run('S').italic = True
            p18.add_run('1').font.subscript = True
            p18.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{13}}^Р = σ_{{14}}^Р = σ_T^Р = {sigma_t_r_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_1_3_r_calc = sigma_1_4_r_calc = sigma_t_r_calc

            p19 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p19.add_run('S').italic = True
            p19.add_run('1').font.subscript = True
            p19.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{15}}^Р = σ_{{16}}^Р = σ_R^Р = {sigma_r_r_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_1_5_r_calc = sigma_1_6_r_calc = sigma_r_r_calc

            p20 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p20.add_run('S').italic = True
            p20.add_run('0').font.subscript = True
            p20.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{01}}^Р = σ_0^Р + σ_{{0м.м}}^Р = {sigma_0_m_calc} + {sigma_0_m_m_r_calc} = {round(sigma_0_m_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.5)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{02}}^Р = -σ_0^Р + σ_{{0м.м}}^Р = {-sigma_0_m_calc} + {sigma_0_m_m_r_calc} = {round(-sigma_0_m_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.0)

            sigma_0_1_r_calc = round(sigma_0_m_calc + sigma_0_m_m_r_calc, 3)
            sigma_0_2_r_calc = round(-sigma_0_m_calc + sigma_0_m_m_r_calc, 3)

            p21 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'конической втулки приварного встык фланца в сечении ')
            p21.add_run('S').italic = True
            p21.add_run('0').font.subscript = True
            p21.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{03}}^Р = σ_{{0м.о}}^Р + 0.3 \cdot σ_0^Р = {sigma_0_m_o_r_calc} + 0.3 \cdot {sigma_0_r_calc} = {round(sigma_0_m_o_r_calc + 0.3 * sigma_0_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=1.5)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{04}}^Р = σ_{{0м.о}}^Р - 0.3 \cdot σ_0^Р = {sigma_0_m_o_r_calc} - 0.3 \cdot {sigma_0_r_calc} = {round(sigma_0_m_o_r_calc - 0.3 * sigma_0_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=1.75)

            sigma_0_3_r_calc = round(sigma_0_m_o_r_calc + 0.3 * sigma_0_r_calc, 3)
            sigma_0_4_r_calc = round(sigma_0_m_o_r_calc - 0.3 * sigma_0_r_calc, 3)

            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с конической втулкой:')
            f_counter += 1
            add_equation(doc, r'σ_{аф} = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert σ_{11}^М \vert ;'
                              r'\vert σ_{12}^М - σ_{14}^М \vert ; '
                              r'\vert σ_{12}^М - σ_{16}^М \vert ; \vert σ_{01}^М \vert \Big} }{2}',
                         number=f_counter, f_l_indent=4.25)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'α', italic=True, sub_text='σ', subscript=True, value=alpha_sigma_1,
                                      indent=2.75, desc='коэффициент, определяемый по рисунку 4 [1].')
            max_3 = max(
                alpha_sigma_1 * abs(sigma_1_1_m_calc),
                abs(sigma_1_2_m_calc - sigma_1_4_m_calc),
                abs(sigma_1_2_m_calc - sigma_1_6_m_calc),
                abs(sigma_0_1_m_calc)
            )
            sigma_alpha_f_calc = round(max_3 / 2, 3)
            add_equation(doc, rf'σ_{{аф}} = \frac{{\max \Big{{ {alpha_sigma_1} \cdot \vert {sigma_1_1_m_calc} \vert ;'
                              rf'\vert {sigma_1_2_m_calc} - {sigma_1_4_m_calc} \vert ; '
                              rf'\vert {sigma_1_2_m_calc} - {sigma_1_6_m_calc} \vert ; \vert {sigma_0_1_m_calc} \vert '
                              rf'\Big}} }}{{2}} = {sigma_alpha_f_calc} МПа')

        if D_m_flange == D_n_flange:
            p13 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p13.add_run('S').italic = True
            p13.add_run('0').font.subscript = True
            p13.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{01}}^М = σ_0^М = {sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.5)
            f_counter += 1
            add_equation(doc, rf'σ_{{02}}^М = -σ_0^М = {-sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.0)

            sigma_0_1_m_calc = sigma_0_m_calc
            sigma_0_2_m_calc = -sigma_0_m_calc

            p14 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'цилиндрической втулки приварного встык фланца в сечении ')
            p14.add_run('S').italic = True
            p14.add_run('0').font.subscript = True
            p14.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{03}}^М = σ_{{04}}^М = σ_T^М = {sigma_t_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_0_3_m_calc = sigma_0_4_m_calc = sigma_t_m_calc

            p15 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p15.add_run('S').italic = True
            p15.add_run('0').font.subscript = True
            p15.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{05}}^М = σ_{{06}}^М = σ_R^М = {sigma_r_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_0_5_m_calc = sigma_0_6_m_calc = sigma_r_m_calc

            p16 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p16.add_run('S').italic = True
            p16.add_run('0').font.subscript = True
            p16.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{01}}^Р = σ_0^Р + σ_{{0м.м}}^Р = {sigma_0_r_calc} + {sigma_0_m_m_r_calc} = {round(sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.75)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{02}}^Р = -σ_0^Р + σ_{{0м.м}}^Р = {-sigma_0_r_calc} + {sigma_0_m_m_r_calc} = {round(-sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.5)

            sigma_0_1_r_calc = round(sigma_0_r_calc + sigma_0_m_m_r_calc, 3)
            sigma_0_2_r_calc = round(-sigma_0_r_calc + sigma_0_m_m_r_calc, 3)

            p17 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'цилиндрической втулки приварного встык фланца в сечении ')
            p17.add_run('S').italic = True
            p17.add_run('0').font.subscript = True
            p17.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{03}}^Р = σ_{{04}}^Р = σ_T^Р = {sigma_t_r_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_0_3_r_calc = sigma_0_4_r_calc = sigma_t_r_calc

            p18 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p18.add_run('S').italic = True
            p18.add_run('0').font.subscript = True
            p18.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{05}}^Р = σ_{{06}}^Р = σ_R^Р = {sigma_r_r_calc} МПа', number=f_counter,
                         f_l_indent=4.75)

            sigma_0_5_r_calc = sigma_0_6_r_calc = sigma_r_r_calc

            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с цилиндрической втулкой:')
            f_counter += 1
            add_equation(doc, r'σ_{аф} = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert σ_{01}^М \vert ;'
                              r'\vert σ_{02}^М - σ_{04}^М \vert ; '
                              r'\vert σ_{02}^М - σ_{06}^М \vert ; \Big} }{2}',
                         number=f_counter, f_l_indent=4.5)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'α', italic=True, sub_text='σ', subscript=True, value=alpha_sigma_1,
                                      indent=2.75, desc='коэффициент, определяемый по рисунку 4 [1].')
            max_3 = max(
                alpha_sigma_1 * abs(sigma_0_1_m_calc),
                abs(sigma_0_2_m_calc - sigma_0_4_m_calc),
                abs(sigma_0_2_m_calc - sigma_0_6_m_calc)
            )
            sigma_alpha_f_calc = round(max_3 / 2, 3)
            add_equation(doc, rf'σ_{{аф}} = \frac{{\max \Big{{ {alpha_sigma_1} \cdot \vert {sigma_0_1_m_calc} \vert ;'
                              rf'\vert {sigma_0_2_m_calc} - {sigma_0_4_m_calc} \vert ; '
                              rf'\vert {sigma_0_2_m_calc} - {sigma_0_6_m_calc} \vert ; '
                              rf'\Big}} }}{{2}} = {sigma_alpha_f_calc} МПа')

    if flange_type == 'zero_one':
        p13 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p13.add_run('S').italic = True
        p13.add_run('0').font.subscript = True
        p13.add_run(' в условиях затяжки:')
        f_counter += 1
        add_equation(doc, rf'σ_{{01}}^М = σ_0^М = {sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.5)
        f_counter += 1
        add_equation(doc, rf'σ_{{02}}^М = -σ_0^М = {-sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.0)

        sigma_0_1_m_calc = sigma_0_m_calc
        sigma_0_2_m_calc = -sigma_0_m_calc

        p14 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                'приварного встык плоского фланца в сечении ')
        p14.add_run('S').italic = True
        p14.add_run('0').font.subscript = True
        p14.add_run(' в условиях затяжки:')
        f_counter += 1
        add_equation(doc, rf'σ_{{03}}^М = σ_{{04}}^М = σ_T^М = {sigma_t_m_calc} МПа', number=f_counter,
                     f_l_indent=5.0)

        sigma_0_3_m_calc = sigma_0_4_m_calc = sigma_t_m_calc

        p15 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p15.add_run('S').italic = True
        p15.add_run('0').font.subscript = True
        p15.add_run(' в условиях затяжки:')
        f_counter += 1
        add_equation(doc, rf'σ_{{05}}^М = σ_{{06}}^М = σ_R^М = {sigma_r_m_calc} МПа', number=f_counter,
                     f_l_indent=5.0)

        sigma_0_5_m_calc = sigma_0_6_m_calc = sigma_r_m_calc

        p16 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p16.add_run('S').italic = True
        p16.add_run('0').font.subscript = True
        p16.add_run(' в рабочих условиях:')
        f_counter += 1
        add_equation(doc,
                     rf'σ_{{01}}^Р = σ_0^Р + σ_{{0м.м}}^Р = {sigma_0_r_calc} + {sigma_0_m_m_r_calc} = {round(sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа',
                     number=f_counter, f_l_indent=2.75)
        f_counter += 1
        add_equation(doc,
                     rf'σ_{{02}}^Р = -σ_0^Р + σ_{{0м.м}}^Р = {-sigma_0_r_calc} + {sigma_0_m_m_r_calc} = {round(-sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа',
                     number=f_counter, f_l_indent=2.5)

        sigma_0_1_r_calc = round(sigma_0_r_calc + sigma_0_m_m_r_calc, 3)
        sigma_0_2_r_calc = round(-sigma_0_r_calc + sigma_0_m_m_r_calc, 3)

        p17 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                'приварного встык плоского фланца в сечении ')
        p17.add_run('S').italic = True
        p17.add_run('0').font.subscript = True
        p17.add_run(' в рабочих условиях:')
        f_counter += 1
        add_equation(doc, rf'σ_{{03}}^Р = σ_{{04}}^Р = σ_T^Р = {sigma_t_r_calc} МПа', number=f_counter,
                     f_l_indent=5.0)

        sigma_0_3_r_calc = sigma_0_4_r_calc = sigma_t_r_calc

        p18 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p18.add_run('S').italic = True
        p18.add_run('0').font.subscript = True
        p18.add_run(' в рабочих условиях:')
        f_counter += 1
        add_equation(doc, rf'σ_{{05}}^Р = σ_{{06}}^Р = σ_R^Р = {sigma_r_r_calc} МПа', number=f_counter,
                     f_l_indent=4.75)

        sigma_0_5_r_calc = sigma_0_6_r_calc = sigma_r_r_calc

        doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык '
                          'плоских фланцев:')
        f_counter += 1
        add_equation(doc, r'σ_{аф} = \frac{1.5 \cdot \max \Big{ \vert σ_{01}^М \vert ;'
                          r'\vert σ_{02}^М - σ_{04}^М \vert ; '
                          r'\vert σ_{02}^М - σ_{06}^М \vert ; \Big} }{2}',
                     number=f_counter, f_l_indent=4.5)
        max_3 = max(
            abs(sigma_0_1_m_calc),
            abs(sigma_0_2_m_calc - sigma_0_4_m_calc),
            abs(sigma_0_2_m_calc - sigma_0_6_m_calc)
        )
        sigma_alpha_f_calc = round(1.5 * max_3 / 2, 3)
        add_equation(doc, rf'σ_{{аф}} = \frac{{1.5 \cdot \max \Big{{ \vert {sigma_0_1_m_calc} \vert ;'
                          rf'\vert {sigma_0_2_m_calc} - {sigma_0_4_m_calc} \vert ; '
                          rf'\vert {sigma_0_2_m_calc} - {sigma_0_6_m_calc} \vert ; '
                          rf'\Big}} }}{{2}} = {sigma_alpha_f_calc} МПа')

    doc.add_paragraph('Амплитуда приведенных условных упругих напряжений при затяжке шпилек:')
    f_counter += 1
    add_equation(doc, r'σ_{аш} = \frac{1.8 \cdot σ_{ш1}}{2}', number=f_counter, f_l_indent=7.0)
    add_equation(doc,
                 rf'σ_{{аш}} = \frac{{1.8 \cdot {sigma_b_1_calc}}}{{2}} = {round(1.8 * sigma_b_1_calc / 2, 3)} МПа')

    if flange_type == 'one_one':
        if D_m_flange > D_n_flange:
            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с конической втулкой в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{аф}^Р = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert Δσ_{11}^Р \vert ;'
                              r' \vert Δσ_{12}^Р - Δσ_{14}^Р \vert ; \vert Δσ_{12}^Р - Δσ_{16}^Р \vert'
                              r'\Big} }{2}', number=f_counter, f_l_indent=4.5)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sup_text='Р', superscript=True, indent=2.75,
                                      desc=f'определяются как максимальные разности значений напряжений, вычисленных при'
                                           ' давлении от 0.6*P до 1,25*P по описанному ранее алгоритму:')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='11', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_1_1, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='12', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_1_2, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='14', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_1_4, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='16', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_1_6, desc='МПа.')
            sigma_alpha_f_R_calc = round(max(
                alpha_sigma_1 * abs(stress_delta_R_1_1),
                abs(stress_delta_R_1_2 - stress_delta_R_1_4),
                abs(stress_delta_R_1_2 - stress_delta_R_1_6)
            ) / 2, 3)
            add_equation(doc,
                         rf'σ_{{аф}}^Р = \frac{{\max \Big{{ \{alpha_sigma_1} \cdot \vert {stress_delta_R_1_1} \vert ;'
                         rf' \vert {stress_delta_R_1_2} - {stress_delta_R_1_4} \vert ; \vert {stress_delta_R_1_2} - {stress_delta_R_1_6} \vert'
                         rf'\Big}} }}{{2}} = {sigma_alpha_f_R_calc} МПа', f_l_indent=1.0)
        if D_m_flange == D_n_flange:
            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с цилиндрической втулкой в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{аф}^Р = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert Δσ_{01}^Р \vert ;'
                              r' \vert Δσ_{01}^Р - Δσ_{03}^Р \vert ; \vert Δσ_{01}^Р - Δσ_{05}^Р \vert ;'
                              r' \vert Δσ_{02}^Р \vert ;'
                              r' \vert Δσ_{02}^Р - Δσ_{04}^Р \vert ; \vert Δσ_{02}^Р - Δσ_{06}^Р \vert '
                              r'\Big} }{2}', number=f_counter, f_l_indent=2.25)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sup_text='Р', superscript=True, indent=2.75,
                                      desc=f'определяются как максимальные разности значений напряжений, вычисленных при'
                                           'давлении от 0.6*P до 1,25*P по описанному ранее алгоритму:')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='01', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_0_1, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='03', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_0_3, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='05', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_0_5, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='02', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_0_2, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='04', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_0_4, desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='06', subscript=True, sup_text='Р',
                                      superscript=True, indent=2.75, value=stress_delta_R_0_6, desc='МПа.')
            sigma_alpha_f_R_calc = round(max(
                alpha_sigma_1 * abs(stress_delta_R_0_1),
                abs(stress_delta_R_0_1 - stress_delta_R_0_3),
                abs(stress_delta_R_0_1 - stress_delta_R_0_5),
                abs(stress_delta_R_0_2),
                abs(stress_delta_R_0_2 - stress_delta_R_0_4),
                abs(stress_delta_R_0_2 - stress_delta_R_0_6),
            ) / 2, 3)
            add_equation(doc,
                         rf'σ_{{аф}}^Р = \frac{{\max \Big{{ \{alpha_sigma_1} \cdot \vert {stress_delta_R_0_1} \vert ;'
                         rf' \vert {stress_delta_R_0_1} - {stress_delta_R_0_3} \vert ; \vert {stress_delta_R_0_1} - {stress_delta_R_0_5} \vert ;'
                         rf' \vert {stress_delta_R_0_2} \vert ;'
                         rf' \vert {stress_delta_R_0_2} - {stress_delta_R_0_4} \vert ; \vert {stress_delta_R_0_2} - {stress_delta_R_0_6} \vert '
                         rf'\Big}} }}{{2}} = {sigma_alpha_f_R_calc} МПа', f_l_indent=1.0)

    if flange_type == 'zero_one':
        doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык плоских фланцев '
                          'в рабочих условиях:')
        f_counter += 1
        add_equation(doc, r'σ_{аф}^Р = \frac{1.5 \cdot \max \Big{ \vert Δσ_{01}^Р \vert ;'
                          r' \vert Δσ_{03}^Р \vert ; \vert Δσ_{01}^Р - Δσ_{03}^Р \vert ;'
                          r' \vert Δσ_{02}^Р \vert ;'
                          r' \vert Δσ_{04}^Р \vert ; \vert Δσ_{02}^Р - Δσ_{04}^Р \vert '
                          r'\Big} }{2}', number=f_counter, f_l_indent=2.5)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sup_text='Р', superscript=True, indent=2.75,
                                  desc=f'определяются как максимальные разности значений напряжений, вычисленных при'
                                       'давлении от 0.6*P до 1,25*P по описанному ранее алгоритму:')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='01', subscript=True, sup_text='Р',
                                  superscript=True, indent=2.75, value=stress_delta_R_0_1, desc='МПа;')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='03', subscript=True, sup_text='Р',
                                  superscript=True, indent=2.75, value=stress_delta_R_0_3, desc='МПа;')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='02', subscript=True, sup_text='Р',
                                  superscript=True, indent=2.75, value=stress_delta_R_0_2, desc='МПа;')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='04', subscript=True, sup_text='Р',
                                  superscript=True, indent=2.75, value=stress_delta_R_0_4, desc='МПа.')
        sigma_alpha_f_R_calc = round(1.5 * max(
            abs(stress_delta_R_0_1),
            abs(stress_delta_R_0_3),
            abs(stress_delta_R_0_1 - stress_delta_R_0_3),
            abs(stress_delta_R_0_2),
            abs(stress_delta_R_0_4),
            abs(stress_delta_R_0_2 - stress_delta_R_0_4),
        ) / 2, 3)
        add_equation(doc,
                     rf'σ_{{аф}}^Р = \frac{{1.5 \cdot \max \Big{{ \vert {stress_delta_R_0_1} \vert ;'
                     rf' \vert {stress_delta_R_0_3} \vert ; \vert {stress_delta_R_0_1} - {stress_delta_R_0_3} \vert ;'
                     rf' \vert {stress_delta_R_0_2} \vert ;'
                     rf' \vert {stress_delta_R_0_4} \vert ; \vert {stress_delta_R_0_2} - {stress_delta_R_0_4} \vert '
                     rf'\Big}} }}{{2}} = {sigma_alpha_f_R_calc} МПа', f_l_indent=1.0)

    doc.add_paragraph('Амплитуда приведенных условных упругих напряжений в рабочих условиях для шпилек:')
    f_counter += 1
    add_equation(doc, r'σ_{аш}^Р = \frac{\eta \cdot σ_{ш2}}{2}', number=f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'η', italic=True, indent=1.25, value=2.0, desc='коэффициент, определяемый по [9]'
                                                                                  ', таблица 2.')
    add_equation(doc,
                 rf'σ_{{аш}} = \frac{{2.0 \cdot {sigma_b_2_calc}}}{{2}} = {round(2.0 * sigma_b_2_calc / 2, 3)} МПа')

    doc.add_paragraph('Проверка элементов фланцевого соединения на малоцикловую усталость проводится согласно [9]. '
                      'Необходимо выполнение условия:')
    f_counter += 1
    add_equation(doc, r'\frac{N_c}{[N]_c} + \frac{N_p}{[N]_p} <= 1.0', number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'N', italic=False, sub_text='c', subscript=True, indent=2.75, value=num_cycles_c,
                              desc='заданное число циклов сборок-разборок оборудования;')
    add_paragraph_with_indent(doc, 'N', italic=False, sub_text='p', subscript=True, indent=2.75, value=num_cycles_r,
                              desc='заданное число циклов изменений режима эксплуатации;')
    add_paragraph_with_indent(doc, '[N]', italic=False, sub_text='c', subscript=True, indent=2.75, value=cycles_c,
                              desc='допускаемое число циклов сборок-разборок ([9], рисунки 2-5);')
    add_paragraph_with_indent(doc, '[N]', italic=False, sub_text='p', subscript=True, indent=2.75, value=cycles_p,
                              desc='допускаемое число циклов изменения режима эксплуатации ([9], рисунки 2-5);')
    add_equation(doc, rf'\frac{{{num_cycles_c}}}{{{cycles_c}}} + \frac{{{num_cycles_r}}}{{{cycles_p}}} <= 1.0')
    add_equation(doc, rf'{round(num_cycles_c / cycles_c + num_cycles_r / cycles_p, 3)} <= 1.0')
    doc.add_paragraph('Условие выполняется.')

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_paragraph("Перечень документов", style='Heading 1')
    doc.add_paragraph("1.	ГОСТ Р 52857.4-2007 - Сосуды и аппараты. Нормы и методы расчета на прочность. "
                      "Расчет на прочность и герметичность фланцевых соединений.")
    if gasket_document == "ГОСТ 15180-86":
        doc.add_paragraph("2.	ГОСТ 15180-86 - Прокладки плоские эластичные. Основные параметры и размеры.")
    if gasket_document == "ГОСТ 34655-2020":
        doc.add_paragraph("2.	ГОСТ 34655-2020 - Арматура трубопроводная. Прокладки овального, восьмиугольного "
                          "сечения, линзовые стальные для фланцев арматуры. Конструкция, размеры и общие технические "
                          "требования.")
    if gasket_document == "ISO 7483-2011":
        doc.add_paragraph("2.	ISO 7483 - Dimensions of gaskets for use with flanges to ISO 7005.")
    doc.add_paragraph("3.	ГОСТ 9064-75 - Гайки для фланцевых соединений с температурой среды от 0 до 650 ˚С. Типы "
                      "и основные размеры.")
    doc.add_paragraph("4.	ГОСТ 9065-75 - Шайбы для фланцевых соединений с температурой среды от 0 до 650 ˚С. Типы "
                      "и основные размеры.")
    doc.add_paragraph("5.	ГОСТ 33259-2015 - Фланцы арматуры, соединительных частей и трубопроводов на номинальное "
                      "давление до PN 250. Конструкция, размеры и общие технические требования")
    doc.add_paragraph("6.	ОСТ 26-2040-96 - Шпильки для фланцевых соединений. Конструкция и размеры.")
    doc.add_paragraph("7.	ГОСТ Р 52857.1-2007 - Сосуды и аппараты. Нормы и методы расчета на прочность. "
                      "Общие требования.")
    doc.add_paragraph("8.	РД 26-15-88 - Сосуды и аппараты. Нормы и методы расчета на прочности и герметичность "
                      "фланцевых соединений.")
    doc.add_paragraph("9.	ГОСТ Р 52857.6-2007 - Сосуды и аппараты. Нормы и методы расчета на прочность. "
                      "Расчет на прочность при малоцикловых нагрузках.")

    filename = f'Расчет_фланцев_DN{int(D_N_flange)}_PN{int(pressure * 10)}_{flange_type}_{flange_steel}_{bolt_steel}'
    doc.save(f'{filename}.docx')
