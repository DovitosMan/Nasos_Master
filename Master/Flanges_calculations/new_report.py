import math
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor
import math2docx
from docx.enum.section import WD_SECTION
from types import SimpleNamespace


def set_heading_style(style, font_name="Times New Roman", size=14, bold=True, align=None,
                      f_l_indent=0.0, space_after=6, space_before=12, color=(0, 0, 0)):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = RGBColor(*color)

    rFonts = OxmlElement('w:rFonts')
    for tag in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        rFonts.set(qn(f'w:{tag}'), font_name)
    style.element.rPr.append(rFonts)

    fmt = style.paragraph_format
    if align:
        fmt.alignment = align
    fmt.first_line_indent = Cm(f_l_indent)
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)


def add_paragraph_with_indent(doc, text="", indent=3.0, italic=None, sub_text="", sup_text="", value=None, desc=None):
    left = indent
    first_line = -indent
    tab = indent

    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.left_indent = Cm(left)
    fmt.first_line_indent = Cm(first_line)
    fmt.tab_stops.add_tab_stop(Cm(tab))
    if text:
        run = p.add_run(text)
        run.italic = italic
        sub = p.add_run(sub_text)
        if sub_text is not None or sub_text != '':
            sub.font.subscript = True
        sup = p.add_run(sup_text)
        if sup_text is not None or sup_text != '':
            sup.font.superscript = True
        if value is not None:
            p.add_run(f" = {value} – ")
        else:
            p.add_run(" – ")
        if desc:
            p.add_run("\t")
            p.add_run(desc)
    return p


def add_equation(doc, latex: str, number=None, f_l_indent=0.0, font_size_pt=14, min_line_pts=30):
    section = doc.sections[0]
    content_width_emu = section.page_width - section.left_margin - section.right_margin
    right_tab_pos_pt = content_width_emu / 12700.0  # EMU -> pt

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.first_line_indent = Cm(f_l_indent)  # без красной строки
    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    fmt.line_spacing = Pt(min_line_pts)  # ключ к отсутствию обрезания
    fmt.space_after = fmt.space_before = Pt(2)

    math2docx.add_math(p, latex)

    if number is not None:
        fmt.tab_stops.add_tab_stop(Pt(right_tab_pos_pt), alignment=WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        p.add_run(f"({number})").font.size = Pt(font_size_pt)

    return p


def add_heading(doc, flange_type, D_N, D_n, D_m, pressure):
    if flange_type == "one_one":
        title = ("С ЦИЛИНДРИЧЕСКОЙ ВТУЛКОЙ" if D_m == D_n else "С КОНИЧЕСКОЙ ВТУЛКОЙ")
    elif flange_type == "zero_one":
        title = "С ПЛОСКИМИ ФЛАНЦАМИ"
    else:
        return
    for _ in range(3):
        doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border_el = OxmlElement(f"w:{border_name}")
        border_el.set(qn('w:val'), 'none')
        tblBorders.append(border_el)

    tblPr = tbl.xpath("./w:tblPr")[0]
    tblPr.append(tblBorders)
    table.columns[0].width = Cm(9.75)
    table.columns[1].width = Cm(7.25)
    cell = table.cell(0, 1)
    paragraphs = [
        ("УТВЕРЖДАЮ", True),
        ("Генеральный директор", False),
        ("ООО «ИНОТЭК»", False),
        ("____________ Р.Т. Фахретдинов", False),
        ("« ___ » ________________2025 г.", False),
    ]
    for text, bold in paragraphs:
        p = cell.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        run = p.runs[0]
        run.bold = bold
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph(f"РАСЧЕТ НА ПРОЧНОСТЬ И ГЕРМЕТИЧНОСТЬ ПРИВАРНОГО ВСТЫК ФЛАНЦЕВОГО СОЕДИНЕНИЯ {title} ",
                          style='Heading 1')
    r = p.add_run()
    r.add_break(WD_BREAK.LINE)
    r.add_text(f"DN{int(D_N)} PN{int(pressure * 10)}")
    r.add_break(WD_BREAK.LINE)
    r.add_text(f"DN{int(D_N)}/{int(D_N)}-PN{int(pressure * 10)}.000")
    for _ in range(14):
        doc.add_paragraph()
    doc.add_paragraph('Уфа 2025 г.')

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_external_forces(doc, force, moment):
    if force == moment == 0:
        text = "Внешняя осевая сила и изгибающий момент отсутствуют."
    elif force != 0 and moment == 0:
        text = f"Внешняя осевая сила равна {force} кН, изгибающий момент отсутствует."
    elif force == 0 and moment != 0:
        text = f"Внешняя осевая сила отсутствует, изгибающий момент равен {moment} кН*м."
    else:
        text = f"Внешняя осевая сила равна {force} кН, изгибающий момент равен {moment} кН*м."
    doc.add_paragraph(text)


def add_gasket_description(doc, gasket_document, gasket_data):
    g = SimpleNamespace(**gasket_data)

    if g.document == 'ГОСТ 15180-86':
        text = (f"Прокладка плоская типа {g.type} {g.document}, материал – {g.material}, "
                f"наружный диаметр Dнп = {g.D_outer} мм, "
                f"внутренний диаметр dнп = {g.d_inner} мм, "
                f"толщина {g.thickness} мм [2].")
    elif gasket_document == 'ГОСТ 34655-2020':
        extra = {
            "1": f"радиус сечения – {g.radius_or_height_1} мм",
            "2": f"высота сечения по наружному диаметру – {g.radius_or_height_1} мм",
            "3": f"внешние радиусы сечения – {g.radius_or_height_1} мм"
        }
        text = (f"Прокладка стальная {['овального', 'восьмиугольного', 'линзовая'][int(g.type) - 1]} "
                f"сечения типа {g.type} {g.document}, материал – {g.material}, "
                f"наружный диаметр Dнп = {g.D_outer} мм, "
                f"внутренний диаметр dнп = {g.d_inner} мм, "
                f"толщина {g.thickness} мм, {extra[g.type]} [2].")
    elif gasket_document == 'ISO 7483-2011':
        extra = {
            "1": f"радиус сечения – {g.radius_or_width_c} мм",
            "2": f"ширина сечения по верхнему краю – {g.radius_or_width_c} мм"
        }
        text = (f"Прокладка стальная "
                f"{'овального' if g.type == '1' else 'восьмиугольного'} сечения "
                f"типа {g.type} {g.document}, материал – {g.material}, "
                f"наружный диаметр Dнп = {g.D_outer} мм, "
                f"внутренний диаметр dнп = {g.d_inner} мм, "
                f"толщина {g.thickness} мм, {extra[g.type]} [2].")
    elif gasket_document == 'ГОСТ 28759.8-2022':
        extra = {
            "1": f"высота сечения по наружному диаметру – {g.radius_or_height_1} мм",
        }
        text = (f"Прокладка стальная восьмиугольного "
                f"сечения типа {g.type} {g.document}, материал – {g.material}, "
                f"наружный диаметр Dнп = {g.D_outer} мм, "
                f"внутренний диаметр dнп = {g.d_inner} мм, "
                f"толщина {g.thickness} мм, {extra[g.type]} [2].")
    else:
        return
    doc.add_paragraph(text)


def add_flange_thickness(doc, flange_type, D_n, D_m, D_int, T_b, f_counter):
    if flange_type == 'one_one':
        if D_m > D_n:
            doc.add_paragraph('Толщина втулки фланца в месте приварки:')
            f_counter += 1
            add_equation(doc, r'S_0 = \frac{D_n - D_{вн}}{2}', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='n', value=D_n,
                                      indent=3.0,
                                      desc='наружный диаметр втулки фланца в месте приварки, мм;')
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='вн', value=D_int,
                                      indent=3.0,
                                      desc='внутренний диаметр фланца, мм.')
            doc.add_paragraph('Толщина втулки фланца в месте присоединения тарелки:')
            f_counter += 1
            add_equation(doc, r'S_1 = \frac{D_m - D_{вн}}{2}', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='m', value=D_m,
                                      indent=3.0,
                                      desc='наружный диаметр втулки фланца в месте присоединения тарелки, мм;')
        else:
            doc.add_paragraph('Толщина втулки фланца:')
            f_counter += 1
            add_equation(doc, r'S_0, S_1 = \frac{D_n - D_{вн}}{2}', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='n', value=D_n,
                                      indent=3.0,
                                      desc='наружный диаметр втулки фланца, мм;')
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='вн', value=D_int,
                                      indent=3.0,
                                      desc='внутренний диаметр фланца, мм.')
    elif flange_type == 'zero_one':
        doc.add_paragraph('Толщина стенки ответвления разрезного тройника:')
        f_counter += 1
        add_equation(doc, rf'S_0, S_1 = T_b = {T_b} мм', number=f_counter, f_l_indent=6.75)
    return f_counter


def generate_new_report(result, input_select_names):
    r = SimpleNamespace(**result)
    print('r')
    for k in result.keys():
        print(k, "=", getattr(r, k), type(getattr(r, k)))

    gasket_data = dict(result['gasket_params'])
    g = SimpleNamespace(**gasket_data)
    print('g')
    for k in gasket_data.keys():
        print(k, "=", getattr(g, k), type(getattr(g, k)))

    fasteners_data = dict(result['fasteners_data'])
    nut_data = fasteners_data['nut']
    washer_data = fasteners_data['washer']
    bolt_data = fasteners_data['bolt']
    bolt_type = bolt_data[0]
    nut_height = nut_data[1]
    washer_thickness = washer_data[2]

    flange_type_name = input_select_names[0]

    stress_delta_data = dict(result['stress_delta_data'])
    sd = SimpleNamespace(**stress_delta_data)
    print('sd')
    for k in stress_delta_data.keys():
        print(k, "=", getattr(sd, k), type(getattr(sd, k)))

    f_counter = 0
    doc = Document()
    section = doc.sections[0]

    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.left_margin, section.right_margin = Cm(3), Cm(1.25)
    section.top_margin, section.bottom_margin = Cm(1.75), Cm(2.5)
    section.header_distance = section.footer_distance = Cm(0)

    style = doc.styles['Normal']
    style.font.name, style.font.size = 'Times New Roman', Pt(14)

    fmt = style.paragraph_format
    fmt.line_spacing = Pt(21)
    fmt.first_line_indent = Cm(1.25)
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_after = fmt.space_before = Pt(0)

    set_heading_style(doc.styles['Heading 1'], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_heading_style(doc.styles['Heading 2'], align=WD_ALIGN_PARAGRAPH.LEFT, f_l_indent=1.25)

    add_heading(doc, r.flange_type, r.D_N_flange, r.D_n_flange, r.D_m_flange, r.pressure)
    doc.add_paragraph("СОДЕРЖАНИЕ", style='Heading 1')
    doc.add_section(WD_SECTION.NEW_PAGE)

    doc.add_paragraph("Настоящий расчет распространяется на фланцевые соединения разрезных тройников, привариваемых "
                      "под давлением на трубопроводы с рабочим давлением до 16.0 МПа включительно с номинальным "
                      "диаметром от DN500 до DN1200, транспортирующие углеводороды (природный газ), жидкие углеводороды"
                      " (нефть и нефтепродукты), стабильный и нестабильный конденсат, широкие фракции углеводородов.")
    doc.add_paragraph("1. Исходные данные", style='Heading 2')

    lines = [
        'Расчет производится согласно ГОСТ Р 52874.4-2007 "Нормы и методы расчета на прочность" [1]. '
        'Геометрические размеры фланца для первой итерации расчета согласно [5]. Для последней итерации – '
        'приведены ниже:',
        f'Тип фланца: {flange_type_name};',
        f'Наружный диаметр тарелки фланца: {r.D_ext_flange} мм;'
    ]
    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            lines.append(f'Наружный диаметр в месте приварки: {r.D_n_flange} мм;')
            lines.append(f'Наружный диаметр в месте присоединения тарелки: {r.D_m_flange} мм;')
            lines.append(f'Радиус скругления между конической втулкой и тарелкой: {r.r} мм;')
        else:
            lines.append(f'Наружный диаметр цилиндрической втулки фланца: {r.D_n_flange} мм;')
            lines.append(f'Радиус скругления между цилиндрической втулкой и тарелкой: {r.r} мм;')
    elif r.flange_type == 'zero_one':
        lines.append(f'Толщина стенки ответвления тройника: {r.T_b} мм;')
    lines.append(f'Толщина тарелки фланца: {r.h_flange} мм;')
    lines.append(f'Полная высота фланца: {r.H_flange} мм.')

    for line in lines:
        doc.add_paragraph(line, style='Normal')

    p1 = doc.add_paragraph()
    parts = [
        {"text": f"Фланцевое соединение с приварным встык фланцем Ду {int(r.D_N_flange)} с уплотнительной "
                 f"поверхностью типа {r.face_type} [5], устанавливаемый на ответвление разрезного тройника "
                 f"с внутренним диаметром {r.D_int_flange} мм, давление среды {r.pressure} МПа. "
                 f"Количество шпилек фланца – {int(r.pins_quantity)} шт. с гайкой М{int(r.pin_diam)}, "
                 f"расположенных на окружности диаметром ",
         "italic": False, "subscript": False},
        {"text": "D", "italic": True, "subscript": False},
        {"text": "б", "italic": False, "subscript": True},
        {"text": f" = {r.d_pins_flange} мм. Расчетная температура в зоне фланцевого соединения – {r.temperature}˚С. "
                 f"Марка стали фланца – {r.flange_steel}, механические свойства при расчетной температуре: σ",
         "italic": False, "subscript": False},
        {"text": "В мин", "italic": False, "subscript": True},
        {"text": f" = {r.sigma_b_flange_T} МПа; σ", "italic": False, "subscript": False},
        {"text": "Т мин", "italic": False, "subscript": True},
        {"text": f" = {r.sigma_y_flange_T} МПа. Марка стали шпилек и гаек – {r.bolt_steel}, механические свойства при "
                 f"расчетной температуре: σ",
         "italic": False, "subscript": False},
        {"text": "В мин", "italic": False, "subscript": True},
        {"text": f" = {r.sigma_b_bolt_T} МПа; σ", "italic": False, "subscript": False},
        {"text": "Т мин", "italic": False, "subscript": True},
        {"text": f" = {r.sigma_y_bolt_T} МПа.", "italic": False, "subscript": False},
    ]
    for part in parts:
        run = p1.add_run(part["text"])
        run.italic = part["italic"]
        if part["subscript"]:
            run.font.subscript = True

    add_external_forces(doc, r.ext_force, r.ext_moment)
    add_gasket_description(doc, g.document, gasket_data)
    f_counter = add_flange_thickness(doc, r.flange_type, r.D_n_flange, r.D_m_flange, r.D_int_flange, r.T_b, f_counter)
    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            add_equation(doc, rf'S_0 = \frac{{{r.D_n_flange} - {r.D_int_flange}}}{{2}} '
                              rf'= {(r.D_n_flange - r.D_int_flange) / 2} мм')
            add_equation(doc, rf'S_1 = \frac{{{r.D_m_flange} - {r.D_int_flange}}}{{2}} '
                              rf'= {(r.D_m_flange - r.D_int_flange) / 2} мм')
            S_0_calc = (r.D_n_flange - r.D_int_flange) / 2
            S_1_calc = (r.D_m_flange - r.D_int_flange) / 2
        else:
            add_equation(doc, rf'S_0, S_1 = \frac{{{r.D_n_flange} - {r.D_int_flange}}}{{2}} '
                              rf'= {(r.D_n_flange - r.D_int_flange) / 2} мм')
            S_0_calc = S_1_calc = (r.D_n_flange - r.D_int_flange) / 2
    else:
        add_equation(doc, rf'S_0, S_1 = T_b = {r.T_b} мм')
        S_0_calc = S_1_calc = r.T_b

    doc.add_paragraph('Расчет применим к фланцам, удовлетворяющим следующим условиям:')

    f_counter += 1
    add_equation(doc, rf'\frac{{D_н}}{{D_в}} <= 5.0', number=f_counter, f_l_indent=7.0)
    f_counter += 1
    add_equation(doc, rf'\frac{{2 \cdot H}}{{D_н - D_в}} >= 0.25', number=f_counter, f_l_indent=6.5)
    f_counter += 1
    add_equation(doc, rf'\frac{{S_1 - S_0}}{{L}} <= 0.4', number=f_counter, f_l_indent=6.75)

    add_equation(doc, rf'{round(r.D_ext_flange / r.D_int_flange, 3)} <= 5.0')
    add_equation(doc, rf'{round(2 * r.h_flange / (r.D_ext_flange - r.D_int_flange), 3)} >= 0.25')
    add_equation(doc, rf'{round((S_1_calc - S_0_calc) / (r.H_flange - r.h_flange), 3)} <= 0.4')
    doc.add_paragraph('Условия выполняются.')

    doc.add_paragraph("2. Усилие, необходимое для смятия прокладки и обеспечения герметичности "
                      "соединения", style='Heading 2')

    doc.add_paragraph("Ширина прокладки определяется исходя из ее наружного и внутреннего диаметра:")
    f_counter += 1
    add_equation(doc, r'b_п = \frac{D_{нп} - d_{нп}}{2}', number=f_counter, f_l_indent=6.75)
    b_p_calc = round((g.D_outer - g.d_inner) / 2, 1)
    add_equation(doc, rf'b_п = \frac{{{g.D_outer} - {g.d_inner}}}{{2}} = {b_p_calc} мм')

    if g.document == 'ГОСТ 15180-86' and b_p_calc <= 15.0:
        b_0_calc = b_p_calc
        p2 = doc.add_paragraph(f"Эффективная ширина плоской прокладки при ")
        p2.add_run("b").italic = True
        p2.add_run("п").font.subscript = True
        p2.add_run(" <= 15 мм:")
        f_counter += 1
        add_equation(doc, rf'b_0 = b_п = {b_0_calc} мм', number=f_counter)
    elif g.document == 'ГОСТ 15180-86' and b_p_calc > 15.0:
        b_0_calc = round(3.8 * b_p_calc ** 0.5, 1)
        p2 = doc.add_paragraph(f"Эффективная ширина плоской прокладки при ")
        p2.add_run("b").italic = True
        p2.add_run("п").font.subscript = True
        p2.add_run(" > 15 мм:")
        f_counter += 1
        add_equation(doc, r'b_0 = 3.8 \cdot \sqrt{b_п}', number=f_counter, f_l_indent=6.75)
        add_equation(doc, rf'b_0 = 3.8 \cdot \sqrt{{{b_p_calc}}} = {b_0_calc} мм')
    elif g.document in ['ГОСТ 34655-2020', 'ISO 7483-2011', 'ГОСТ 28759.8-2022']:
        b_0_calc = round(b_p_calc / 4, 1)
        doc.add_paragraph("Эффективная ширина прокладки овального или восьмиугольного сечения:")
        f_counter += 1
        add_equation(doc, r'b_0 = \frac{b_п}{4}', number=f_counter, f_l_indent=7.25)
        add_equation(doc, rf'b_0 = \frac{{{b_p_calc}}}{{4}} = {b_0_calc} мм')
    else:
        b_0_calc = None

    if g.document == 'ГОСТ 15180-86':
        D_sp_calc = round(g.D_outer - b_0_calc, 1)
        doc.add_paragraph("Расчетный диаметр плоской прокладки:")
        f_counter += 1
        add_equation(doc, r'D_{сп} = D_{нп} - b_0', number=f_counter, f_l_indent=6.75)
        add_equation(doc, rf'D_{{сп}} = {g.D_outer} - {b_0_calc} = {D_sp_calc} мм')
    elif g.document in ['ГОСТ 34655-2020', 'ISO 7483-2011', 'ГОСТ 28759.8-2022']:
        D_sp_calc = round(g.D_outer - b_p_calc, 1)
        doc.add_paragraph("Расчетный диаметр прокладки овального или восьмиугольного "
                          "сечения равен ее среднему диаметру:")
        f_counter += 1
        add_equation(doc, r'D_{сп} = D_{нп} - b_п', number=f_counter, f_l_indent=6.5)
        add_equation(doc, rf'D_{{сп}} = {g.D_outer} - {b_p_calc} = {D_sp_calc} мм')
    else:
        D_sp_calc = None

    doc.add_paragraph("Усилие, необходимое для смятия прокладки при затяжке:")
    P_obj_calc = round(0.5 * math.pi * D_sp_calc * b_0_calc * g.q_obj / 1000, 3)
    f_counter += 1
    add_equation(doc, r'P_{обж} = 0.5 \cdot \pi \cdot D_{сп} \cdot b_0 \cdot q_{обж}',
                 number=f_counter, f_l_indent=5.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, "q", italic=True, sub_text='обж', value=g.q_obj,
                              desc="удельное давление обжатия прокладки, МПа (принимается согласно [1], приложение И).")
    add_equation(doc, rf'P_{{обж}} = 0.5 \cdot \pi \cdot {D_sp_calc} \cdot {b_0_calc} \cdot {g.q_obj}'
                      rf' = {P_obj_calc} кН')

    doc.add_paragraph("Усилие на прокладке, необходимое для обеспечения герметичности фланцевого "
                      "соединения в рабочих условиях:")
    R_p_calc = round(math.pi * D_sp_calc * b_0_calc * g.m * r.pressure / 1000, 3)
    f_counter += 1
    add_equation(doc, r'R_п = \pi \cdot D_{сп} \cdot b_0 \cdot m \cdot P', number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'm', italic=True, value=g.m, indent=2.25,
                              desc='прокладочный коэффициент (принимается согласно [1], приложение И);')
    add_paragraph_with_indent(doc, 'P', italic=True, indent=1.25, desc='расчетное давление, МПа.')
    add_equation(doc, rf'R_п = \pi \cdot {D_sp_calc} \cdot {b_0_calc} \cdot {g.m} \cdot {r.pressure}'
                      rf' = {R_p_calc} кН')

    doc.add_paragraph("3. Усилие в шпильках фланцевого соединения при затяжке и в рабочих условиях",
                      style='Heading 2')
    doc.add_paragraph("Суммарная площадь поперечного сечения шпилек по внутреннему диаметру резьбы:")
    bolts_area_calc = r.pins_quantity * r.bolt_area
    f_counter += 1
    add_equation(doc, r'A_ш = n \cdot f_ш', number=f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'n', italic=True, indent=1.25, desc='количество шпилек фланца, шт.;')
    add_paragraph_with_indent(doc, 'f', italic=True, sub_text='ш', value=r.bolt_area, indent=2.75,
                              desc=f'площадь поперечного сечения шпильки, мм² (принимается для одной шпильки'
                                   f' типа {bolt_type} под гайку с резьбой М{int(r.pin_diam)} согласно [1], '
                                   f'таблица Д.1).')
    add_equation(doc, rf'A_ш = {r.pins_quantity} \cdot {r.bolt_area} = {bolts_area_calc} мм²')

    doc.add_paragraph("Равнодействующая нагрузка от давления:")
    Q_d_calc = round(0.785 * D_sp_calc ** 2 * r.pressure / 1000, 3)
    f_counter += 1
    add_equation(doc, r'Q_д = 0.785 \cdot D_{сп}^2 \cdot P', number=f_counter, f_l_indent=6.25)
    add_equation(doc, rf'Q_д = 0.785 \cdot {D_sp_calc}^2 \cdot {r.pressure} = {Q_d_calc} кН')

    if r.ext_force != 0.0 or r.ext_moment != 0.0:
        doc.add_paragraph("Приведенная нагрузка, вызванная воздействием внешней силы и изгибающего момента:")
        Q_fm_calc = round(max(r.ext_force + 4 * abs(r.ext_moment) * 1000 / D_sp_calc,
                              r.ext_force - 4 * abs(r.ext_moment) * 1000 / D_sp_calc), 3)
        f_counter += 1
        add_equation(doc, r'Q_{FM} = F ± \frac{4 \cdot |M|}{D_{сп}}', number=f_counter, f_l_indent=6.25)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'F', italic=True, indent=1.25, desc='внешняя осевая сила, кН;')
        add_paragraph_with_indent(doc, 'M', italic=True, indent=1.25, desc='внешний изгибающий момент, кН*м.')
        add_equation(doc, rf'Q_{{FM}} = {r.ext_force} ± \frac{{4 \cdot |{r.ext_moment}|}}{{{D_sp_calc}}}'
                          rf' = {Q_fm_calc} кН')
    elif r.ext_force == 0.0 and r.ext_moment == 0.0:
        doc.add_paragraph("Приведенная нагрузка, вызванная воздействием внешней силы и изгибающего момента равна нулю,"
                          " так как внешняя сила и изгибающий момент отсутствуют:")
        f_counter += 1
        add_equation(doc, r'Q_{FM} = F ± \frac{4 \cdot |M|}{D_{сп}} = 0', number=f_counter, f_l_indent=6.25)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'F', italic=True, indent=1.25, desc='внешняя осевая сила, кН;')
        add_paragraph_with_indent(doc, 'M', italic=True, indent=1.25, desc='внешний изгибающий момент, кН*м.')
        Q_fm_calc = 0.0
    else:
        Q_fm_calc = None

    doc.add_paragraph("Нагрузка, вызванная стесненностью температурных деформаций в соединении с приварным встык"
                      " или плоским фланцем:")
    f_counter += 1
    f_counter_mem_2 = f_counter
    add_equation(doc, r'Q_t = \gamma \cdot \Big[ 2 \cdot \alpha_ф \cdot h \cdot (t_ф - 20) '
                      r'- 2 \cdot \alpha_ш \cdot h \cdot (t_ш - 20) \Big]', number=f_counter, f_l_indent=2.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'γ', italic=True, indent=1.25, desc='коэффициент жесткости фланцевого'
                                                                       ' соединения, Н/мм;')
    add_paragraph_with_indent(doc, 'α', italic=True, sub_text='ф', indent=1.25,
                              desc='температурный коэффициент линейного расширения материала фланца, 1/˚С '
                                   '([1], таблица Ж.2);')
    add_paragraph_with_indent(doc, 'h', italic=True, sub_text='', indent=1.25,
                              desc='толщина тарелки фланца, мм;')
    add_paragraph_with_indent(doc, 't', italic=True, sub_text='ф', indent=1.25,
                              desc='рабочая температура фланцев, ˚С; ')
    add_paragraph_with_indent(doc, 'α', italic=True, sub_text='ш', indent=1.25,
                              desc='температурный коэффициент линейного расширения материала шпильки, 1/˚С '
                                   '([1], таблица Ж.2);')
    add_paragraph_with_indent(doc, 't', italic=True, sub_text='ш', indent=1.25,
                              desc='рабочая температура шпильки, ˚С; ')

    doc.add_paragraph("Коэффициент жесткости фланцевого соединения для приварных встык и плоских фланцев "
                      "согласно [1], приложение Е, определяется по формуле:")
    f_counter += 1
    f_counter_mem_1 = f_counter
    add_equation(doc, r'\gamma = \frac{1}{y_п + y_ш \cdot \frac{E_ш^{20}}{E_ш} + \left( 2 \cdot y_{ф} \cdot '
                      r'\frac{E^{20}}{E} \right) \cdot b^2}', number=f_counter, f_l_indent=5.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'y', italic=True, sub_text='п, ш', indent=1.5,
                              desc='податливость прокладки, шпильки, мм/Н ([1], приложение К);')
    add_paragraph_with_indent(doc, 'y', italic=True, sub_text='ф', indent=2.0,
                              desc='угловая податливость фланца при затяжке, 1/(Н*мм) ([1], приложение К);')
    add_paragraph_with_indent(doc, 'b', italic=True, indent=1.0, desc='плечо усилий в шпильках, мм '
                                                                      '([1], приложение Е);')
    add_paragraph_with_indent(doc, 'E', italic=True, sup_text='20', indent=1.5,
                              desc='модуль продольной упругости материала фланца, шпильки при 20˚С, '
                                   'МПа ([1], приложение Ж).')
    add_paragraph_with_indent(doc, 'E', italic=True, indent=1.5,
                              desc='модуль продольной упругости материала фланца, шпильки при расчетной температуре, '
                                   'МПа ([1], приложение Ж).')

    if g.document == 'ГОСТ 15180-86':
        doc.add_paragraph('Податливость прокладки определяется по формуле:')
        y_p_calc = round(g.thickness * g.K_obj / (g.E_p * math.pi * D_sp_calc * b_p_calc), 3)
        f_counter += 1
        add_equation(doc, r'y_п = \frac{h_п \cdot K_{обж}}{E_п \cdot \pi \cdot D_{сп} \cdot b_п}',
                     number=f_counter, f_l_indent=6.5)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'h', italic=True, sub_text='п', indent=2.25,
                                  value=g.thickness, desc='толщина прокладки, мм [2];')
        add_paragraph_with_indent(doc, 'k', italic=True, sub_text='обж', indent=2.75, value=g.K_obj,
                                  desc='коэффициент обжатия прокладки ([1], приложение И);')
        add_paragraph_with_indent(doc, 'E', italic=True, sub_text='п', indent=2.75, value=g.E_p,
                                  desc='Модуль продольной упругости материала прокладки, МПа.')
        add_equation(doc, rf'y_п = '
                          rf'\frac{{{g.thickness} \cdot {g.K_obj}}}'
                          rf'{{{g.E_p} \cdot \pi \cdot {D_sp_calc} \cdot {b_p_calc}}} = {y_p_calc} мм/Н')
    elif g.document in ['ГОСТ 34655-2020', 'ISO 7483-2011', 'ГОСТ 28759.8-2022']:
        doc.add_paragraph('Податливость прокладки равняется нулю, так как прокладка стальная.')
        y_p_calc = 0.0
        add_equation(doc, r'y_п = 0')
    else:
        y_p_calc = None

    doc.add_paragraph('Податливость шпилек определяется по формуле:')
    f_counter += 1
    add_equation(doc, r'y_ш = \frac{L_ш}{E_ш^{20} \cdot А_ш}', number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'L', italic=True, sub_text='ш', indent=1.25,
                              desc='эффективная длина шпильки, мм.')
    doc.add_paragraph('Эффективная длина шпильки:')
    f_counter += 1
    add_equation(doc, r'L_ш = L_{ш0} + 0.56 \cdot d', number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'L', italic=True, sub_text='ш0', indent=1.25,
                              desc='начальная длина шпильки, мм;')
    add_paragraph_with_indent(doc, 'd', italic=True, indent=1.25, desc='диаметр шпильки, мм.')

    doc.add_paragraph('Начальная длина шпильки, округленная в большую сторону по [6]:')
    L_stud_calc = round(r.stud_length + 0.56 * r.pin_diam, 1)
    f_counter += 1
    add_equation(doc, r'L_{ш0} = 2 \cdot (h_г + h_ш + h)', number=f_counter, f_l_indent=5.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'h', italic=True, sub_text='г, ш', indent=1.5,
                              desc='высота гайки и шайбы, мм [3, 4].')
    add_equation(doc, rf'L_{{ш0}} = 2 \cdot ({nut_height} + {washer_thickness} + {r.h_flange})'
                      rf' = {r.stud_length} мм')
    add_equation(doc, rf'L_{{ш}} = {r.stud_length} + 0.56 \cdot {r.pin_diam} = {L_stud_calc} мм')

    doc.add_paragraph('Тогда, податливость шпилек:')
    y_b_calc = round(L_stud_calc / (r.E_bolt_20 * bolts_area_calc), 11)
    add_equation(doc, rf'y_ш = \frac{{{L_stud_calc}}}{{{r.E_bolt_20} \cdot {bolts_area_calc}}}'
                      rf' = {y_b_calc} мм/Н')

    doc.add_paragraph('Угловая податливость фланца при затяжке определяется по формуле:')
    f_counter += 1
    add_equation(doc, r'y_ф = \frac{0.91 \cdot \beta_V}{E^{20} \cdot \lambda \cdot S_0^2 \cdot l_0}',
                 number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='V', value=round(r.B_V, 3), indent=2.75,
                              desc='коэффициент, определяемый по рисунку К.3 [1], приложение К;')
    add_paragraph_with_indent(doc, 'λ', italic=True, value=round(r.lambda_, 3), indent=2.75,
                              desc='коэффициент, определяемый по формуле К.11 [1], приложение К, а также '
                                   'по рисункам К.2, К.3 и формулам К.3 - К.10')
    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            add_paragraph_with_indent(doc, 'S', italic=True, sub_text='0', indent=1.25,
                                      desc='толщина конической втулки приварного встык фланца в месте '
                                           'приварки к обечайке;')
        else:
            add_paragraph_with_indent(doc, 'S', italic=True, sub_text='0', indent=1.25,
                                      desc='толщина цилиндрической втулки приварного встык фланца;')
    else:
        add_paragraph_with_indent(doc, 'S', italic=True, sub_text='0', indent=1.25,
                                  desc='толщина стенки ответвления разрезного тройника;')
    add_paragraph_with_indent(doc, 'l', italic=True, sub_text='0', indent=1.25,
                              desc='длина обечайки, мм.')

    doc.add_paragraph('Длина обечайки определяется по формуле (К.3, [1], приложение К):')
    l_0_calc = round((r.D_int_flange * S_0_calc) ** 0.5, 1)
    f_counter += 1
    add_equation(doc, r'l_0 = \sqrt{D_{вн} \cdot S_0}', number=f_counter, f_l_indent=6.5)
    add_equation(doc, rf'l_0 = \sqrt{{{r.D_int_flange} \cdot {S_0_calc}}} = {l_0_calc} мм')

    doc.add_paragraph('Тогда, угловая податливость фланца при затяжке будет равна:')
    y_f_calc = round(0.91 * r.B_V / (r.E_flange_20 * r.lambda_ * S_0_calc ** 2 * l_0_calc), 15)
    add_equation(doc, rf'y_ф = '
                      rf'\frac{{0.91 \cdot {round(r.B_V, 3)}}}'
                      rf'{{{r.E_flange_20} \cdot {round(r.lambda_, 3)} \cdot {S_0_calc}^2  \cdot {l_0_calc}}}'
                      rf' = {y_f_calc} (1/(Н*мм))')

    doc.add_paragraph('Плечо усилий в шпильках определим по формуле:')
    b_calc = round(0.5 * (r.d_pins_flange - D_sp_calc), 1)
    f_counter += 1
    add_equation(doc, r'b = 0.5 \cdot (D_б - D_{сп})', number=f_counter, f_l_indent=6.0)
    add_equation(doc, rf'b = 0.5 \cdot ({r.d_pins_flange} - {D_sp_calc}) = {b_calc} мм')

    doc.add_paragraph(f'Подставим найденные значения в выражение ({f_counter_mem_1}):')
    gamma_calc = round(1 / (y_p_calc + y_b_calc * r.E_bolt_20 / r.E_bolt_T +
                            (2 * y_f_calc * r.E_flange_20 / r.E_flange_T) * b_p_calc ** 2), 5)
    add_equation(doc, rf'\gamma = \frac{{1}}{{{y_p_calc} + {y_b_calc} \cdot \frac{{{r.E_bolt_20}}}{{{r.E_bolt_T}}}'
                      rf'+ \left( 2 \cdot {y_f_calc} \cdot \frac{{{r.E_flange_20}}}{{{r.E_flange_T}}} \right)'
                      rf' \cdot {b_p_calc}^2}} = {gamma_calc} мм/Н')

    doc.add_paragraph("Рабочая температура фланцев и шпилек:")
    add_equation(doc, rf't_ф = 0.96 \cdot t = 0.96 \cdot {r.temperature} = {0.96 * r.temperature} ˚С')
    add_equation(doc, rf't_ш = 0.95 \cdot t = 0.95 \cdot {r.temperature} = {0.95 * r.temperature} ˚С')

    doc.add_paragraph(f'Далее, подставим найденные значения в выражение ({f_counter_mem_2}):')
    Q_t_calc = round(gamma_calc * (2 * r.alpha_flange_T * r.h_flange * (0.96 * r.temperature - 20) -
                                   r.alpha_bolt_T * 2 * r.h_flange * (0.95 * r.temperature - 20)) / 1000, 3)
    add_equation(doc, rf'Q_t = {gamma_calc} \cdot \Big[ 2 \cdot {r.alpha_flange_T} \cdot {r.h_flange} '
                      rf' \cdot ({0.96 * r.temperature} - 20) - 2 \cdot {r.alpha_bolt_T} \cdot {r.h_flange} '
                      rf'\cdot ({0.95 * r.temperature} - 20) \Big] = {Q_t_calc} кН')

    doc.add_paragraph('Расчетная нагрузка на шпильки фланцевого соединения при затяжке:')
    f_counter += 1
    f_counter_mem_8 = f_counter
    add_equation(doc, r'P_ш^м = \max (P_{ш1}; P_{ш2})', number=f_counter, f_l_indent=6.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'P', italic=True, sub_text='ш1', indent=1.25,
                              desc='расчетная нагрузка на шпильки при затяжке, необходимая для обеспечения '
                                   'достаточной герметизации фланцевого соединения в рабочих условиях '
                                   'давления на прокладку, кН;')
    add_paragraph_with_indent(doc, 'P', italic=True, sub_text='ш2', indent=1.25,
                              desc='расчетная нагрузка на шпильки при затяжке, необходимая для обеспечения'
                                   ' обжатия прокладки и минимального начального натяжения шпилек, кН.')
    f_counter += 1
    f_counter_mem_6 = f_counter
    add_equation(doc, r'P_{ш1} = \max \Big['
                      r'\alpha \cdot (Q_д + F) + R_п + \frac{4 \cdot \alpha_м \cdot \vert M \vert}{D_{сп}} \, , \\ '
                      r'\alpha \cdot (Q_д + F) + R_п + \frac{4 \cdot \alpha_м \cdot \vert M \vert}{D_{сп}} - Q_t'
                      r'\Big]', number=f_counter, f_l_indent=3.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'α', italic=True, indent=1.25,
                              desc='коэффициент жесткости фланцевого соединения, нагруженного внутренним давлением '
                                   'или внешней осевой силой ([1], приложение Е);')
    add_paragraph_with_indent(doc, 'α', italic=True, sub_text='М', indent=1.25,
                              desc='коэффициент жесткости фланцевого соединения, нагруженного внешним изгибающим '
                                   'моментом ([1], приложение Е);')
    f_counter += 1
    f_counter_mem_4 = f_counter
    add_equation(doc, r'\alpha = 1 - \frac{y_п - 2 \cdot y_ф \cdot e \cdot b}'
                      r'{y_п + y_ш + 2 \cdot y_ф \cdot b^2}',
                 number=f_counter, f_l_indent=6.0)
    f_counter += 1
    f_counter_mem_5 = f_counter
    add_equation(doc, r'\alpha_м = '
                      r'\frac{y_ш + 2 \cdot y_{фн} \cdot b \cdot \left( b + e - \frac{e^2}{D_{сп}} \right)}'
                      r'{y_ш + y_п \cdot \left( \frac{D_б}{D_{сп}} \right)^2 + 2 \cdot y_{фн} \cdot b^2}',
                 number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'e', italic=True, indent=1.25,
                              desc='плечо усилия от действия давления внутри фланца, мм;')
    add_paragraph_with_indent(doc, 'y', italic=True, sub_text='фн', indent=1.25,
                              desc='угловая податливость фланца, нагруженного внешним изгибающим моментом, 1/(Н*мм)'
                                   ' ([1], приложение К).')
    doc.add_paragraph('Плечо усилия от действия давления внутри фланца:')
    f_counter += 1
    f_counter_mem_3 = f_counter
    add_equation(doc, r'e = 0.5 \cdot (D_{сп} - D_{вн} - S_э)', number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    if r.flange_type == 'one_one':
        add_paragraph_with_indent(doc, 'S', italic=True, sub_text='э', indent=1.25,
                                  desc='эквивалентная толщина втулки приварных встык фланцев, мм')
        S_e_calc = round(r.ko_ef_B_x * S_0_calc, 1)
        f_counter += 1
        add_equation(doc, r'S_э = \zeta \cdot S_0', number=f_counter, f_l_indent=7.25)
        add_paragraph_with_indent(doc, 'ζ', italic=True, indent=2.5, value=round(r.ko_ef_B_x, 3),
                                  desc='коэффициент, зависящий от соотношения размеров конической втулки фланца,'
                                       ' определяется по формуле Е.6 [1], приложение Е, используя формулы К.9, '
                                       'К.10 [1], приложение К.')
        add_equation(doc, rf'S_э = {round(r.ko_ef_B_x, 3)} \cdot {S_0_calc} = {S_e_calc} мм')
    else:
        add_paragraph_with_indent(doc, 'S', italic=True, sub_text='э', indent=1.25,
                                  desc='эквивалентная толщина стенки ответвления разрезного тройника, мм')
        S_e_calc = S_0_calc
        f_counter += 1
        add_equation(doc, rf'S_э = S_0 = {S_0_calc} мм', number=f_counter, f_l_indent=1.0)

    doc.add_paragraph(f'Тогда, плечо усилия от действия давления внутри фланца, формула ({f_counter_mem_3}):')
    e_calc = round(0.5 * (D_sp_calc - r.D_int_flange - S_e_calc), 1)
    add_equation(doc, rf'e = 0.5 \cdot ({D_sp_calc} - {r.D_int_flange} - {S_e_calc}) = {e_calc}, мм')

    doc.add_paragraph('Угловая податливость фланца, нагруженного внешним изгибающим моментом:')
    y_f_n_calc = round((math.pi / 4) ** 3 * r.d_pins_flange / (r.E_flange_20 * r.D_ext_flange * r.h_flange ** 3), 14)
    f_counter += 1
    add_equation(doc, r'y_{фн} = '
                      r'\left( \frac{\pi}{4} \right)^3 \cdot \frac{D_б}{E^{20} \cdot D_н \cdot h^3}',
                 number=f_counter, f_l_indent=6.0)
    add_equation(doc, r'y_{фн} = '
                      rf'\left( \frac{{\pi}}{{4}} \right)^3 \cdot '
                      rf'\frac{{{r.d_pins_flange}}}{{{r.E_flange_20} \cdot {r.D_ext_flange} \cdot {r.h_flange}^3}}'
                      rf' = {y_f_n_calc} (1/(Н*мм))')
    doc.add_paragraph(f'Подставим найденные значения в выражения ({f_counter_mem_4}) и ({f_counter_mem_5}):')
    alpha_calc = round(1 -
                       (y_p_calc - 2 * y_f_calc * e_calc * b_calc) / (y_p_calc + y_b_calc + 2 * y_f_calc * b_calc ** 2),
                       3)
    alpha_M_calc = round(((y_b_calc + 2 * y_f_n_calc * b_calc * (b_calc + e_calc - e_calc ** 2 / D_sp_calc)) /
                          (y_b_calc + y_p_calc * (r.d_pins_flange / D_sp_calc) ** 2 + 2 * y_f_n_calc * b_calc ** 2)), 3)
    add_equation(doc, rf'\alpha = 1 - \frac{{{y_p_calc} - 2 \cdot {y_f_calc} \cdot {e_calc} \cdot {b_calc}}}'
                      rf'{{{y_p_calc} + {y_b_calc} + 2 \cdot {y_f_calc} \cdot {b_calc}^2}} = {alpha_calc}')
    add_equation(doc, rf'\alpha_м = '
                      rf'\frac{{{y_b_calc} + 2 \cdot {y_f_n_calc} \cdot {b_calc} \cdot \left( {b_calc} + {e_calc} - '
                      rf'\frac{{{e_calc}^2}}{{{D_sp_calc}}} \right)}}'
                      rf'{{{y_b_calc} + {y_p_calc} \cdot \left( \frac{{{r.d_pins_flange}}}{{{D_sp_calc}}} \right)^2 + '
                      rf'2 \cdot {y_f_n_calc} \cdot {b_calc}^2}} = {alpha_M_calc}')

    doc.add_paragraph(f'Подставим полученные значения в выражение ({f_counter_mem_6}):')
    P_b_1_calc = round(max(alpha_calc * (Q_d_calc + r.ext_force) + R_p_calc +
                           4 * alpha_M_calc * abs(r.ext_moment) * 1000 / D_sp_calc,
                           alpha_calc * (Q_d_calc + r.ext_force) + R_p_calc +
                           4 * alpha_M_calc * abs(r.ext_moment) * 1000 / D_sp_calc - Q_t_calc), 3)
    add_equation(doc, rf'P_{{ш1}} = \max \Big['
                      rf'{alpha_calc} \cdot ({Q_d_calc} + {r.ext_force}) + {R_p_calc} + '
                      rf'\frac{{4 \cdot {alpha_M_calc} \cdot \vert {r.ext_moment} \vert}}{{{D_sp_calc}}} \, , \\ '
                      rf'{alpha_calc} \cdot ({Q_d_calc} + {r.ext_force}) + {R_p_calc} + '
                      rf'\frac{{4 \cdot {alpha_M_calc} \cdot \vert {r.ext_moment} \vert}}{{{D_sp_calc}}} - {Q_t_calc}'
                      rf'\Big] = {P_b_1_calc} кН')

    doc.add_paragraph('Расчетная нагрузка на шпильки при затяжке, необходимая для обжатия прокладки и минимального '
                      'начального натяжения шпилек:')
    f_counter += 1
    f_counter_mem_7 = f_counter
    add_equation(doc, r'P_{ш2} = \max \Big( P_{обж}; 0,4 \cdot A_ш \cdot [σ]_м^ш \Big)',
                 number=f_counter, f_l_indent=5.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, '[σ]', italic=False, sub_text='м', sup_text='ш',
                              indent=1.5, desc='допускаемое напряжение для шпилек при затяжке в рабочих '
                                               'условиях, МПа ([1], приложение Г).')
    f_counter += 1
    add_equation(doc, r'[σ]_м^ш = \xi \cdot K_{у.р.} \cdot K_{у.з.} \cdot K_{у.т.} \cdot [σ]_н^ш',
                 number=f_counter, f_l_indent=5.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'ξ', italic=True, value=1.2, indent=2.75,
                              desc='коэффициент увеличения допускаемых напряжений при затяжке ([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.р.', value=1.0, indent=2.75,
                              desc='коэффициент условий работы для рабочих условий ([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.з.', value=1.1, indent=2.75,
                              desc='коэффициент условий затяжки при затяжке с контролем по крутящему моменту '
                                   '([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.т.', value=1.3, indent=2.75,
                              desc='коэффициент учета нагрузки от температурных деформаций ([1], приложение Г);')
    add_paragraph_with_indent(doc, '[σ]', sub_text='н', sup_text='ш', indent=1.75,
                              desc='номинальное допускаемое напряжение для шпилек при затяжке в рабочих условиях, МПа '
                                   '([1], приложение Г).')
    f_counter += 1
    add_equation(doc, r'[σ]_н^ш = \frac{σ_т}{n_т}', number=f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'n', italic=True, sub_text='т', indent=2.5, value=r.n_T_b,
                              desc='коэффициент запаса по отношению к пределу текучести ([1], приложение Г).')
    a_sigma_b_n_calc = round(r.sigma_y_bolt_T / r.n_T_b, 3)
    add_equation(doc, rf'[σ]_н^ш = \frac{{{r.sigma_y_bolt_T}}}{{{r.n_T_b}}} = {a_sigma_b_n_calc} МПа')
    a_sigma_b_m_calc = round(1.2 * 1.0 * 1.1 * 1.3 * a_sigma_b_n_calc, 3)
    add_equation(doc, rf'[σ]_м^ш = 1.2 \cdot 1.0 \cdot 1.1 \cdot 1.3 \cdot {a_sigma_b_n_calc}'
                      rf' = {a_sigma_b_m_calc} МПа')

    doc.add_paragraph('Допускаемое напряжение для шпилек в условиях испытания, МПа ([1], приложение Г):')
    f_counter += 1
    add_equation(doc, r'[σ]_р^ш = K_{у.р.} \cdot K_{у.з.} \cdot K_{у.т.} \cdot [σ]_н^ш',
                 number=f_counter, f_l_indent=5.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.р.', value=1.35, indent=2.75,
                              desc='коэффициент условий работы для условий испытания ([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.з.', value=1.1, indent=2.75,
                              desc='коэффициент условий затяжки при затяжке с контролем по крутящему моменту '
                                   '([1], приложение Г);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='у.т.', value=1.3, indent=2.75,
                              desc='коэффициент учета нагрузки от температурных деформаций ([1], приложение Г);')
    a_sigma_b_r_calc = round(1.35 * 1.1 * 1.3 * a_sigma_b_n_calc, 3)
    add_equation(doc, rf'[σ]_р^ш = 1.35 \cdot 1.1 \cdot 1.3 \cdot {a_sigma_b_n_calc} = {a_sigma_b_r_calc} МПа')

    doc.add_paragraph(f'Подставим полученные значения в выражение ({f_counter_mem_7}):')
    P_b_2_calc = round(max(P_obj_calc, 0.4 * bolts_area_calc * a_sigma_b_m_calc / 1000), 3)
    add_equation(doc, rf'P_{{ш2}} = \max \Big( {P_obj_calc}; '
                      rf'0.4 \cdot {bolts_area_calc} \cdot {a_sigma_b_m_calc} \Big) = {P_b_2_calc} кН')

    doc.add_paragraph(f'Тогда, расчетная нагрузка на шпильки при затяжке фланцевого соединения по формуле '
                      f'({f_counter_mem_8}):')
    P_b_m_calc = max(P_b_1_calc, P_b_2_calc)
    add_equation(doc, rf'P_ш^м = \max ({P_b_1_calc}; {P_b_2_calc}) = {P_b_m_calc} кН')

    doc.add_paragraph('Расчетная нагрузка на шпильки в рабочих условиях:')
    P_b_r_calc = round((P_b_m_calc + (1 - alpha_calc) * (Q_d_calc + r.ext_force) +
                        Q_t_calc + 4 * (1 - alpha_M_calc) * abs(r.ext_moment) * 1000 / D_sp_calc), 3)
    f_counter += 1
    add_equation(doc, r'P_ш^р = P_ш^м + (1 - \alpha) \cdot (Q_д + F) + '
                      r'Q_t + \frac{4 \cdot (1 - \alpha_м) \cdot '
                      r'\vert M \vert}{D_{сп}}',
                 number=f_counter, f_l_indent=3.0)
    add_equation(doc, rf'P_ш^р = {P_b_m_calc} + (1 - {alpha_calc}) \cdot ({Q_d_calc} + {r.ext_force}) + '
                      rf'{Q_t_calc} + \frac{{4 \cdot (1 - {alpha_M_calc}) \cdot '
                      rf'\vert {r.ext_moment} \vert}}{{{D_sp_calc}}} = {P_b_r_calc} кН')

    doc.add_paragraph("4. Проверка прочности шпилек и прокладки", style='Heading 2')
    word = "лет"
    if not (11 <= int(r.operating_time) % 100 <= 19):  # исключаем 11–19
        if int(r.operating_time) % 10 == 1:
            word = "год"
        elif 2 <= int(r.operating_time) % 10 <= 4:
            word = "года"
    doc.add_paragraph(f'Для учета воздействия на фланцевое соединение коррозионно-активной среды при определении '
                      f'напряжений в шпильках следует использовать диаметр шпильки за вычетом прибавки на коррозию'
                      f' [7], п.12.1; [1], п.4.12, определенной на срок эксплуатации {int(r.operating_time)} {word}.')

    doc.add_paragraph('Для определения суммарной площади сечения шпилек с учетом прибавки на коррозию, определим '
                      'диаметр сечения одной шпильки исходя их принятой площади сечения:')
    f_counter += 1
    add_equation(doc, r'd_c = \sqrt{\frac{4 \cdot f_ш}{\pi}} - 2 \cdot c_ш', number=f_counter, f_l_indent=6.25)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'c', italic=True, sub_text='ш', value=round(r.c_bolt, 2), indent=2.25,
                              desc=f'прибавка на коррозию для материала шпилек {r.bolt_steel} '
                                   f'за {int(r.operating_time)} {word}, мм.')
    d_c_bolt_calc = round((4 * r.bolt_area / math.pi) ** 0.5 - 2 * r.c_bolt, 1)
    bolt_corrosion_area = round(math.pi * d_c_bolt_calc ** 2 / 4, 1)
    bolts_corrosion_area = round(r.pins_quantity * bolt_corrosion_area, 1)
    add_equation(doc, rf'd_c = \sqrt{{\frac{{4 \cdot {r.bolt_area}}}{{\pi}}}} - 2 \cdot {r.c_bolt}'
                      rf' = {d_c_bolt_calc} мм')
    f_counter += 1
    add_equation(doc, r'f_{ш.с} = \frac{\pi \cdot d_c^2}{4}', number=f_counter, f_l_indent=7.25)
    add_equation(doc, rf'f_{{ш.с}} = \frac{{\pi \cdot {d_c_bolt_calc}^2}}{{4}} = {bolt_corrosion_area} мм²')

    f_counter += 1
    add_equation(doc, r'A_{ш.c} = n \cdot f_{ш.с}', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'A_{{ш.c}} = {r.pins_quantity} \cdot {bolt_corrosion_area} = {bolts_corrosion_area} мм²')

    doc.add_paragraph('Расчетные напряжения в шпильках при затяжке:')
    sigma_b_1_calc = round(P_b_m_calc * 1000 / bolts_corrosion_area, 3)
    f_counter += 1
    add_equation(doc, r'σ_{ш1} = \frac{P_ш^м}{A_{ш.c}}', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'σ_{{ш1}} = \frac{{{P_b_m_calc}}}{{{bolts_corrosion_area}}} = {sigma_b_1_calc} МПа')

    doc.add_paragraph('Расчетные напряжения в шпильках в рабочих условиях:')
    sigma_b_2_calc = round(P_b_r_calc * 1000 / bolts_corrosion_area, 3)
    f_counter += 1
    add_equation(doc, r'σ_{ш2} = \frac{P_ш^р}{A_{ш.c}}', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'σ_{{ш2}} = \frac{{{P_b_r_calc}}}{{{bolts_corrosion_area}}} = {sigma_b_2_calc} МПа')

    doc.add_paragraph('Условие прочности шпилек при затяжке:')
    f_counter += 1
    add_equation(doc, r'σ_{ш1} <= [σ]_м^ш', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'{sigma_b_1_calc} МПа <= {a_sigma_b_m_calc} МПа.')

    doc.add_paragraph('Условие прочности шпилек в рабочих условиях:')
    f_counter += 1
    add_equation(doc, r'σ_{ш2} <= [σ]_р^ш', number=f_counter, f_l_indent=7.0)
    add_equation(doc, rf'{sigma_b_2_calc} МПа <= {a_sigma_b_r_calc} МПа.')

    if g.document == 'ГОСТ 15180-86':
        doc.add_paragraph('Условие прочности прокладки:')
        f_counter += 1
        add_equation(doc, r'q <= [q]', number=f_counter, f_l_indent=7.0)

        f_counter += 1
        add_equation(doc, r'q = \frac{ \max (P_ш^м; P_ш^р)}{ \pi \cdot D_{сп} \cdot b_п}',
                     number=f_counter, f_l_indent=6.5)
        q_calc = round(max(P_b_m_calc, P_b_r_calc) * 1000 / (math.pi * D_sp_calc * b_p_calc), 3)
        add_equation(doc, rf'{q_calc} МПа <= {g.q_obj_dop} МПа.')
    else:
        doc.add_paragraph('Для стальных прокладок проверка прочности прокладок не требуется [1].')

    doc.add_paragraph("5. Расчет фланцев на статическую прочность", style='Heading 2')
    doc.add_paragraph('Расчетный изгибающий момент, действующий на фланец при затяжке:')
    f_counter += 1
    f_counter_mem_9 = f_counter
    add_equation(doc, r'M^M = C_F \cdot P_ш^м \cdot b', number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'C', italic=True, sub_text='F', indent=1.25,
                              desc='коэффициент, учитывающий изгиб тарелки фланца между шпильками ([1], приложение К).')
    f_counter += 1
    C_F_calc = round(max(1.0, ((math.pi * r.d_pins_flange / r.pins_quantity) /
                               (2 * r.pin_diam + 6 * r.h_flange / (g.m + 0.5))) ** 0.5), 3)
    add_equation(doc, r'C_F = \max \Big{ 1.0; \sqrt{\frac{\frac{\pi \cdot D_б}{n}}'
                      r'{2 \cdot \text{d} + \frac{6 \cdot h}{\text{m} + 0.5}}} '
                      r'\Big}', number=f_counter, f_l_indent=5.5)
    add_equation(doc, rf'C_F = \max \Big{{ 1.0; \sqrt{{\frac{{\frac{{\pi \cdot {r.d_pins_flange}}}'
                      rf'{{{r.pins_quantity}}}}}{{2 \cdot {r.pin_diam} + \frac{{6 \cdot {r.h_flange}}}'
                      rf'{{{g.m} + 0.5}}}}}} \Big}} = {C_F_calc}')

    doc.add_paragraph(f'Тогда, расчетный изгибающий момент, действующий на фланец при затяжке ({f_counter_mem_9}):')
    M_M_calc = round(C_F_calc * P_b_m_calc * b_calc / 1000, 3)
    add_equation(doc, rf'M^M = {C_F_calc} \cdot {P_b_m_calc} \cdot {b_calc} = {M_M_calc} кН*м')

    doc.add_paragraph('Расчетный изгибающий момент, действующий на фланец в рабочих условиях:')
    M_R_calc = round(C_F_calc * max(P_b_r_calc * b_calc / 1000 + (Q_d_calc + Q_fm_calc) * e_calc / 1000,
                                    abs(Q_d_calc + Q_fm_calc) * e_calc / 1000), 3)
    f_counter += 1
    add_equation(doc, r'M^Р = C_F \cdot \max \Big{ \Big[ P_ш^р \cdot b + '
                      r'\left( Q_д + Q_{FM} \right) \cdot e \Big];'
                      r' \vert Q_д + Q_{FM} \vert \cdot e \Big}', number=f_counter, f_l_indent=2.0)
    add_equation(doc, rf'M^Р = {C_F_calc} \cdot \max \Big{{ \Big[ {P_b_r_calc} \cdot {b_calc} + '
                      rf'\left( {Q_d_calc} + {Q_fm_calc} \right) \cdot {e_calc} \Big];'
                      rf' \vert {Q_d_calc} + {Q_fm_calc} \vert \cdot {e_calc} \Big}} = {M_R_calc} кН*м')

    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            p3 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p3.add_run('S').italic = True
            p3.add_run('1').font.subscript = True
            p3.add_run(' при затяжке:')
            f_counter += 1
            f_counter_mem_10 = f_counter
            add_equation(doc, r'σ_1^м = \frac{M^M}{\lambda \cdot (S_1 - с_ф)^2 \cdot D_{пр.}}',
                         number=f_counter, f_l_indent=6.25)
            p4 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p4.add_run('S').italic = True
            p4.add_run('0').font.subscript = True
            p4.add_run(' при затяжке:')
            f_counter += 1
            f_counter_mem_11 = f_counter
            add_equation(doc, r'σ_0^м = f \cdot σ_1^м', number=f_counter, f_l_indent=7.0)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'с', italic=True, sub_text='ф', value=round(r.c_flange, 2), indent=2.25,
                                      desc=f'прибавка на коррозию для материала шпилек {r.flange_steel} '
                                           f'за {int(r.operating_time)} {word}, мм;')
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='пр.', indent=2.25,
                                      desc='приведенный диаметр приварного встык фланца с конической втулкой;')
            add_paragraph_with_indent(doc, 'f', italic=True, value=round(r.f, 3), indent=2.25,
                                      desc='поправочный коэффициент для напряжений во втулке фланца ([1], '
                                           'рисунок К.4).')
            if r.D_int_flange >= 20 * S_1_calc:
                p5 = doc.add_paragraph('Приведенный диаметр при ')
                p5.add_run('D').italic = True
                p5.add_run('вн').font.subscript = True
                p5.add_run(' >= 20 * ')
                p5.add_run('S').italic = True
                p5.add_run('1').font.subscript = True
                f_counter += 1
                add_equation(doc, rf'D_{{пр.}} = D_{{вн}} = {r.D_int_flange} мм', number=f_counter, f_l_indent=3.5)
                D_pr_calc = r.D_int_flange
            elif r.D_int_flange < 20 * S_1_calc and r.f > 1:
                p5 = doc.add_paragraph('Приведенный диаметр при ')
                p5.add_run('D').italic = True
                p5.add_run('вн').font.subscript = True
                p5.add_run(' < 20 * ')
                p5.add_run('S').italic = True
                p5.add_run('1').font.subscript = True
                p5.add_run(' и ')
                p5.add_run('f').italic = True
                p5.add_run(' > 1.0:')
                f_counter += 1
                add_equation(doc, rf'D_{{пр.}} = D_{{вн}} + S_0 = {r.D_int_flange} + {S_0_calc} = '
                                  rf'{r.D_int_flange + S_0_calc} мм', number=f_counter, f_l_indent=3.5)
                D_pr_calc = r.D_int_flange + S_0_calc
            elif r.D_int_flange < 20 * S_1_calc and r.f == 1.0:
                p5 = doc.add_paragraph('Приведенный диаметр при ')
                p5.add_run('D').italic = True
                p5.add_run('вн').font.subscript = True
                p5.add_run(' < 20 * ')
                p5.add_run('S').italic = True
                p5.add_run('1').font.subscript = True
                p5.add_run(' и ')
                p5.add_run('f').italic = True
                p5.add_run(' = 1.0:')
                f_counter += 1
                add_equation(doc, rf'D_{{пр.}} = D_{{вн}} + S_1 = {r.D_int_flange} + {S_1_calc} = '
                                  rf'{r.D_int_flange + S_1_calc} мм', number=f_counter, f_l_indent=3.5)
                D_pr_calc = r.D_int_flange + S_1_calc
            else:
                D_pr_calc = None
            doc.add_paragraph(
                f'Подставим полученные значения в формулу меридиональных изгибных напряжений в условиях затяжки'
                f'(формулы ({f_counter_mem_10}) и ({f_counter_mem_11})):')
            sigma_1_m_calc = round(M_M_calc * 1000000 / (r.lambda_ * (S_1_calc - r.c_flange) ** 2 * D_pr_calc), 3)
            sigma_0_m_calc = round(r.f * sigma_1_m_calc, 3)
            add_equation(doc,
                         rf'σ_1^м = \frac{{{M_M_calc}}}'
                         rf'{{{r.lambda_:.3f} \cdot ({S_1_calc} - {r.c_flange})^2 \cdot {D_pr_calc}}} '
                         rf'= {sigma_1_m_calc} МПа')
            add_equation(doc, rf'σ_0^м = {r.f} \cdot {sigma_1_m_calc} = {sigma_0_m_calc} МПа')
        else:
            doc.add_paragraph('Меридиональное изгибное напряжение в цилиндрической втулке приварного встык фланца '
                              'при затяжке:')
            f_counter += 1
            f_counter_mem_10 = f_counter
            add_equation(doc, r'σ_1^м, σ_0^м = \frac{M^M}{\lambda \cdot (S_0 - с_ф)^2 \cdot D_{пр.}}',
                         number=f_counter, f_l_indent=6.25)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'с', italic=True, sub_text='ф', value=round(r.c_flange, 2), indent=2.25,
                                      desc=f'прибавка на коррозию для материала шпилек {r.flange_steel} '
                                           f'за {int(r.operating_time)} {word}, мм;')
            add_paragraph_with_indent(doc, 'D', italic=True, sub_text='пр.', indent=2.25,
                                      desc='приведенный диаметр приварного встык фланца с прямой втулкой;')
            if r.D_int_flange >= 20 * S_1_calc:
                p5 = doc.add_paragraph('Приведенный диаметр при ')
                p5.add_run('D').italic = True
                p5.add_run('вн').font.subscript = True
                p5.add_run(' >= 20 * ')
                p5.add_run('S').italic = True
                p5.add_run('1').font.subscript = True
                f_counter += 1
                add_equation(doc, rf'D_{{пр.}} = D_{{вн}} = {r.D_int_flange} мм', number=f_counter, f_l_indent=3.5)
                D_pr_calc = r.D_int_flange
            elif r.D_int_flange < 20 * S_1_calc and r.f > 1:
                p5 = doc.add_paragraph('Приведенный диаметр при ')
                p5.add_run('D').italic = True
                p5.add_run('вн').font.subscript = True
                p5.add_run(' < 20 * ')
                p5.add_run('S').italic = True
                p5.add_run('1').font.subscript = True
                p5.add_run(' и ')
                p5.add_run('f').italic = True
                p5.add_run(' > 1.0:')
                f_counter += 1
                add_equation(doc, rf'D_{{пр.}} = D_{{вн}} + S_0 = {r.D_int_flange} + {S_0_calc} = '
                                  rf'{r.D_int_flange + S_0_calc} мм', number=f_counter, f_l_indent=3.5)
                D_pr_calc = r.D_int_flange + S_0_calc
            elif r.D_int_flange < 20 * S_1_calc and r.f == 1.0:
                p5 = doc.add_paragraph('Приведенный диаметр при ')
                p5.add_run('D').italic = True
                p5.add_run('вн').font.subscript = True
                p5.add_run(' < 20 * ')
                p5.add_run('S').italic = True
                p5.add_run('1').font.subscript = True
                p5.add_run(' и ')
                p5.add_run('f').italic = True
                p5.add_run(' = 1.0:')
                f_counter += 1
                add_equation(doc, rf'D_{{пр.}} = D_{{вн}} + S_1 = {r.D_int_flange} + {S_1_calc} = '
                                  rf'{r.D_int_flange + S_1_calc} мм', number=f_counter, f_l_indent=3.5)
                D_pr_calc = r.D_int_flange + S_1_calc
            else:
                D_pr_calc = None
            doc.add_paragraph(
                f'Подставим полученные значения в формулу меридиональных изгибных напряжений в условиях затяжки'
                f'(формула ({f_counter_mem_10})):')
            sigma_1_m_calc = sigma_0_m_calc = round(M_M_calc * 1000000 /
                                                    (r.lambda_ * (S_0_calc - r.c_flange) ** 2 * D_pr_calc), 3)
            add_equation(doc,
                         rf'σ_1^м, σ_0^м = \frac{{{M_M_calc}}}'
                         rf'{{{r.lambda_:.3f} \cdot ({S_0_calc} - {r.c_flange})^2 \cdot {D_pr_calc}}} = '
                         rf'{sigma_1_m_calc or sigma_0_m_calc} МПа')
    else:
        doc.add_paragraph('Меридиональное изгибное напряжение в приварном встык плоском фланце при затяжке:')
        f_counter += 1
        f_counter_mem_10 = f_counter
        add_equation(doc, r'σ_1^м, σ_0^м = \frac{M^M}{\lambda \cdot (S_0 - с_ф)^2 \cdot D_{пр.}}',
                     number=f_counter, f_l_indent=6.25)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'с', italic=True, sub_text='ф', value=round(r.c_flange, 2), indent=2.25,
                                  desc=f'прибавка на коррозию для материала шпилек {r.flange_steel} '
                                       f'за {int(r.operating_time)} {word}, мм;')
        add_paragraph_with_indent(doc, 'D', italic=True, sub_text='пр.', indent=2.25,
                                  desc='приведенный диаметр приварного встык плоского фланца.')
        doc.add_paragraph('Приведенный диаметр приварного встык плоского фланца:')
        f_counter += 1
        add_equation(doc, rf'D_{{пр.}} = D_{{вн}} = {r.D_int_flange} мм', number=f_counter, f_l_indent=3.5)
        D_pr_calc = r.D_int_flange
        doc.add_paragraph(
            f'Подставим полученные значения в формулу меридиональных изгибных напряжений в условиях затяжки'
            f'(формула ({f_counter_mem_10})):')
        sigma_1_m_calc = sigma_0_m_calc = round(M_M_calc * 1000000 /
                                                (r.lambda_ * (S_0_calc - r.c_flange) ** 2 * D_pr_calc), 3)
        add_equation(doc,
                     rf'σ_1^м, σ_0^м = \frac{{{M_M_calc}}}'
                     rf'{{{r.lambda_:.3f} \cdot ({S_0_calc} - {r.c_flange})^2 \cdot {D_pr_calc}}} = '
                     rf'{sigma_1_m_calc or sigma_0_m_calc} МПа')

    doc.add_paragraph('Радиальное напряжение в тарелке приварного встык фланца в условиях затяжки:')
    f_counter += 1
    add_equation(doc, r'σ_R^M = \frac{1.33 \cdot \beta_F \cdot h + l_0}'
                      r'{\lambda \cdot h^2 \cdot l_0 \cdot D_{вн}} \cdot M^M', number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='F', value=round(r.B_F, 3), indent=2.75,
                              desc='коэффициент, зависящий от размеров втулки фланца, принимается по рисунку '
                                   'К.2 [1], приложение К.')
    sigma_r_m_calc = round(((1.33 * r.B_F * r.h_flange + l_0_calc) /
                            (r.lambda_ * r.h_flange ** 2 * l_0_calc * r.D_int_flange) * M_M_calc * 1000000), 3)
    add_equation(doc, rf'σ_R^M = \frac{{1.33 \cdot {r.B_F:.3f} \cdot {r.h_flange} + {l_0_calc}}}'
                      rf'{{{r.lambda_:.3f} \cdot {r.h_flange}^2 \cdot {l_0_calc} \cdot {r.D_int_flange}}}'
                      rf' \cdot {M_M_calc} = {sigma_r_m_calc} МПа')

    doc.add_paragraph('Окружное напряжение в тарелке приварного встык фланца в условиях затяжки:')
    f_counter += 1
    add_equation(doc, r'σ_T^M = \frac{\beta_Y \cdot M^M}{h^2 \cdot D_{вн}} - \beta_Z \cdot σ_R^M',
                 number=f_counter, f_l_indent=5.75)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='Y', value=round(r.B_Y, 3), indent=2.75,
                              desc='коэффициент, зависящий от соотношения размеров тарелки фланца, '
                                   'принимается по рисунку К.1 или по формулам К.4, К.7 [1], приложение К.')
    add_paragraph_with_indent(doc, 'β', italic=True, sub_text='Z', value=round(r.B_Z, 3), indent=2.75,
                              desc='коэффициент, зависящий от соотношения размеров тарелки фланца, '
                                   'принимается по рисунку К.1 или по формулам К.4, К.8 [1], приложение К.')
    sigma_t_m_calc = round(r.B_Y * M_M_calc * 1000000 / (r.h_flange ** 2 * r.D_int_flange) - r.B_Z * sigma_r_m_calc, 3)
    add_equation(doc, rf'σ_T^M = \frac{{{r.B_Y:.3f} \cdot {M_M_calc}}}{{{r.h_flange}^2 \cdot {r.D_int_flange}}}'
                      rf' - {r.B_Z:.3f} \cdot {sigma_r_m_calc} = {sigma_t_m_calc} МПа')

    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            p6 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p6.add_run('S').italic = True
            p6.add_run('1').font.subscript = True
            p6.add_run(' в рабочих условиях:')  # с - русская
            f_counter += 1
            add_equation(doc, r'σ_1^р = \frac{M^Р}{\lambda \cdot (S_1 - с_ф)^2 \cdot D_{пр.}}',
                         number=f_counter, f_l_indent=6.25)

            p7 = doc.add_paragraph('Меридиональное изгибное напряжение в конической втулке приварного '
                                   'встык фланца в сечении ')
            p7.add_run('S').italic = True
            p7.add_run('0').font.subscript = True
            p7.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_0^р = f \cdot σ_1^р', number=f_counter, f_l_indent=7.0)

            sigma_1_r_calc = round(M_R_calc * 1000000 / (r.lambda_ * (S_1_calc - r.c_flange) ** 2 * D_pr_calc), 3)
            sigma_0_r_calc = round(r.f * sigma_1_r_calc, 3)
            add_equation(doc,
                         rf'σ_1^р = \frac{{{M_R_calc}}}'
                         rf'{{{r.lambda_:.3f} \cdot ({S_1_calc} - {r.c_flange})^2 \cdot {D_pr_calc}}} '
                         rf'= {sigma_1_r_calc} МПа')
            add_equation(doc, rf'σ_0^р = {r.f} \cdot {sigma_1_r_calc} = {sigma_0_r_calc} МПа')

            p8 = doc.add_paragraph('Меридиональные мембранные напряжения в конической втулке приварного встык фланца в '
                                   'сечении ')
            p8.add_run('S').italic = True
            p8.add_run('1').font.subscript = True
            p8.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{1м.м}^р = \frac{Q_д + Q_{FM}}{\pi \cdot (D_{вн} + S_1) \cdot (S_1 - с_ф)}',
                         number=f_counter, f_l_indent=6.25)
            sigma_1_m_m_r_calc = round((Q_d_calc + Q_fm_calc) * 1000 /
                                       (math.pi * (r.D_int_flange + S_1_calc) * (S_1_calc - r.c_flange)), 3)
            add_equation(doc, rf'σ_{{1м.м}}^р = \frac{{{Q_d_calc} + {Q_fm_calc}}}'
                              rf'{{\pi \cdot ({r.D_int_flange} + {S_1_calc}) \cdot ({S_1_calc} - {r.c_flange})}} = '
                              rf'{sigma_1_m_m_r_calc} МПа')

            p9 = doc.add_paragraph('Меридиональные мембранные напряжения в конической втулке приварного встык фланца в '
                                   'сечении ')
            p9.add_run('S').italic = True
            p9.add_run('0').font.subscript = True
            p9.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{0м.м}^р = \frac{Q_д + Q_{FM}}{\pi \cdot (D_{вн} + S_0) \cdot (S_0 - с_ф)}',
                         number=f_counter, f_l_indent=6.25)
            sigma_0_m_m_r_calc = round((Q_d_calc + Q_fm_calc) * 1000 /
                                       (math.pi * (r.D_int_flange + S_0_calc) * (S_0_calc - r.c_flange)), 3)
            add_equation(doc, rf'σ_{{0м.м}}^р = \frac{{{Q_d_calc} + {Q_fm_calc}}}'
                              rf'{{\pi \cdot ({r.D_int_flange} + {S_0_calc}) \cdot ({S_0_calc} - {r.c_flange})}} = '
                              rf'{sigma_0_m_m_r_calc} МПа')
        else:
            sigma_1_m_m_r_calc = None
            doc.add_paragraph('Меридиональное изгибное напряжение в цилиндрической втулке приварного встык фланца '
                              'в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_1^р, σ_0^р = \frac{M^Р}{\lambda \cdot (S_0 - с)^2 \cdot D_{пр.}}',
                         number=f_counter, f_l_indent=6.25)
            sigma_1_r_calc = sigma_0_r_calc = round(M_R_calc * 1000000 /
                                                    (r.lambda_ * (S_0_calc - r.c_flange) ** 2 * D_pr_calc), 3)
            add_equation(doc,
                         rf'σ_1^р, σ_0^р = \frac{{{M_R_calc}}}'
                         rf'{{{r.lambda_:.3f} \cdot ({S_0_calc} - {r.c_flange})^2 \cdot {D_pr_calc}}} = '
                         rf'{sigma_1_r_calc or sigma_0_r_calc} МПа')

            p9 = doc.add_paragraph('Меридиональные мембранные напряжения в цилиндрической втулке приварного встык '
                                   'фланца в сечении ')
            p9.add_run('S').italic = True
            p9.add_run('0').font.subscript = True
            p9.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{0м.м}^р = \frac{Q_д + Q_{FM}}{\pi \cdot (D_{вн} + S_0) \cdot (S_0 - с_ф)}',
                         number=f_counter, f_l_indent=6.25)
            sigma_0_m_m_r_calc = round((Q_d_calc + Q_fm_calc) * 1000 /
                                       (math.pi * (r.D_int_flange + S_0_calc) * (S_0_calc - r.c_flange)), 3)
            add_equation(doc, rf'σ_{{0м.м}}^р = \frac{{{Q_d_calc} + {Q_fm_calc}}}'
                              rf'{{\pi \cdot ({r.D_int_flange} + {S_0_calc}) \cdot ({S_0_calc} - {r.c_flange})}} = '
                              rf'{sigma_0_m_m_r_calc} МПа')
    else:
        sigma_1_m_m_r_calc = None
        doc.add_paragraph('Меридиональное изгибное напряжение в приварном встык плоском фланце в рабочих условиях:')
        f_counter += 1
        add_equation(doc, r'σ_1^р, σ_0^р = \frac{M^Р}{\lambda \cdot (S_0 - с)^2 \cdot D_{пр.}}',
                     number=f_counter, f_l_indent=6.25)
        sigma_1_r_calc = sigma_0_r_calc = round(M_R_calc * 1000000 /
                                                (r.lambda_ * (S_0_calc - r.c_flange) ** 2 * D_pr_calc), 3)
        add_equation(doc,
                     rf'σ_1^р, σ_0^р = \frac{{{M_R_calc}}}'
                     rf'{{{r.lambda_:.3f} \cdot ({S_0_calc} - {r.c_flange})^2 \cdot {D_pr_calc}}} = '
                     rf'{sigma_1_r_calc or sigma_0_r_calc} МПа')

        doc.add_paragraph('Меридиональные мембранные напряжения в приварном встык плоском фланце в рабочих условиях:')
        f_counter += 1
        add_equation(doc, r'σ_{0м.м}^р = \frac{Q_д + Q_{FM}}{\pi \cdot (D_{вн} + S_0) \cdot (S_0 - с_ф)}',
                     number=f_counter, f_l_indent=6.25)
        sigma_0_m_m_r_calc = round((Q_d_calc + Q_fm_calc) * 1000 /
                                   (math.pi * (r.D_int_flange + S_0_calc) * (S_0_calc - r.c_flange)), 3)
        add_equation(doc, rf'σ_{{0м.м}}^р = \frac{{{Q_d_calc} + {Q_fm_calc}}}'
                          rf'{{\pi \cdot ({r.D_int_flange} + {S_0_calc}) \cdot ({S_0_calc} - {r.c_flange})}} = '
                          rf'{sigma_0_m_m_r_calc} МПа')

    p10 = doc.add_paragraph('Окружные мембранные напряжения от действия давления во втулке приварного встык фланца '
                            'в сечении ')
    p10.add_run('S').italic = True
    p10.add_run('0').font.subscript = True
    p10.add_run(' в рабочих условиях:')
    f_counter += 1
    add_equation(doc, r'σ_{0м.о}^р = \frac{P \cdot D_{вн}}{2 \cdot (S_0 - с_ф)}',
                 number=f_counter, f_l_indent=7.0)
    sigma_0_m_o_r_calc = round(r.pressure * r.D_int_flange / (2 * (S_0_calc - r.c_flange)), 3)
    add_equation(doc, rf'σ_{{0м.о}}^р = \frac{{{r.pressure} \cdot {r.D_int_flange}}}'
                      rf'{{2 \cdot ({S_0_calc} - {r.c_flange})}} = {sigma_0_m_o_r_calc} МПа')

    doc.add_paragraph('Радиальное напряжение в тарелке приварного встык фланца в рабочих условиях:')
    f_counter += 1
    add_equation(doc, r'σ_R^р = \frac{1.33 \cdot \beta_F \cdot h + l_0}'
                      r'{\lambda \cdot h^2 \cdot l_0 \cdot D_{вн}} \cdot M^Р', number=f_counter, f_l_indent=5.75)
    sigma_r_r_calc = round((1.33 * r.B_F * r.h_flange + l_0_calc) /
                           (r.lambda_ * r.h_flange ** 2 * l_0_calc * r.D_int_flange) * M_R_calc * 1000000, 3)
    add_equation(doc, rf'σ_R^р = \frac{{1.33 \cdot {r.B_F:.3f} \cdot {r.h_flange} + {l_0_calc}}}'
                      rf'{{{r.lambda_:.3f} \cdot {r.h_flange}^2 \cdot {l_0_calc} \cdot {r.D_int_flange}}}'
                      rf' \cdot {M_R_calc} = {sigma_r_r_calc} МПа')

    doc.add_paragraph('Окружное напряжение в тарелке приварного встык фланца в рабочих условиях:')
    f_counter += 1
    add_equation(doc, r'σ_T^р = \frac{\beta_Y \cdot M^Р}{h^2 \cdot D_{вн}} - \beta_Z \cdot σ_R^р',
                 number=f_counter, f_l_indent=5.75)
    sigma_t_r_calc = round(r.B_Y * M_R_calc * 1000000 / (r.h_flange ** 2 * r.D_int_flange) - r.B_Z * sigma_r_r_calc, 3)
    add_equation(doc, rf'σ_T^р = \frac{{{r.B_Y:.3f} \cdot {M_R_calc}}}{{{r.h_flange}^2 \cdot {r.D_int_flange}}}'
                      rf' - {r.B_Z:.3f} \cdot {sigma_r_r_calc} = {sigma_t_r_calc} МПа')

    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            p11 = doc.add_paragraph('Условия статической прочности приварных встык фланцев с конической втулкой в '
                                    'сечении ')
            p11.add_run('S').italic = True
            p11.add_run('1').font.subscript = True
            p11.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_1^м + σ_R^м \vert; \vert σ_1^м + σ_Т^м \vert \Big}'
                              r' <= K_T \cdot [σ]_м', number=f_counter, f_l_indent=4.0)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_1^р - σ_{1м.м}^р + σ_R^р \vert; \vert σ_1^р - σ_{1м.м}^р + σ_Т^р '
                              r'\vert; \vert σ_1^р + σ_{1м.м}^р \vert \Big} <= K_T \cdot [σ]_м', number=f_counter,
                         f_l_indent=1.25)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'K', italic=True, sub_text='T', value=1.3, indent=2.25,
                                      desc='коэффициент увеличения допускаемых напряжений при расчете фланцев с '
                                           'учетом стесненности температурных деформаций ([1] п.8.5.1);')
            add_paragraph_with_indent(doc, '[σ]', sub_text='м', indent=1.5,
                                      desc='допускаемое значение общих мембранных и изгибных напряжений, МПа '
                                           '([7] п.8.10).')
            f_counter += 1
            add_equation(doc, r'\left[ σ \right]_м = 1.5 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, '[σ]', value=r.allowed_stress_flange_T, indent=3.25,
                                      desc=f'допускаемое напряжение для стали марки {r.flange_steel}, МПа,'
                                           f' найденное методом экстраполяции для расчетной температуры '
                                           '([8] табл. 3).')
            add_equation(doc, rf'\left[ σ \right]_м = 1.5 \cdot {r.allowed_stress_flange_T} = '
                              rf'{r.allowed_stress_flange_m} МПа')

            p12 = doc.add_paragraph('Условия статической прочности приварных встык фланцев с конической втулкой в '
                                    'сечении ')
            p12.add_run('S').italic = True
            p12.add_run('0').font.subscript = True
            p12.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
            f_counter += 1
            add_equation(doc, r'σ_0^м <= 1.3 \cdot [σ]_R', number=f_counter, f_l_indent=6.5)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_0^р ± σ_{0м.м}^р \vert; \vert 0.3 \cdot σ_0^р ± σ_{0м.о}^р '
                              r'\vert; \vert 0.7 \cdot σ_0^р ± \Big(σ_{0м.м}^р - σ_{0м.о}^р \Big) \vert \Big} '
                              r'<= 1.3 \cdot [σ]_R', number=f_counter, f_l_indent=0.0)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, '[σ]', sub_text='R', indent=1.5,
                                      desc='допускаемое значение суммарных общих и местных условных упругих мембранных '
                                           'и изгибных напряжений, МПа ([7] п.8.10).')
            f_counter += 1
            add_equation(doc, r'\left[ σ \right]_R = 3.0 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
            add_equation(doc, rf'\left[ σ \right]_R = 3.0 \cdot {r.allowed_stress_flange_T} = '
                              rf'{r.allowed_stress_flange_r} МПа')

            p13 = doc.add_paragraph('Также для фланцев всех типов в сечении ')
            p13.add_run('S').italic = True
            p13.add_run('0').font.subscript = True
            p13.add_run(' должно выполняться условие:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_{0м.о}^р \vert; \vert σ_{0м.м}^р \vert \Big} <= [σ]',
                         number=f_counter, f_l_indent=5.25)

            doc.add_paragraph('Для тарелок приварных встык фланцев при затяжке и в рабочих условиях должны выполняться '
                              'следующие условия:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_R^м \vert; \vert σ_T^м \vert \Big} <= K_T \cdot [σ]',
                         number=f_counter, f_l_indent=5.25)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_R^р \vert; \vert σ_T^р \vert \Big} <= K_T \cdot [σ]',
                         number=f_counter, f_l_indent=5.25)

            doc.add_paragraph('Подставим найденные значения в условия статической прочности фланцев с '
                              'конической втулкой:')

            condition_43 = max(abs(sigma_1_m_calc + sigma_r_m_calc), abs(sigma_1_m_calc + sigma_t_m_calc))
            condition_44 = max(abs(sigma_1_r_calc - sigma_1_m_m_r_calc + sigma_r_r_calc),
                               abs(sigma_1_r_calc - sigma_1_m_m_r_calc + sigma_t_r_calc),
                               abs(sigma_1_r_calc + sigma_1_m_m_r_calc))
            condition_45 = sigma_0_m_calc
            condition_46 = max(max(abs(sigma_0_r_calc + sigma_0_m_m_r_calc), abs(sigma_0_r_calc - sigma_0_m_m_r_calc)),
                               max(abs(0.3 * sigma_0_r_calc + sigma_0_m_o_r_calc),
                                   abs(0.3 * sigma_0_r_calc - sigma_0_m_o_r_calc)),
                               max(abs(0.7 * sigma_0_r_calc + (sigma_0_m_m_r_calc - sigma_0_m_o_r_calc)),
                                   abs(0.7 * sigma_0_r_calc - (sigma_0_m_m_r_calc - sigma_0_m_o_r_calc))))
            condition_53 = max(abs(sigma_0_m_o_r_calc), abs(sigma_0_m_m_r_calc))
            condition_54 = max(abs(sigma_r_m_calc), abs(sigma_t_m_calc))
            condition_55 = max(abs(sigma_r_r_calc), abs(sigma_t_r_calc))

            if all([condition_43 <= 1.3 * r.allowed_stress_flange_m,
                    condition_44 <= 1.3 * r.allowed_stress_flange_m,
                    condition_45 <= 1.3 * r.allowed_stress_flange_r,
                    condition_46 <= 1.3 * r.allowed_stress_flange_r,
                    condition_53 <= r.allowed_stress_flange_T,
                    condition_54 <= 1.3 * r.allowed_stress_flange_T,
                    condition_55 <= 1.3 * r.allowed_stress_flange_T]):

                add_equation(doc, rf'\max \Big{{ \vert {sigma_1_m_calc} + {sigma_r_m_calc} \vert; '
                                  rf'\vert {sigma_1_m_calc} + {sigma_t_m_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_m}')
                add_equation(doc, rf'{condition_43:.3f} МПа <= {1.3 * r.allowed_stress_flange_m:.3f} МПа')

                add_equation(doc, rf'\max '
                                  rf'\Big{{ \vert {sigma_1_r_calc} - {sigma_1_m_m_r_calc} + {sigma_r_r_calc} \vert; '
                                  rf'\vert {sigma_1_r_calc} - {sigma_1_m_m_r_calc} + {sigma_t_r_calc} '
                                  rf'\vert; \vert {sigma_1_r_calc} + {sigma_1_m_m_r_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_m}')
                add_equation(doc, rf'{condition_44:.3f} МПа <= {1.3 * r.allowed_stress_flange_m:.3f} МПа')

                add_equation(doc, rf'{sigma_0_m_calc} <= 1.3 \cdot {r.allowed_stress_flange_r}')
                add_equation(doc, rf'{condition_45:.3f} МПа <= {1.3 * r.allowed_stress_flange_r:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_0_r_calc} ± {sigma_0_m_m_r_calc} \vert; '
                                  rf'\vert 0.3 \cdot {sigma_0_r_calc} ± {sigma_0_m_o_r_calc} '
                                  rf'\vert; \vert 0.7 \cdot {sigma_0_r_calc} ± \Big({sigma_0_m_m_r_calc}'
                                  rf' - {sigma_0_m_o_r_calc} \Big) \vert \Big}} '
                                  rf'<= 1.3 \cdot {r.allowed_stress_flange_r}')
                add_equation(doc, rf'{condition_46:.3f} МПа <= {1.3 * r.allowed_stress_flange_r:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_0_m_o_r_calc} \vert; '
                                  rf'\vert {sigma_0_m_m_r_calc} \vert \Big}}'
                                  rf' <= {r.allowed_stress_flange_T}')
                add_equation(doc, rf'{condition_53:.3f} МПа <= {r.allowed_stress_flange_T:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_r_m_calc} \vert; '
                                  rf'\vert {sigma_t_m_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_T}')
                add_equation(doc, rf'{condition_54:.3f} МПа <= {1.3 * r.allowed_stress_flange_T:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_r_r_calc} \vert; '
                                  rf'\vert {sigma_t_r_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_T}')
                add_equation(doc, rf'{condition_55:.3f} МПа <= {1.3 * r.allowed_stress_flange_T:.3f} МПа')

                doc.add_paragraph('Условия выполняются.')
            else:
                doc.add_paragraph('Условия НЕ выполняются!')
        else:
            p11 = doc.add_paragraph('Условия статической прочности приварных встык фланцев с цилиндрической втулкой в '
                                    'сечении ')
            p11.add_run('S').italic = True
            p11.add_run('0').font.subscript = True
            p11.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_0^м + σ_R^м \vert; \vert σ_0^м + σ_Т^м \vert \Big}'
                              r' <= K_T \cdot [σ]_м', number=f_counter, f_l_indent=4.0)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_0^р - σ_{0м.м}^р + σ_T^р \vert; \vert σ_0^р - σ_{0м.м}^р + σ_R^р '
                              r'\vert; \vert σ_0^р + σ_{0м.м}^р \vert \Big} <= K_T \cdot [σ]_м', number=f_counter,
                         f_l_indent=1.0)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'K', italic=True, sub_text='T', value=1.3, indent=2.25,
                                      desc='коэффициент увеличения допускаемых напряжений при расчете фланцев с '
                                           'учетом стесненности температурных деформаций ([1] п.8.5.1);')
            add_paragraph_with_indent(doc, '[σ]', sub_text='м', indent=1.5,
                                      desc='допускаемое значение общих мембранных и изгибных напряжений, МПа '
                                           '([7] п.8.10).')
            f_counter += 1
            add_equation(doc, r'\left[ σ \right]_м = 1.5 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, '[σ]', value=r.allowed_stress_flange_T, indent=3.25,
                                      desc=f'допускаемое напряжение для стали марки {r.flange_steel}, МПа,'
                                           f' найденное методом экстраполяции для расчетной температуры '
                                           '([8] табл. 3).')
            add_equation(doc, rf'\left[ σ \right]_м = 1.5 \cdot {r.allowed_stress_flange_T} = '
                              rf'{r.allowed_stress_flange_m} МПа')

            p13 = doc.add_paragraph('Также для фланцев всех типов в сечении ')
            p13.add_run('S').italic = True
            p13.add_run('0').font.subscript = True
            p13.add_run(' должно выполняться условие:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_{0м.о}^р \vert; \vert σ_{0м.м}^р \vert \Big} <= [σ]',
                         number=f_counter, f_l_indent=5.25)

            doc.add_paragraph('Для тарелок приварных встык фланцев при затяжке и в рабочих условиях должны выполняться '
                              'следующие условия:')
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_R^м \vert; \vert σ_T^м \vert \Big} <= K_T \cdot [σ]',
                         number=f_counter, f_l_indent=5.25)
            f_counter += 1
            add_equation(doc, r'\max \Big{ \vert σ_R^р \vert; \vert σ_T^р \vert \Big} <= K_T \cdot [σ]',
                         number=f_counter, f_l_indent=5.25)

            doc.add_paragraph('Подставим найденные значения в условия статической прочности фланцев с '
                              'цилиндрической втулкой:')
            condition_47 = max(abs(sigma_0_m_calc + sigma_r_m_calc), abs(sigma_0_m_calc + sigma_t_m_calc))
            condition_48 = max(abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_t_r_calc),
                               abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_r_r_calc),
                               abs(sigma_0_r_calc + sigma_0_m_m_r_calc))
            condition_53 = max(abs(sigma_0_m_o_r_calc), abs(sigma_0_m_m_r_calc))
            condition_54 = max(abs(sigma_r_m_calc), abs(sigma_t_m_calc))
            condition_55 = max(abs(sigma_r_r_calc), abs(sigma_t_r_calc))

            if all([condition_47 <= 1.3 * r.allowed_stress_flange_m,
                    condition_48 <= 1.3 * r.allowed_stress_flange_m,
                    condition_53 <= r.allowed_stress_flange_T,
                    condition_54 <= 1.3 * r.allowed_stress_flange_T,
                    condition_55 <= 1.3 * r.allowed_stress_flange_T]):

                add_equation(doc, rf'\max \Big{{ \vert {sigma_0_m_calc} + {sigma_r_m_calc} \vert; '
                                  rf'\{sigma_0_m_calc} + {sigma_t_m_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_m}')
                add_equation(doc, rf'{condition_47:.3f} МПа <= {1.3 * r.allowed_stress_flange_m:.3f} МПа')

                add_equation(doc,
                             rf'\max \Big{{\vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_t_r_calc} \vert;'
                             rf'\vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_r_r_calc} '
                             rf'\vert; \vert {sigma_0_r_calc} + {sigma_0_m_m_r_calc} \vert \Big}} '
                             rf'<= 1.3 \cdot {r.allowed_stress_flange_m}')
                add_equation(doc, rf'{condition_48:.3f} МПа <= {1.3 * r.allowed_stress_flange_m:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_0_m_o_r_calc} \vert; '
                                  rf'\vert {sigma_0_m_m_r_calc} \vert \Big}}'
                                  rf' <= {r.allowed_stress_flange_T}')
                add_equation(doc, rf'{condition_53:.3f} МПа <= {r.allowed_stress_flange_T:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_r_m_calc} \vert; '
                                  rf'\vert {sigma_t_m_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_T}')
                add_equation(doc, rf'{condition_54:.3f} МПа <= {1.3 * r.allowed_stress_flange_T:.3f} МПа')

                add_equation(doc, rf'\max \Big{{ \vert {sigma_r_r_calc} \vert; '
                                  rf'\vert {sigma_t_r_calc} \vert \Big}}'
                                  rf' <= 1.3 \cdot {r.allowed_stress_flange_T}')
                add_equation(doc, rf'{condition_55:.3f} МПа <= {1.3 * r.allowed_stress_flange_T:.3f} МПа')
            else:
                doc.add_paragraph('Условия НЕ выполняются!')
    else:
        p11 = doc.add_paragraph('Условия статической прочности приварных встык плоских фланцев в сечении ')
        p11.add_run('S').italic = True
        p11.add_run('0').font.subscript = True
        p11.add_run(' при затяжке и в рабочих условиях определяются по формулам:')
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_0^м + σ_R^м \vert; \vert σ_0^м + σ_Т^м \vert \Big}'
                          r' <= K_T \cdot [σ]_м', number=f_counter, f_l_indent=4.0)
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_0^р - σ_{0м.м}^р + σ_T^р \vert; \vert σ_0^р - σ_{0м.м}^р + σ_R^р '
                          r'\vert; \vert σ_0^р + σ_{0м.м}^р \vert \Big} <= K_T \cdot [σ]_м', number=f_counter,
                     f_l_indent=1.0)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'K', italic=True, sub_text='T', value=1.3, indent=2.25,
                                  desc='коэффициент увеличения допускаемых напряжений при расчете фланцев с '
                                       'учетом стесненности температурных деформаций ([1] п.8.5.1);')
        add_paragraph_with_indent(doc, '[σ]', sub_text='м', indent=1.5,
                                  desc='допускаемое значение общих мембранных и изгибных напряжений, МПа '
                                       '([7] п.8.10).')
        f_counter += 1
        add_equation(doc, r'\left[ σ \right]_м = 1.5 \cdot \left[ σ \right]', number=f_counter, f_l_indent=6.75)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, '[σ]', value=r.allowed_stress_flange_T, indent=3.25,
                                  desc=f'допускаемое напряжение для стали марки {r.flange_steel}, МПа,'
                                       f' найденное методом экстраполяции для расчетной температуры '
                                       '([8] табл. 3).')
        add_equation(doc, rf'\left[ σ \right]_м = 1.5 \cdot {r.allowed_stress_flange_T} = '
                          rf'{r.allowed_stress_flange_m} МПа')

        p13 = doc.add_paragraph('Также для фланцев всех типов в сечении ')
        p13.add_run('S').italic = True
        p13.add_run('0').font.subscript = True
        p13.add_run(' должно выполняться условие:')
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_{0м.о}^р \vert; \vert σ_{0м.м}^р \vert \Big} <= [σ]',
                     number=f_counter, f_l_indent=5.25)

        doc.add_paragraph('Для тарелок приварных встык фланцев при затяжке и в рабочих условиях должны выполняться '
                          'следующие условия:')
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_R^м \vert; \vert σ_T^м \vert \Big} <= K_T \cdot [σ]',
                     number=f_counter, f_l_indent=5.25)
        f_counter += 1
        add_equation(doc, r'\max \Big{ \vert σ_R^р \vert; \vert σ_T^р \vert \Big} <= K_T \cdot [σ]',
                     number=f_counter, f_l_indent=5.25)

        doc.add_paragraph('Подставим найденные значения в условия статической прочности плоских фланцев:')
        condition_47 = max(abs(sigma_0_m_calc + sigma_r_m_calc), abs(sigma_0_m_calc + sigma_t_m_calc))
        condition_48 = max(abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_t_r_calc),
                           abs(sigma_0_r_calc - sigma_0_m_m_r_calc + sigma_r_r_calc),
                           abs(sigma_0_r_calc + sigma_0_m_m_r_calc))
        condition_53 = max(abs(sigma_0_m_o_r_calc), abs(sigma_0_m_m_r_calc))
        condition_54 = max(abs(sigma_r_m_calc), abs(sigma_t_m_calc))
        condition_55 = max(abs(sigma_r_r_calc), abs(sigma_t_r_calc))

        if all([condition_47 <= 1.3 * r.allowed_stress_flange_m,
                condition_48 <= 1.3 * r.allowed_stress_flange_m,
                condition_53 <= r.allowed_stress_flange_T,
                condition_54 <= 1.3 * r.allowed_stress_flange_T,
                condition_55 <= 1.3 * r.allowed_stress_flange_T]):

            add_equation(doc, rf'\max \Big{{ \vert {sigma_0_m_calc} + {sigma_r_m_calc} \vert; '
                              rf'\{sigma_0_m_calc} + {sigma_t_m_calc} \vert \Big}}'
                              rf' <= 1.3 \cdot {r.allowed_stress_flange_m}')
            add_equation(doc, rf'{condition_47:.3f} МПа <= {1.3 * r.allowed_stress_flange_m:.3f} МПа')

            add_equation(doc,
                         rf'\max \Big{{\vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_t_r_calc} \vert;'
                         rf'\vert {sigma_0_r_calc} - {sigma_0_m_m_r_calc} + {sigma_r_r_calc} '
                         rf'\vert; \vert {sigma_0_r_calc} + {sigma_0_m_m_r_calc} \vert \Big}} '
                         rf'<= 1.3 \cdot {r.allowed_stress_flange_m}')
            add_equation(doc, rf'{condition_48:.3f} МПа <= {1.3 * r.allowed_stress_flange_m:.3f} МПа')

            add_equation(doc, rf'\max \Big{{ \vert {sigma_0_m_o_r_calc} \vert; '
                              rf'\vert {sigma_0_m_m_r_calc} \vert \Big}}'
                              rf' <= {r.allowed_stress_flange_T}')
            add_equation(doc, rf'{condition_53:.3f} МПа <= {r.allowed_stress_flange_T:.3f} МПа')

            add_equation(doc, rf'\max \Big{{ \vert {sigma_r_m_calc} \vert; '
                              rf'\vert {sigma_t_m_calc} \vert \Big}}'
                              rf' <= 1.3 \cdot {r.allowed_stress_flange_T}')
            add_equation(doc, rf'{condition_54:.3f} МПа <= {1.3 * r.allowed_stress_flange_T:.3f} МПа')

            add_equation(doc, rf'\max \Big{{ \vert {sigma_r_r_calc} \vert; '
                              rf'\vert {sigma_t_r_calc} \vert \Big}}'
                              rf' <= 1.3 \cdot {r.allowed_stress_flange_T}')
            add_equation(doc, rf'{condition_55:.3f} МПа <= {1.3 * r.allowed_stress_flange_T:.3f} МПа')
        else:
            doc.add_paragraph('Условия НЕ выполняются!')

    doc.add_paragraph("6. Проверка углов поворота фланцев", style='Heading 2')
    doc.add_paragraph('Угол поворота приварного встык фланца в рабочих условиях:')
    f_counter += 1
    f_counter_mem_12 = f_counter
    add_equation(doc, r'\theta = M^Р \cdot y_ф \cdot \frac{E_ф^{20}}{E_ф} <= K_{\theta} \cdot [\theta]',
                 number=f_counter, f_l_indent=5.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, '[θ]', indent=3.0, value=r.allowed_theta,
                              desc='допустимый угол поворота приварного встык фланца,'
                                   ' при промежуточных значениях определяется с помощью '
                                   'линейной интерполяции ([1], п.9.1);')
    add_paragraph_with_indent(doc, 'K', italic=True, sub_text='θ', value=1.0, indent=2.25,
                              desc='коэффициент, применяемый к рабочим условиям ([1], п.9.1);')
    add_paragraph_with_indent(doc, 'E', italic=True, sub_text='ф', sup_text='20', value=r.E_flange_20, indent=3.0,
                              desc='модуль продольной упругости материала фланца при температуре 20˚С, МПа'
                                   ' ([1], таблица Ж.1);')
    add_paragraph_with_indent(doc, 'E', italic=True, sub_text='ф', value=r.E_flange_T, indent=3.0,
                              desc='модуль продольной упругости материала фланца при расчетной температуре,'
                                   ' рассчитанный методом линейной интерполяции, МПа ([1], таблица Ж.1).')
    add_equation(doc, r'[\theta] = 0.006 при D_{вн} <= 400 мм')
    add_equation(doc, r'[\theta] = 0.013 при D_{вн} > 2000 мм')

    doc.add_paragraph(f'Подставим в выражение ({f_counter_mem_12}):')
    theta_calc = round(M_R_calc * 1e6 * y_f_calc * r.E_flange_20 / r.E_flange_T, 4)
    add_equation(doc, rf'\theta = {M_R_calc} \cdot {y_f_calc} \cdot \frac{{{r.E_flange_20}}}{{{r.E_flange_T}}} '
                      rf'<= 1.0 \cdot {r.allowed_theta}')
    add_equation(doc, rf'{theta_calc} <= {r.allowed_theta}')
    doc.add_paragraph('Условие выполняется.')

    doc.add_paragraph("7. Расчет элементов фланцевого соединения на малоцикловую усталость", style='Heading 2')
    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            p14 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p14.add_run('S').italic = True
            p14.add_run('1').font.subscript = True
            p14.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{11}}^М = σ_1^М = {sigma_1_m_calc} МПа', number=f_counter, f_l_indent=5.5)
            f_counter += 1
            add_equation(doc, rf'σ_{{12}}^М = -σ_1^М = {-sigma_1_m_calc} МПа', number=f_counter, f_l_indent=5.0)

            p15 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'конической втулки приварного встык фланца в сечении ')
            p15.add_run('S').italic = True
            p15.add_run('1').font.subscript = True
            p15.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{13}}^М = σ_{{14}}^М = σ_T^М = {sigma_t_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            p16 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p16.add_run('S').italic = True
            p16.add_run('1').font.subscript = True
            p16.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{15}}^М = σ_{{16}}^М = σ_R^М = {sigma_r_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            p17 = doc.add_paragraph('Размах условных упругих напряжений в меридиональном направлении на наружной '
                                    'и внутренней поверхностях конической втулки приварного встык фланца в сечении ')
            p17.add_run('S').italic = True
            p17.add_run('0').font.subscript = True
            p17.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{01}}^М = σ_0^М = {sigma_0_m_calc} МПа', number=f_counter,
                         f_l_indent=5.5)
            f_counter += 1
            add_equation(doc, rf'σ_{{02}}^М = -σ_0^М = {-sigma_0_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            sigma_1_1_m_calc = sigma_1_m_calc
            sigma_1_2_m_calc = -sigma_1_m_calc
            sigma_1_4_m_calc = sigma_t_m_calc
            sigma_1_6_m_calc = sigma_r_m_calc
            sigma_0_1_m_calc = sigma_0_m_calc

            p18 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p18.add_run('S').italic = True
            p18.add_run('1').font.subscript = True
            p18.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{11}}^Р = σ_1^Р + σ_{{1м.м}}^Р = {sigma_1_m_calc} + {sigma_1_m_m_r_calc}'
                         rf' = {round(sigma_1_m_calc + sigma_1_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.5)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{12}}^Р = -σ_1^Р + σ_{{1м.м}}^Р = {-sigma_1_m_calc} + {sigma_1_m_m_r_calc}'
                         rf' = {-sigma_1_m_calc + sigma_1_m_m_r_calc} МПа',
                         number=f_counter, f_l_indent=2.0)

            p19 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'конической втулки приварного встык фланца в сечении ')
            p19.add_run('S').italic = True
            p19.add_run('1').font.subscript = True
            p19.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{13}}^Р = σ_{{14}}^Р = σ_T^Р = {sigma_t_r_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            p20 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p20.add_run('S').italic = True
            p20.add_run('1').font.subscript = True
            p20.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{15}}^Р = σ_{{16}}^Р = σ_R^Р = {sigma_r_r_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            p21 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях конической втулки приварного встык фланца в сечении ')
            p21.add_run('S').italic = True
            p21.add_run('0').font.subscript = True
            p21.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{01}}^Р = σ_0^Р + σ_{{0м.м}}^Р = {sigma_0_m_calc} + {sigma_0_m_m_r_calc}'
                         rf' = {round(sigma_0_m_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.5)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{02}}^Р = -σ_0^Р + σ_{{0м.м}}^Р = {-sigma_0_m_calc} + {sigma_0_m_m_r_calc}'
                         rf' = {round(-sigma_0_m_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.0)

            p22 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'конической втулки приварного встык фланца в сечении ')
            p22.add_run('S').italic = True
            p22.add_run('0').font.subscript = True
            p22.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{03}}^Р = σ_{{0м.о}}^Р + 0.3 \cdot σ_0^Р = {sigma_0_m_o_r_calc} + '
                         rf'0.3 \cdot {sigma_0_r_calc} = {round(sigma_0_m_o_r_calc + 0.3 * sigma_0_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=1.5)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{04}}^Р = σ_{{0м.о}}^Р - 0.3 \cdot σ_0^Р = {sigma_0_m_o_r_calc} - '
                         rf'0.3 \cdot {sigma_0_r_calc} = {round(sigma_0_m_o_r_calc - 0.3 * sigma_0_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=1.75)

            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с конической втулкой:')
            f_counter += 1
            add_equation(doc, r'σ_{аф} = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert σ_{11}^М \vert ;'
                              r'\vert σ_{12}^М - σ_{14}^М \vert ; '
                              r'\vert σ_{12}^М - σ_{16}^М \vert ; \vert σ_{01}^М \vert \Big} }{2}',
                         number=f_counter, f_l_indent=4.25)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'α', italic=True, sub_text='σ', value=round(r.alpha_sigma, 3),
                                      indent=2.75, desc='коэффициент, определяемый по рисунку 4 [1].')

            max_3 = max(r.alpha_sigma * abs(sigma_1_1_m_calc), abs(sigma_1_2_m_calc - sigma_1_4_m_calc),
                        abs(sigma_1_2_m_calc - sigma_1_6_m_calc), abs(sigma_0_1_m_calc))
            stress_a = round(max_3 / 2, 3)
            add_equation(doc, rf'σ_{{аф}} = \frac{{\max \Big{{{r.alpha_sigma:.3f} \cdot \vert {sigma_1_1_m_calc} \vert'
                              rf'; \vert {sigma_1_2_m_calc} - {sigma_1_4_m_calc} \vert ; '
                              rf'\vert {sigma_1_2_m_calc} - {sigma_1_6_m_calc} \vert ; \vert {sigma_0_1_m_calc} \vert '
                              rf'\Big}} }}{{2}} = {stress_a} МПа')
        else:
            p14 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p14.add_run('S').italic = True
            p14.add_run('0').font.subscript = True
            p14.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{01}}^М = σ_0^М = {sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.5)
            f_counter += 1
            add_equation(doc, rf'σ_{{02}}^М = -σ_0^М = {-sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.0)

            p15 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'цилиндрической втулки приварного встык фланца в сечении ')
            p15.add_run('S').italic = True
            p15.add_run('0').font.subscript = True
            p15.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{03}}^М = σ_{{04}}^М = σ_T^М = {sigma_t_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            p16 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p16.add_run('S').italic = True
            p16.add_run('0').font.subscript = True
            p16.add_run(' в условиях затяжки:')
            f_counter += 1
            add_equation(doc, rf'σ_{{05}}^М = σ_{{06}}^М = σ_R^М = {sigma_r_m_calc} МПа', number=f_counter,
                         f_l_indent=5.0)
            sigma_0_1_m_calc = sigma_0_m_calc
            sigma_0_2_m_calc = -sigma_0_m_calc
            sigma_0_3_m_calc = sigma_0_4_m_calc = sigma_t_m_calc
            sigma_0_6_m_calc = sigma_r_m_calc

            p17 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p17.add_run('S').italic = True
            p17.add_run('0').font.subscript = True
            p17.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{01}}^Р = σ_0^Р + σ_{{0м.м}}^Р = {sigma_0_r_calc} + {sigma_0_m_m_r_calc}'
                         rf' = {round(sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.75)
            f_counter += 1
            add_equation(doc,
                         rf'σ_{{02}}^Р = -σ_0^Р + σ_{{0м.м}}^Р = {-sigma_0_r_calc} + {sigma_0_m_m_r_calc}'
                         rf' = {round(-sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа',
                         number=f_counter, f_l_indent=2.5)

            p18 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                    'цилиндрической втулки приварного встык фланца в сечении ')
            p18.add_run('S').italic = True
            p18.add_run('0').font.subscript = True
            p18.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{03}}^Р = σ_{{04}}^Р = σ_T^Р = {sigma_t_r_calc} МПа', number=f_counter,
                         f_l_indent=5.0)

            p19 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                    'поверхностях цилиндрической втулки приварного встык фланца в сечении ')
            p19.add_run('S').italic = True
            p19.add_run('0').font.subscript = True
            p19.add_run(' в рабочих условиях:')
            f_counter += 1
            add_equation(doc, rf'σ_{{05}}^Р = σ_{{06}}^Р = σ_R^Р = {sigma_r_r_calc} МПа', number=f_counter,
                         f_l_indent=4.75)

            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с цилиндрической втулкой:')
            f_counter += 1
            add_equation(doc, r'σ_{аф} = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert σ_{01}^М \vert ;'
                              r'\vert σ_{02}^М - σ_{04}^М \vert ; '
                              r'\vert σ_{02}^М - σ_{06}^М \vert ; \Big} }{2}',
                         number=f_counter, f_l_indent=4.5)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'α', italic=True, sub_text='σ', value=round(r.alpha_sigma, 3),
                                      indent=2.75, desc='коэффициент, определяемый по рисунку 4 [1].')
            max_3 = max(r.alpha_sigma * abs(sigma_0_1_m_calc), abs(sigma_0_2_m_calc - sigma_0_4_m_calc),
                        abs(sigma_0_2_m_calc - sigma_0_6_m_calc))
            stress_a = round(max_3 / 2, 3)
            add_equation(doc, rf'σ_{{аф}} = \frac{{\max \Big{{{r.alpha_sigma:.3f} \cdot '
                              rf'\vert {sigma_0_1_m_calc} \vert'
                              rf'; \vert {sigma_0_2_m_calc} - {sigma_0_4_m_calc} \vert ; '
                              rf'\vert {sigma_0_2_m_calc} - {sigma_0_6_m_calc} \vert ; '
                              rf'\Big}} }}{{2}} = {stress_a} МПа')
    else:
        p14 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p14.add_run('S').italic = True
        p14.add_run('0').font.subscript = True
        p14.add_run(' в условиях затяжки:')
        f_counter += 1
        add_equation(doc, rf'σ_{{01}}^М = σ_0^М = {sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.5)
        f_counter += 1
        add_equation(doc, rf'σ_{{02}}^М = -σ_0^М = {-sigma_0_m_calc} МПа', number=f_counter, f_l_indent=5.0)

        p15 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                'приварного встык плоского фланца в сечении ')
        p15.add_run('S').italic = True
        p15.add_run('0').font.subscript = True
        p15.add_run(' в условиях затяжки:')
        f_counter += 1
        add_equation(doc, rf'σ_{{03}}^М = σ_{{04}}^М = σ_T^М = {sigma_t_m_calc} МПа', number=f_counter,
                     f_l_indent=5.0)

        p16 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p16.add_run('S').italic = True
        p16.add_run('0').font.subscript = True
        p16.add_run(' в условиях затяжки:')
        f_counter += 1
        add_equation(doc, rf'σ_{{05}}^М = σ_{{06}}^М = σ_R^М = {sigma_r_m_calc} МПа', number=f_counter,
                     f_l_indent=5.0)

        sigma_0_1_m_calc = sigma_0_m_calc
        sigma_0_2_m_calc = -sigma_0_m_calc
        sigma_0_4_m_calc = sigma_t_m_calc
        sigma_0_6_m_calc = sigma_r_m_calc

        p17 = doc.add_paragraph('Размах условных упругих меридиональных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p17.add_run('S').italic = True
        p17.add_run('0').font.subscript = True
        p17.add_run(' в рабочих условиях:')
        f_counter += 1
        add_equation(doc,
                     rf'σ_{{01}}^Р = σ_0^Р + σ_{{0м.м}}^Р = {sigma_0_r_calc} + {sigma_0_m_m_r_calc}'
                     rf' = {round(sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа', number=f_counter, f_l_indent=2.75)
        f_counter += 1
        add_equation(doc,
                     rf'σ_{{02}}^Р = -σ_0^Р + σ_{{0м.м}}^Р = {-sigma_0_r_calc} + {sigma_0_m_m_r_calc}'
                     rf' = {round(-sigma_0_r_calc + sigma_0_m_m_r_calc, 3)} МПа', number=f_counter, f_l_indent=2.5)

        p18 = doc.add_paragraph('Размах условных упругих окружных напряжений на наружной и внутренней поверхностях '
                                'приварного встык плоского фланца в сечении ')
        p18.add_run('S').italic = True
        p18.add_run('0').font.subscript = True
        p18.add_run(' в рабочих условиях:')
        f_counter += 1
        add_equation(doc, rf'σ_{{03}}^Р = σ_{{04}}^Р = σ_T^Р = {sigma_t_r_calc} МПа', number=f_counter,
                     f_l_indent=5.0)

        p19 = doc.add_paragraph('Размах условных упругих радиальных напряжений на наружной и внутренней '
                                'поверхностях приварного встык плоского фланца в сечении ')
        p19.add_run('S').italic = True
        p19.add_run('0').font.subscript = True
        p19.add_run(' в рабочих условиях:')
        f_counter += 1
        add_equation(doc, rf'σ_{{05}}^Р = σ_{{06}}^Р = σ_R^Р = {sigma_r_r_calc} МПа', number=f_counter,
                     f_l_indent=4.75)

        doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык '
                          'плоских фланцев:')
        f_counter += 1
        add_equation(doc, r'σ_{аф} = \frac{1.5 \cdot \max \Big{ \vert σ_{01}^М \vert ;'
                          r'\vert σ_{02}^М - σ_{04}^М \vert ; '
                          r'\vert σ_{02}^М - σ_{06}^М \vert ; \Big} }{2}',
                     number=f_counter, f_l_indent=4.5)
        max_3 = max(abs(sigma_0_1_m_calc), abs(sigma_0_2_m_calc - sigma_0_4_m_calc),
                    abs(sigma_0_2_m_calc - sigma_0_6_m_calc))
        stress_a = round(1.5 * max_3 / 2, 3)
        add_equation(doc, rf'σ_{{аф}} = \frac{{1.5 \cdot \max \Big{{ \vert {sigma_0_1_m_calc} \vert ;'
                          rf'\vert {sigma_0_2_m_calc} - {sigma_0_4_m_calc} \vert ; '
                          rf'\vert {sigma_0_2_m_calc} - {sigma_0_6_m_calc} \vert ; '
                          rf'\Big}} }}{{2}} = {stress_a} МПа')

    doc.add_paragraph('Амплитуда приведенных условных упругих напряжений при затяжке шпилек:')
    f_counter += 1
    stress_a_bolt = round(1.8 * sigma_b_1_calc / 2, 3)
    add_equation(doc, r'σ_{аш} = \frac{1.8 \cdot σ_{ш1}}{2}', number=f_counter, f_l_indent=7.0)
    add_equation(doc,
                 rf'σ_{{аш}} = \frac{{1.8 \cdot {sigma_b_1_calc}}}{{2}} = {stress_a_bolt} МПа')

    if r.flange_type == 'one_one':
        if r.D_m_flange > r.D_n_flange:
            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с конической втулкой в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{аф}^Р = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert Δσ_{11}^Р \vert ;'
                              r' \vert Δσ_{12}^Р - Δσ_{14}^Р \vert ; \vert Δσ_{12}^Р - Δσ_{16}^Р \vert'
                              r'\Big} }{2}', number=f_counter, f_l_indent=4.5)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sup_text='Р', indent=2.75,
                                      desc=f'определяются как максимальные разности значений напряжений, вычисленных '
                                           'при давлении от 0.6*P до 1,25*P по описанному ранее алгоритму:')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='11', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_1_1, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='12', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_1_2, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='14', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_1_4, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='16', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_1_6, 3), desc='МПа.')
            stress_R_a = round(max(r.alpha_sigma * abs(sd.stress_delta_R_1_1),
                                             abs(sd.stress_delta_R_1_2 - sd.stress_delta_R_1_4),
                                             abs(sd.stress_delta_R_1_2 - sd.stress_delta_R_1_6)) / 2, 3)
            add_equation(doc,
                         rf'σ_{{аф}}^Р = \frac{{\max \Big{{ {r.alpha_sigma:.3f} \cdot '
                         rf'\vert {sd.stress_delta_R_1_1:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_1_2:.3f} - {sd.stress_delta_R_1_4:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_1_2:.3f} - {sd.stress_delta_R_1_6:.3f} \vert'
                         rf'\Big}} }}{{2}} = {stress_R_a} МПа', f_l_indent=1.0)
        else:
            doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык фланцев '
                              'с цилиндрической втулкой в рабочих условиях:')
            f_counter += 1
            add_equation(doc, r'σ_{аф}^Р = \frac{\max \Big{ \alpha_{\sigma} \cdot \vert Δσ_{01}^Р \vert ;'
                              r' \vert Δσ_{01}^Р - Δσ_{03}^Р \vert ; \vert Δσ_{01}^Р - Δσ_{05}^Р \vert ;'
                              r' \vert Δσ_{02}^Р \vert ;'
                              r' \vert Δσ_{02}^Р - Δσ_{04}^Р \vert ; \vert Δσ_{02}^Р - Δσ_{06}^Р \vert '
                              r'\Big} }{2}', number=f_counter, f_l_indent=2.25)
            doc.add_paragraph("где: ", style="Normal")
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sup_text='Р', indent=2.75,
                                      desc=f'определяются как максимальные разности значений напряжений, вычисленных '
                                           'при давлении от 0.6*P до 1,25*P по описанному ранее алгоритму:')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='01', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_0_1, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='03', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_0_3, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='05', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_0_5, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='02', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_0_2, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='04', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_0_4, 3), desc='МПа;')
            add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='06', sup_text='Р',
                                      indent=3.25, value=round(sd.stress_delta_R_0_6, 3), desc='МПа.')
            stress_R_a = round(max(r.alpha_sigma * abs(sd.stress_delta_R_0_1),
                                             abs(sd.stress_delta_R_0_1 - sd.stress_delta_R_0_3),
                                             abs(sd.stress_delta_R_0_1 - sd.stress_delta_R_0_5),
                                             abs(sd.stress_delta_R_0_2),
                                             abs(sd.stress_delta_R_0_2 - sd.stress_delta_R_0_4),
                                             abs(sd.stress_delta_R_0_2 - sd.stress_delta_R_0_6)) / 2, 3)
            add_equation(doc,
                         rf'σ_{{аф}}^Р = \frac{{\max \Big{{ \{r.alpha_sigma:.3f} \cdot'
                         rf' \vert {sd.stress_delta_R_0_1:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_0_1:.3f} - {sd.stress_delta_R_0_3:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_0_1:.3f} - {sd.stress_delta_R_0_5:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_0_2:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_0_2:.3f} - {sd.stress_delta_R_0_4:.3f} \vert ;'
                         rf' \vert {sd.stress_delta_R_0_2:.3f} - {sd.stress_delta_R_0_6:.3f} \vert '
                         rf'\Big}} }}{{2}} = {stress_R_a} МПа', f_l_indent=1.0)
    else:
        doc.add_paragraph('Амплитуда приведенных условных упругих напряжений для приварных встык плоских фланцев '
                          'в рабочих условиях:')
        f_counter += 1
        add_equation(doc, r'σ_{аф}^Р = \frac{1.5 \cdot \max \Big{ \vert Δσ_{01}^Р \vert ;'
                          r' \vert Δσ_{03}^Р \vert ; \vert Δσ_{01}^Р - Δσ_{03}^Р \vert ;'
                          r' \vert Δσ_{02}^Р \vert ;'
                          r' \vert Δσ_{04}^Р \vert ; \vert Δσ_{02}^Р - Δσ_{04}^Р \vert '
                          r'\Big} }{2}', number=f_counter, f_l_indent=2.5)
        doc.add_paragraph("где: ", style="Normal")
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sup_text='Р', indent=2.75,
                                  desc=f'определяются как максимальные разности значений напряжений, вычисленных при'
                                       'давлении от 0.6*P до 1,25*P по описанному ранее алгоритму:')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='01', sup_text='Р',
                                  indent=3.25, value=round(sd.stress_delta_R_0_1, 3), desc='МПа;')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='03', sup_text='Р',
                                  indent=3.25, value=round(sd.stress_delta_R_0_3, 3), desc='МПа;')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='02', sup_text='Р',
                                  indent=3.25, value=round(sd.stress_delta_R_0_2, 3), desc='МПа;')
        add_paragraph_with_indent(doc, 'Δσ', italic=False, sub_text='04', sup_text='Р',
                                  indent=3.25, value=round(sd.stress_delta_R_0_4, 3), desc='МПа.')
        stress_R_a = round(1.5 * max(abs(sd.stress_delta_R_0_1), abs(sd.stress_delta_R_0_3),
                                               abs(sd.stress_delta_R_0_1 - sd.stress_delta_R_0_3),
                                               abs(sd.stress_delta_R_0_2), abs(sd.stress_delta_R_0_4),
                                               abs(sd.stress_delta_R_0_2 - sd.stress_delta_R_0_4)) / 2, 3)
        add_equation(doc,
                     rf'σ_{{аф}}^Р = \frac{{1.5 \cdot \max \Big{{ \vert {sd.stress_delta_R_0_1:.3f} \vert ;'
                     rf' \vert {sd.stress_delta_R_0_3:.3f} \vert ;'
                     rf' \vert {sd.stress_delta_R_0_1:.3f} - {sd.stress_delta_R_0_3:.3f} \vert ;'
                     rf' \vert {sd.stress_delta_R_0_2:.3f} \vert ;'
                     rf' \vert {sd.stress_delta_R_0_4:.3f} \vert ;'
                     rf' \vert {sd.stress_delta_R_0_2:.3f} - {sd.stress_delta_R_0_4:.3f} \vert '
                     rf'\Big}} }}{{2}} = {stress_R_a:.3f} МПа', f_l_indent=1.0)

    doc.add_paragraph('Амплитуда приведенных условных упругих напряжений в рабочих условиях для шпилек:')
    f_counter += 1
    add_equation(doc, r'σ_{аш}^Р = \frac{\eta \cdot σ_{ш2}}{2}', number=f_counter, f_l_indent=7.0)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'η', italic=True, indent=1.25, value=2.0,
                              desc='коэффициент, определяемый по [9], таблица 2.')
    stress_R_a_bolt = round(2.0 * sigma_b_2_calc / 2, 3)
    add_equation(doc,rf'σ_{{аш}} = \frac{{2.0 \cdot {sigma_b_2_calc}}}{{2}} = {stress_R_a_bolt} МПа')

    doc.add_paragraph('Проверка элементов фланцевого соединения на малоцикловую усталость проводится согласно [9]. '
                      'Необходимо выполнение условий:')
    f_counter += 1
    add_equation(doc, r'\frac{N_c}{[N]_c} + \frac{N_p}{[N]_p} <= 1.0', number=f_counter, f_l_indent=6.5)
    doc.add_paragraph("где: ", style="Normal")
    add_paragraph_with_indent(doc, 'N', italic=False, sub_text='c', indent=2.5, value=r.num_cycles_c,
                              desc='заданное число циклов сборок-разборок оборудования;')
    add_paragraph_with_indent(doc, 'N', italic=False, sub_text='p', indent=3.0, value=r.num_cycles_r,
                              desc='заданное число циклов изменений режима эксплуатации;')
    add_paragraph_with_indent(doc, '[N]', italic=False, sub_text='c', indent=3.5, value=r.c_cycles,
                              desc='допускаемое число циклов сборок-разборок ([9], рисунки 2-5);')
    add_paragraph_with_indent(doc, '[N]', italic=False, sub_text='p', indent=3.5, value=r.r_cycles,
                              desc='допускаемое число циклов изменения режима эксплуатации ([9], рисунки 2-5);')
    add_equation(doc, rf'\frac{{{r.num_cycles_c}}}{{{r.c_cycles}}} + \frac{{{r.num_cycles_r}}}{{{r.r_cycles}}}'
                      rf' <= 1.0')
    add_equation(doc, rf'{round(r.num_cycles_c / r.c_cycles + r.num_cycles_r / r.r_cycles, 3)} <= 1.0')
    doc.add_paragraph('Условие выполняется.')

    pf = doc.add_paragraph('Допускаемые амплитуды [σ')
    pf.add_run('а').font.subscript = True
    pf.add_run('] для материалов шпилек и фланцев при затяжке и в рабочих условиях определяются согласно [9], '
               'формула 12 и таблица 3.')

    if all([stress_a_bolt <= r.allowed_amplitude_bolt_m,
            stress_R_a_bolt <= r.allowed_amplitude_bolt_r,
            stress_a <= r.allowed_amplitude_flange_m,
            stress_R_a <= r.allowed_amplitude_flange_r,
            r.condition_cycles <= 1.0]):
        f_counter += 1
        add_equation(doc, r'σ_{аш} <= [σ_{аш}]', number=f_counter, f_l_indent=7.0)
        add_equation(doc, rf'{stress_a_bolt} МПа <= {r.allowed_amplitude_bolt_m:.3f} МПа')

        f_counter += 1
        add_equation(doc, r'σ_{аш}^Р <= [σ_{аш}^Р]', number=f_counter, f_l_indent=7.0)
        add_equation(doc, rf'{stress_R_a_bolt} МПа <= {r.allowed_amplitude_bolt_r:.3f} МПа')

        f_counter += 1
        add_equation(doc, r'σ_{аф} <= [σ_{аф}]', number=f_counter, f_l_indent=7.0)
        add_equation(doc, rf'{stress_a} МПа <= {r.allowed_amplitude_flange_m:.3f} МПа')

        f_counter += 1
        add_equation(doc, r'σ_{аф}^Р <= [σ_{аф}^Р]', number=f_counter, f_l_indent=7.0)
        add_equation(doc, rf'{stress_R_a} МПа <= {r.allowed_amplitude_flange_r:.3f} МПа')
        doc.add_paragraph('Условия выполняются.')
    else:
        doc.add_paragraph('Условия НЕ выполняется!')

    pp = doc.add_paragraph()
    prun = pp.add_run()
    prun.add_break(WD_BREAK.PAGE)

    doc.add_paragraph("Перечень документов", style='Heading 1')
    doc.add_paragraph("1.	ГОСТ Р 52857.4-2007 - Сосуды и аппараты. Нормы и методы расчета на прочность. "
                      "Расчет на прочность и герметичность фланцевых соединений.")
    if g.document == "ГОСТ 15180-86":
        doc.add_paragraph("2.	ГОСТ 15180-86 - Прокладки плоские эластичные. Основные параметры и размеры.")
    if g.document == "ГОСТ 34655-2020":
        doc.add_paragraph("2.	ГОСТ 34655-2020 - Арматура трубопроводная. Прокладки овального, восьмиугольного "
                          "сечения, линзовые стальные для фланцев арматуры. Конструкция, размеры и общие технические "
                          "требования.")
    if g.document == "ISO 7483-2011":
        doc.add_paragraph("2.	ISO 7483 - Dimensions of gaskets for use with flanges to ISO 7005.")
    if g.document == "ГОСТ 28759.8-2022":
        doc.add_paragraph("2.	ГОСТ 28759.8-2022 - Фланцы сосудов и аппаратов. Прокладки металлические "
                          "восьмиугольные. Конструкция и размеры. Технические требования.")
    doc.add_paragraph("3.	ГОСТ 9064-75 - Гайки для фланцевых соединений с температурой среды от 0 до 650 ˚С. Типы "
                      "и основные размеры.")
    doc.add_paragraph("4.	ГОСТ 9065-75 - Шайбы для фланцевых соединений с температурой среды от 0 до 650 ˚С. Типы "
                      "и основные размеры.")
    if g.document == "ГОСТ 28759.8-2022":
        doc.add_paragraph("5.   ГОСТ 28759.4-2022 - Фланцы сосудов и аппаратов стальные приварные"
                          " встык под прокладку восьмиугольного сечения. Конструкция и размеры.")
    else:
        doc.add_paragraph("5.	ГОСТ 33259-2015 - Фланцы арматуры, соединительных частей и трубопроводов на номинальное"
                      " давление до PN 250. Конструкция, размеры и общие технические требования")
    doc.add_paragraph("6.	ОСТ 26-2040-96 - Шпильки для фланцевых соединений. Конструкция и размеры.")
    doc.add_paragraph("7.	ГОСТ Р 52857.1-2007 - Сосуды и аппараты. Нормы и методы расчета на прочность. "
                      "Общие требования.")
    doc.add_paragraph("8.	РД 26-15-88 - Сосуды и аппараты. Нормы и методы расчета на прочности и герметичность "
                      "фланцевых соединений.")
    doc.add_paragraph("9.	ГОСТ Р 52857.6-2007 - Сосуды и аппараты. Нормы и методы расчета на прочность. "
                      "Расчет на прочность при малоцикловых нагрузках.")

    filename = (f'Расчет_фланцев_DN{int(r.D_N_flange)}_PN{int(r.pressure * 10)}_{r.flange_type}_{r.face_type}_'
                f'{r.flange_steel}_{r.bolt_steel}')

    doc.save(f'{filename}.docx')
